"""Achird Local — 로컬 논문 읽기 도우미 서버.

AI 백엔드는 Claude Code CLI(`claude -p`) 서브프로세스: 구독 인증 재사용, API 키 불필요.
"""
import asyncio
import hashlib
import json
import mimetypes
import os
import re
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import threading
import time
import traceback
import urllib.error
import urllib.parse
import urllib.request
import uuid
from pathlib import Path

import uvicorn
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.exceptions import RequestValidationError
from fastapi.responses import FileResponse, JSONResponse, PlainTextResponse, Response
from fastapi.staticfiles import StaticFiles

# Windows 레지스트리에 .mjs가 없으면 octet-stream으로 서빙되어 브라우저가 ES 모듈을 거부한다
mimetypes.add_type("text/javascript", ".mjs")
mimetypes.add_type("text/javascript", ".js")
mimetypes.add_type("font/woff2", ".woff2")

ROOT = Path(__file__).resolve().parent
LIB = ROOT / "library"
LIB.mkdir(exist_ok=True)

HOST, PORT = "127.0.0.1", 8766  # 8765는 주식 대시보드(dashboard\server.py)가 점유
URL = f"http://{HOST}:{PORT}"

CLAUDE = shutil.which("claude")
CREATE_NO_WINDOW = getattr(subprocess, "CREATE_NO_WINDOW", 0)
MAX_CTX = 150_000          # 논문 전문 프롬프트 상한 (chars)
MAX_PDF_BYTES = 100 * 1024 * 1024   # PDF 업로드 상한 100MB
ID_RE = re.compile(r"^[0-9a-f]{8}$")

app = FastAPI(title="Achird Local")


@app.exception_handler(Exception)
async def unhandled(request: Request, exc: Exception):
    """HTTPException 이 아닌 예외는 기본값이 본문 없는 500 이라 원인이 통째로 사라진다 —
    화면에는 예외 타입·메시지를, 콘솔에는 traceback 을 남긴다(HTTPException 은 기본 핸들러로 통과)."""
    traceback.print_exception(exc)
    return JSONResponse({"detail": f"서버 오류: {type(exc).__name__}: {exc}"}, status_code=500)


@app.exception_handler(RequestValidationError)
async def bad_request(request: Request, exc: RequestValidationError):
    """FastAPI 기본 핸들러는 detail 을 리스트로 낸다 — 클라이언트가 문자열로 읽어
    "[object Object]" 가 뜨고 어느 필드가 왜 틀렸는지가 사라진다. 문자열로 되돌린다."""
    e = (exc.errors() or [{}])[0]
    where = ".".join(str(x) for x in (e.get("loc") or [])[1:]) or "요청"
    return JSONResponse({"detail": f"요청 형식 오류: {where} — {e.get('msg', '')}"}, status_code=422)


_ALLOWED_HOSTS = (f"{HOST}:{PORT}", f"localhost:{PORT}")


@app.middleware("http")
async def block_dns_rebinding(request: Request, call_next):
    """Host 헤더가 허용 목록이 아니면 403 — DNS 리바인딩(외부 도메인을 127.0.0.1로 resolve시켜
    브라우저가 이 로컬 서버를 두드리게 하는 공격) 차단."""
    if request.headers.get("host") not in _ALLOWED_HOSTS:
        return JSONResponse({"detail": "허용되지 않은 Host"}, status_code=403)
    return await call_next(request)


# ---------------------------------------------------------------- storage

def paper_dir(pid: str) -> Path:
    if not ID_RE.match(pid):
        raise HTTPException(400, "잘못된 논문 id")
    d = LIB / pid
    if not d.is_dir():
        raise HTTPException(404, "논문이 없습니다")
    return d


def read_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return default
    except (OSError, json.JSONDecodeError) as e:   # 있는데 못 읽음 — 조용히 넘기면 원인이 사라진다
        print(f"[read_json] {path}: {type(e).__name__}: {e}", file=sys.stderr)
        return default


def write_json(path: Path, data) -> None:
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=1), encoding="utf-8")
    tmp.replace(path)


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def existing_pdf_hashes() -> dict:
    """{sha256: pid} — 라이브러리 내 PDF 내용 해시. 업로드·임포트 내용중복 판정용.
    해시 없는 기존 논문은 1회 계산해 meta에 백필(이후 캐시)."""
    out = {}
    for d in LIB.iterdir():
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        h = meta.get("pdf_sha256")
        if not h and (d / "paper.pdf").exists():
            h = file_sha256(d / "paper.pdf")
            meta["pdf_sha256"] = h
            write_json(d / "meta.json", meta)
        if h:
            out[h] = d.name
    return out


def _no_text(d: Path) -> str:
    """본문 없음의 원인을 갈라 준다 — 파일이 있는데 read_json 이 default 로 떨어진 경우
    '아직 추출되지 않았다'는 거짓이고, 사용자는 논문을 다시 열며 헛수고한다."""
    if (d / "text.json").exists():
        return "본문 텍스트 파일을 읽지 못했습니다 — text.json 손상 또는 동기화 중"
    return "본문 텍스트가 아직 추출되지 않았습니다. 논문을 먼저 열어주세요."


def full_text(pid: str) -> str:
    d = paper_dir(pid)
    t = read_json(d / "text.json", None)
    if not t or not t.get("pages"):
        raise HTTPException(409, _no_text(d))
    return "\n\n".join(p["text"] for p in t["pages"])


# ---------------------------------------------------------------- claude runner

def ask_claude(prompt: str, system: str, timeout: int = 300, model: str = "sonnet",
               tools: str = "", cwd: str | None = None) -> str:
    """claude -p 1회 호출. 기본은 --tools ''(도구 차단)로 순수 텍스트 입출력.

    cp949 함정 때문에 text 모드를 쓰지 않고 bytes로 주고받는다.
    프롬프트는 stdin으로 전달한다(Windows 32,767자 cmdline 한계).
    참고: 옛 --safe-mode는 CLI 2.x에서 제거됨. 후속 --bare는 OAuth를 끊고 API키를 강제하므로 못 씀.
    model: 번역·요약·채팅은 sonnet(품질), 문장 정렬 같은 단순 매칭은 haiku(속도·비용).
    tools: 그림 해설·스캔 OCR처럼 이미지를 눈으로 읽어야 할 때만 "Read" 허용(그 외엔 '').
    cwd: Read 허용 시 이미지가 있는 폴더로 지정해 파일 접근을 그 논문 폴더로 한정.
    """
    if not CLAUDE:
        raise HTTPException(500, "claude CLI를 찾을 수 없습니다. `claude` 로그인 설치가 필요합니다.")
    # disableAllHooks: 사용자 전역 훅(SessionStart 동기화, SessionEnd transcript-mirror 등)이
    # 헤드리스 호출마다 상속 실행되면 초 단위 오버헤드 + 동시 호출 시 훅 취소("Hook cancelled")로
    # returncode≠0 → 모든 AI 기능이 502로 죽는다. --no-session-persistence라 transcript도 없어
    # 훅이 할 일 자체가 없다.
    cmd = [CLAUDE, "-p", "--model", model, "--output-format", "text",
           "--tools", tools, "--no-session-persistence",
           "--settings", '{"disableAllHooks":true}',
           "--system-prompt", system]
    try:
        p = subprocess.run(
            cmd, input=prompt.encode("utf-8"), capture_output=True,
            timeout=timeout, cwd=cwd or os.environ.get("TEMP", str(ROOT)),
            creationflags=CREATE_NO_WINDOW)
    except subprocess.TimeoutExpired:
        raise HTTPException(504, f"AI 응답 시간 초과({timeout}초)")
    if p.returncode != 0:
        # CLI가 오류를 stdout으로 내는 경우가 있다 — stderr만 보면 'unknown error'로 정보가 사라진다
        err = (p.stderr.decode("utf-8", "replace").strip()
               or p.stdout.decode("utf-8", "replace").strip())[-400:]
        raise HTTPException(502, f"AI 호출 실패: {err or f'종료코드 {p.returncode}'}")
    out = p.stdout.decode("utf-8", "replace").strip()
    if not out:
        # 종료코드 0 + 빈 stdout — 단서(모델 거절 안내·컨텍스트 초과·인증 갱신)는 stderr 에만 남는다
        err = p.stderr.decode("utf-8", "replace").strip()[-200:]
        raise HTTPException(502, f"AI가 빈 응답을 반환했습니다{f' — stderr: {err}' if err else ''}")
    return out


# 동시 claude 서브프로세스 상한 (전역). 번역 2워커에 요약·핵심·용어집·채팅 등이 겹치면
# 무제한일 때 서브프로세스가 쌓여 자원·구독 동시성 한계를 때린다 → 4개로 묶고 초과분은 큐잉.
# 모든 AI 호출은 ask_claude_async / ask_vision_async 를 거쳐 이 세마포어를 공유한다.
_CLAUDE_SEM = asyncio.Semaphore(4)


async def ask_claude_async(*args, **kwargs) -> str:
    """ask_claude는 blocking subprocess.run이다. async 엔드포인트에서 직접 부르면 단일
    이벤트 루프가 AI 호출(수십 초) 내내 멈춰 다른 요청(다른 논문 열기 포함)이 전부 얼어붙는다.
    스레드로 오프로드해 루프를 놓아준다 — AI 작업 중에도 서버가 계속 응답한다.
    _CLAUDE_SEM 으로 동시 실행 수를 제한한다."""
    async with _CLAUDE_SEM:
        return await asyncio.to_thread(ask_claude, *args, **kwargs)


async def ask_vision_async(*args, **kwargs) -> str:
    """이미지 Read 비전 호출(_ask_vision)도 같은 상한을 공유한다."""
    async with _CLAUDE_SEM:
        return await asyncio.to_thread(_ask_vision, *args, **kwargs)


def extract_json(raw: str):
    """모델 출력에서 첫 JSON 배열/객체를 꺼낸다 (앞뒤 산문·코드펜스 허용)."""
    m = re.search(r"```(?:json)?\s*([\[{].*?[\]}])\s*```", raw, re.S)
    cand = m.group(1) if m else None
    if cand is None:
        start = min((i for i in (raw.find("["), raw.find("{")) if i != -1), default=-1)
        if start == -1:
            raise ValueError("no json")
        cand = raw[start:]
        # 뒤에서부터 닫는 괄호를 줄여가며 파싱 시도
        for end in range(len(cand), max(len(cand) - 2000, 0), -1):
            if cand[end - 1] in "]}":
                try:
                    return json.loads(cand[:end])
                except json.JSONDecodeError:
                    continue
        raise ValueError("no json")
    return json.loads(cand)


def parse_json_or_502(raw: str, want=None, what: str = "AI 응답", empty_ok: bool = False):
    """extract_json 을 감싸 실패 원인과 받은 내용 앞부분을 502 에 실어 보낸다 — raw 를 버리면
    거절문·산문·형식 불일치를 구별할 수 없어 재시도가 의미 있는지 판단할 근거가 사라진다.
    want 를 주면 형태와 '비지 않음'까지 본다(빈 결과가 정상인 곳만 empty_ok=True)."""
    try:
        data = extract_json(raw)
    except (ValueError, json.JSONDecodeError) as e:
        raise HTTPException(502, f"{what}을 해석하지 못했습니다({type(e).__name__}). "
                                 f"받은 내용 앞부분: {raw[:200]!r}")
    if want and not isinstance(data, want):
        raise HTTPException(502, f"{what} 형식이 예상과 다릅니다({type(data).__name__}). "
                                 f"받은 내용 앞부분: {raw[:200]!r}")
    if want and not data and not empty_ok:
        raise HTTPException(502, f"{what}이 비어 있습니다. 받은 내용 앞부분: {raw[:200]!r}")
    return data


# ---------------------------------------------------------------- library

@app.get("/api/papers")
def list_papers():
    with JOBS_LOCK:                 # 이 논문을 물고 있는 실행 중 잡(준비·번역) — 서가 "AI 진행 중" 표시
        running_pids = {p for j in JOBS.values() if j["state"] == "running"
                        for p in j.get("pids", [])}
    out = []
    for d in sorted(LIB.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if meta:
            meta["id"] = d.name       # 폴더명이 정본 — meta 의 id 가 없거나 어긋나도 링크가 안 깨진다
            meta["has_thumb"] = (d / "thumb.jpg").exists()
            meta["has_text"] = (d / "text.json").exists()
            meta["ai_busy"] = d.name in running_pids
            if meta["has_text"]:    # pypdf 임시본이면 첫 열람 때 pdf.js 가 재추출해 덮는다
                meta["text_provisional"] = bool(read_json(d / "text.json", {}).get("provisional"))
            st = _prep_status(d)
            meta["prep"] = {"done": sum(st.values()), "total": len(PREP_STEPS)}
            out.append(meta)
    return out


UPLOAD_LOCK = threading.Lock()


@app.post("/api/papers")
def upload_paper(file: UploadFile):
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(400, "PDF 파일만 업로드할 수 있습니다")
    pid = uuid.uuid4().hex[:8]
    d = LIB / pid
    d.mkdir()
    size = 0
    try:
        with open(d / "paper.pdf", "wb") as dst:
            while chunk := file.file.read(1024 * 1024):
                size += len(chunk)
                if size > MAX_PDF_BYTES:
                    raise HTTPException(413, "PDF가 100MB를 초과합니다")
                dst.write(chunk)
    except BaseException:       # 디스크 오류·연결 끊김 포함 — 실패 잔재 폴더가 서가에 유령으로 남지 않게
        shutil.rmtree(d, ignore_errors=True)
        raise
    h = file_sha256(d / "paper.pdf")
    # 검사→기록 사이를 잠근다 — 창 두 개에서 같은 PDF를 동시에 올리면(sync 라우트라 진짜 병렬)
    # 둘 다 검사를 통과해 중복 항목이 생긴다. 해시 계산은 락 밖(위)에서 이미 끝났다.
    with UPLOAD_LOCK:
        dup = existing_pdf_hashes().get(h)    # meta 아직 없어 방금 만든 폴더는 자기매칭 안 됨
        if dup:                               # 내용 동일 PDF 이미 있음 → 새 폴더 버리고 기존 반환
            shutil.rmtree(d, ignore_errors=True)
            meta = read_json(paper_dir(dup) / "meta.json", {})
            meta["duplicate"] = True
            return meta
        title = re.sub(r"\.pdf$", "", file.filename, flags=re.I).replace("_", " ").strip() or "제목 없음"
        meta = {"id": pid, "title": title, "filename": file.filename,
                "added": int(time.time()), "pages": 0, "pdf_sha256": h}
        write_json(d / "meta.json", meta)
    _auto_prep([pid])
    return meta


@app.delete("/api/papers/{pid}")
def delete_paper(pid: str):
    d = paper_dir(pid)                      # 존재 검증(404)
    dropped = _prune_draft_refs(pid)
    last_err = None
    for i in range(3):                      # OneDrive가 방금 연 PDF 핸들을 놓을 시간을 준다
        try:
            shutil.rmtree(d)
            return {"ok": True, "draft_refs_dropped": dropped}
        except OSError as e:
            last_err = e                    # 어떤 파일이 왜 잠겼는지 — 삼키면 '성공했는데 그대로 있다'가 된다
            if i < 2:
                time.sleep(0.3)
    shutil.rmtree(d, ignore_errors=True)    # 그래도 잠겨 있으면 부분삭제 허용 — 잔여는 곧 사라진다
    if (d / "meta.json").exists():          # meta 가 남으면 논문은 서가에 계속 뜬다 — 성공이라 말할 수 없다
        raise HTTPException(409, f"삭제하지 못했습니다 — 파일이 사용 중입니다: {last_err}")
    return {"ok": True, "partial": str(last_err), "draft_refs_dropped": dropped}


META_LOCK = threading.Lock()


def _clean_tags(tags: list) -> list:
    """태그 정제: str화·trim → 빈 값 제거 → 24자 컷 → 중복 제거(순서 유지) → 최대 8개."""
    seen = set()
    out = []
    for t in tags:
        s = str(t).strip()
        if not s:
            continue
        s = s[:24]
        if s in seen:
            continue
        seen.add(s)
        out.append(s)
    return out[:8]


@app.patch("/api/papers/{pid}")
async def patch_paper(pid: str, req: Request):
    """읽기 상태·태그·읽던 위치 부분 수정. body에 온 필드만 meta.json에 반영한다."""
    d = paper_dir(pid)
    body = await req.json()
    with META_LOCK:
        meta = read_json(d / "meta.json", {})
        if "status" in body:
            status = body.get("status")
            if status not in ("none", "reading", "done"):
                raise HTTPException(400, "status는 none/reading/done 중 하나여야 합니다")
            if status == "none":
                meta.pop("status", None)     # 프론트는 falsy/none 둘 다 없음으로 취급하므로 저장해도 무방하지만 정리
            else:
                meta["status"] = status
        if "tags" in body:
            tags = body.get("tags")
            if not isinstance(tags, list):
                raise HTTPException(400, "tags는 배열이어야 합니다")
            meta["tags"] = _clean_tags(tags)
        # 읽던 위치. 여태 localStorage 에만 있어 PC 를 바꾸면 진행률이 0 으로 보였다 —
        # library/ 는 OneDrive 로 두 PC 가 함께 쓰는데 진도만 브라우저에 갇혀 있던 셈이다.
        pos = _clean_read_pos(body.get("read_pos"))
        if pos:
            meta["read_pos"] = pos
            meta["read_at"] = int(time.time())
        write_json(d / "meta.json", meta)
    return meta


def _clean_read_pos(raw) -> dict:
    """{p: 쪽, r: 0~1 비율} 만 남긴다. 값이 이상하면 통째로 버린다 — 진행률은 서가에
    막대로 그려지므로 1 을 넘기면 눈에 바로 띈다."""
    if not isinstance(raw, dict):
        return None
    try:
        p, r = int(raw.get("p") or 0), float(raw.get("r") or 0)
    except (TypeError, ValueError):
        return None
    if p < 1 or not (0 <= r <= 1):
        return None
    return {"p": p, "r": round(r, 4)}


# ---------------------------------------------------------------- 근접 중복 · 복습 큐
# 업로드는 sha256 완전일치만 막는다(upload_paper). 같은 논문이라도 프리프린트본·재다운로드본은
# 바이트가 달라 그물을 빠져나간다. 여기서 DOI·제목으로 한 번 더 본다 — 다만 **막지는 않는다**.
# 오탐으로 진짜 새 논문을 거절하는 손해가, 중복 한 편이 서가에 남는 손해보다 크다.

def _near_dups(pid: str, meta: dict) -> list:
    """이 논문과 같은 논문일 수 있는 서재의 다른 편. doi 완전일치 > 정규화 제목 일치 순."""
    doi = str(meta.get("doi") or "").lower().strip()
    sig = _alnum_sig(meta.get("title"))
    if not doi and len(sig) < 15:
        return []
    out = []
    for d in sorted(LIB.iterdir()):
        if not d.is_dir() or d.name == pid:
            continue
        m = read_json(d / "meta.json", None)
        if not m:
            continue
        mdoi = str(m.get("doi") or "").lower().strip()
        msig = _alnum_sig(m.get("title"))
        if doi and mdoi and doi == mdoi:
            why = "doi"
        elif len(sig) >= 15 and len(msig) >= 15 and (sig == msig or sig in msig or msig in sig):
            why = "title"
        else:
            continue
        out.append({"pid": d.name, "title": m.get("title") or d.name,
                    "year": m.get("year"), "why": why})
    return out


@app.get("/api/papers/{pid}/dupes")
def paper_dupes(pid: str):
    meta = read_json(paper_dir(pid) / "meta.json", {})
    return {"items": _near_dups(pid, meta)}


REVIEW_STALE_DAYS = 30


def _review_queue(now: int, stale_days: int = REVIEW_STALE_DAYS) -> list:
    """다시 볼 논문: 강조·노트를 남겼는데 오래 안 연 편. 표시해 둔 게 없는 논문은 부르지 않는다 —
    다시 볼 이유가 있어야 큐에 오른다. read_at 이 없는 옛 논문은 added 를 기준으로 본다."""
    cut = now - stale_days * 86400
    out = []
    for d in sorted(LIB.iterdir()):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        marks = (len(read_json(d / "highlights.json", {"items": []}).get("items") or [])
                 + len(read_json(d / "notes.json", []) or []))
        if not marks:
            continue
        seen = int(meta.get("read_at") or meta.get("added") or 0)
        if seen > cut:
            continue
        out.append({"pid": d.name, "title": meta.get("title") or d.name,
                    "marks": marks, "seen": seen,
                    "days": (now - seen) // 86400 if seen else None,
                    "status": meta.get("status"),
                    "pos": meta.get("read_pos") or None})
    out.sort(key=lambda x: (-x["marks"], x["seen"]))
    return out


@app.get("/api/library/review")
def library_review(days: int = REVIEW_STALE_DAYS):
    days = max(1, min(3650, int(days or REVIEW_STALE_DAYS)))
    return {"days": days, "items": _review_queue(int(time.time()), days)}


@app.get("/api/papers/{pid}/pdf")
def get_pdf(pid: str):
    p = paper_dir(pid) / "paper.pdf"
    if not p.exists():      # FileResponse 는 없는 파일에 RuntimeError → pdf.js 는 "(500)"만 남기고 원인이 사라진다
        raise HTTPException(404, "PDF 파일이 없습니다 — 업로드가 중단됐거나 파일이 삭제됐습니다")
    # no-store: 동기화 중 잘린 응답이 브라우저 캐시에 고착되는 것을 방지 (로컬 디스크가 곧 캐시)
    return FileResponse(p, media_type="application/pdf", headers={"Cache-Control": "no-store"})


@app.get("/api/papers/{pid}/thumb")
def get_thumb(pid: str):
    p = paper_dir(pid) / "thumb.jpg"
    if not p.exists():
        raise HTTPException(404, "썸네일 없음")
    return FileResponse(p, media_type="image/jpeg")


@app.put("/api/papers/{pid}/thumb")
async def put_thumb(pid: str, req: Request):
    data = await req.body()
    if len(data) > 400_000:
        raise HTTPException(413, "썸네일이 너무 큽니다")
    (paper_dir(pid) / "thumb.jpg").write_bytes(data)
    return {"ok": True}


# ---------------------------------------------------------------- text cache

@app.get("/api/papers/{pid}/text")
def get_text(pid: str):
    d = paper_dir(pid)
    t = read_json(d / "text.json", None)
    if not t:
        raise HTTPException(409, _no_text(d))
    return t


@app.put("/api/papers/{pid}/text")
async def put_text(pid: str, req: Request):
    d = paper_dir(pid)
    body = await req.json()
    pages = body.get("pages")
    if not isinstance(pages, list) or not pages:
        raise HTTPException(400, "pages 배열이 필요합니다")
    with TEXT_LOCK:
        write_json(d / "text.json", {"pages": pages})
    with META_LOCK:
        meta = read_json(d / "meta.json", {})
        meta["pages"] = len(pages)
        # PDF 내부 메타데이터 Title 은 "RSC_CS_C3CS60480B 3..26" 같은 조판 아티팩트인 경우가
        # 많다. Zotero 유래 논문은 이미 검증된 제목을 갖고 있으므로 절대 덮지 않는다.
        if body.get("title") and not (meta.get("zotero_key") or meta.get("citekey")):
            meta["title"] = str(body["title"])[:300]
        write_json(d / "meta.json", meta)
    _auto_prep([pid])       # pdf.js 정본 텍스트 도착 — pypdf 미설치 등으로 밀린 단계 자동 재개
    return {"ok": True, "title": meta.get("title")}


# ---------------------------------------------------------------- summary

@app.get("/api/papers/{pid}/summary")
def get_summary(pid: str):
    p = paper_dir(pid) / "summary.md"
    if not p.exists():
        raise HTTPException(404, "요약 없음")
    return {"markdown": p.read_text(encoding="utf-8")}


@app.post("/api/papers/{pid}/summary")
async def make_summary(pid: str):
    d = paper_dir(pid)
    meta = read_json(d / "meta.json", {})
    text = full_text(pid)[:MAX_CTX]
    system = (
        "너는 논문을 한국어로 요약하는 연구 조교다. 마크다운으로만 답한다. "
        "출력의 첫 글자는 반드시 `## 한 줄 요약`의 `#`이어야 한다 — 그 앞에 어떤 문장도 쓰지 마라. "
        "구조: `## 한 줄 요약` / `## 문제의식` / `## 방법` / `## 핵심 결과`(수치 포함) / "
        "`## 한계와 남은 질문` / `## 섹션별 흐름`(각 섹션 1-2문장). "
        "전문용어는 한국어(원어) 병기. 맺음말 금지."
    )
    md = await ask_claude_async(f"논문 제목: {meta.get('title','')}\n\n===== 논문 전문 =====\n{text}", system)
    cut = md.find("## ")
    if cut > 0:
        md = md[cut:]
    (d / "summary.md").write_text(md, encoding="utf-8")
    return {"markdown": md}


# ---------------------------------------------------------------- 구조 분석 (8섹션)
# 요약(summary.md)은 산문이라 읽기엔 좋지만 편끼리 나란히 놓고 비교할 수는 없다. 여기는 칸이
# 고정돼 있어서 "이 편의 데이터" "저 편의 한계"를 같은 자리에서 꺼낸다 — 근거표와 같은 축이다.
# 절마다 '근거 보드에 담기'가 붙어 그대로 초안의 재료가 된다.

ANALYSIS_SECTIONS = [
    ("basic", "기본 정보"), ("background", "연구 배경"), ("theory", "이론적 토대"),
    ("method", "실험·방법"), ("data", "데이터"), ("findings", "핵심 결과"),
    ("limits", "한계"), ("apply", "적용 지침"),
]
ANALYSIS_MAX = 1800        # 섹션 하나의 상한(자) — 넘으면 요약이 아니라 발췌가 된다


def _clean_analysis(raw: dict) -> list:
    """모델 출력 → 고정 8칸. 빠진 칸은 빈 문자열로 채워 항상 같은 모양이 나오게 한다 —
    화면이 '없는 칸'과 '아직 안 만든 분석'을 구별할 수 있어야 한다."""
    out = []
    for key, title in ANALYSIS_SECTIONS:
        body = raw.get(key) if isinstance(raw, dict) else None
        if isinstance(body, list):
            body = "\n".join(f"- {str(x).strip()}" for x in body if str(x).strip())
        out.append({"key": key, "title": title, "body": str(body or "").strip()[:ANALYSIS_MAX]})
    return out


@app.get("/api/papers/{pid}/analysis")
def get_analysis(pid: str):
    a = read_json(paper_dir(pid) / "analysis.json", None)
    if not a:
        raise HTTPException(404, "구조 분석 없음")
    return a


@app.post("/api/papers/{pid}/analysis")
async def make_analysis(pid: str):
    d = paper_dir(pid)
    meta = read_json(d / "meta.json", {})
    text = full_text(pid)[:MAX_CTX]
    keys = ", ".join(f'"{k}"({t})' for k, t in ANALYSIS_SECTIONS)
    system = (
        "논문을 고정된 8개 칸으로 정리하는 도구다. 출력은 JSON 객체 하나만, 다른 텍스트 금지. "
        f"키는 정확히 이 8개: {keys}. 값은 한국어 문자열(마크다운 목록 허용). "
        "각 칸 3~6문장. 결과·데이터 칸에는 논문에 나온 수치를 그대로 옮긴다. "
        "논문에 없는 칸은 빈 문자열로 둔다 — 채우려고 지어내지 마라. "
        "전문용어는 한국어(원어) 병기. 서두·맺음말 금지."
    )
    raw = await ask_claude_async(f"논문 제목: {meta.get('title','')}\n\n===== 논문 전문 =====\n{text}", system)
    data = parse_json_or_502(raw, dict)
    out = {"items": _clean_analysis(data), "ts": int(time.time())}
    write_json(d / "analysis.json", out)
    return out


# ---------------------------------------------------------------- 마인드맵 (논지 지도)
# 8칸 분석이 "편끼리 비교할 고정 축"이라면 여기는 "이 논문이 실제로 어떻게 논증하는가"다.
# 장식이 아니라 항법 장치로 만든다 — 모든 노드가 본문 한 문장을 가리키고, 누르면 그 쪽으로 뛴다.
# 앵커(p/s/se)는 핵심 4색과 **같은 스키마**라 좌표 매핑·재탐색·점프가 전부 재사용된다.

MINDMAP_MAX_NODES = 40
MINDMAP_MAX_DEPTH = 3          # 뿌리(0) + 축(1) + 근거(2) + 수치(3)
MINDMAP_CATS = ("nov", "met", "res", "lim", "idea")


def _clean_mindmap(raw, root_label: str) -> list:
    """모델 출력(평평한 노드 목록) → 뿌리 하나짜리 트리.

    부모를 못 찾은 노드는 버리지 않고 뿌리에 붙인다(내용은 멀쩡한데 id만 틀린 경우가 흔하다).
    순환은 뿌리에서 너비우선으로 훑어 자연히 떨어져 나간다 — 화면이 원형 배치라 순환이 하나만
    있어도 각도 배분이 끝나지 않는다."""
    root = {"id": "n0", "parent": None, "label": str(root_label or "논문")[:60],
            "c": "root", "p": 0, "s": "", "se": "", "note": "", "depth": 0}
    if not isinstance(raw, list):
        return [root]
    seen, nodes = {"n0"}, {}
    for i, it in enumerate(raw):
        if not isinstance(it, dict):
            continue
        label = " ".join(str(it.get("label") or "").split())[:60]
        if not label:
            continue
        nid = str(it.get("id") or "").strip()[:16] or f"a{i}"
        while nid in seen:
            nid += "_"
        seen.add(nid)
        try:
            pg = int(it.get("p") or 0)
        except (TypeError, ValueError):
            pg = 0
        c = it.get("c") if it.get("c") in MINDMAP_CATS else "idea"
        nodes[nid] = {"id": nid, "parent": str(it.get("parent") or "").strip()[:16] or "n0",
                      "label": label, "c": c, "p": pg,
                      "s": str(it.get("s") or "").strip()[:200],
                      "se": str(it.get("se") or "").strip()[:200],
                      "note": str(it.get("note") or "").strip()[:200]}
    for n in nodes.values():                      # 모르는 부모 → 뿌리로
        if n["parent"] not in nodes or n["parent"] == n["id"]:
            n["parent"] = "n0"
    # 순환 끊기: 뿌리에서 닿지 않는 덩어리는 버리지 않고 그중 하나를 뿌리에 붙인다.
    # 내용은 멀쩡한데 parent 만 어긋난 경우가 대부분이라 통째로 버리면 지도가 텅 빈다.
    for _ in range(len(nodes) + 1):
        reach, grew = {"n0"}, True
        while grew:
            grew = False
            for n in nodes.values():
                if n["parent"] in reach and n["id"] not in reach:
                    reach.add(n["id"])
                    grew = True
        orphans = [n for n in nodes.values() if n["id"] not in reach]
        if not orphans:
            break
        orphans[0]["parent"] = "n0"
    kids = {}
    for n in nodes.values():
        kids.setdefault(n["parent"], []).append(n)
    out, queue = [root], [("n0", 0)]
    while queue:
        pid, depth = queue.pop(0)
        if depth >= MINDMAP_MAX_DEPTH:
            continue
        for ch in kids.get(pid, []):
            if len(out) >= MINDMAP_MAX_NODES:
                return out
            ch["depth"] = depth + 1
            out.append(ch)
            queue.append((ch["id"], depth + 1))
    return out


@app.get("/api/papers/{pid}/mindmap")
def get_mindmap(pid: str):
    m = read_json(paper_dir(pid) / "mindmap.json", None)
    if not m:
        raise HTTPException(404, "마인드맵 없음")
    return m


@app.post("/api/papers/{pid}/mindmap")
async def make_mindmap(pid: str, req: Request):
    d = paper_dir(pid)
    try:
        body = await req.json()
    except Exception:
        body = {}
    cached = read_json(d / "mindmap.json", None)
    if cached and cached.get("nodes") and not body.get("force"):
        return cached
    t = read_json(d / "text.json", None)
    if not t or not t.get("pages"):
        raise HTTPException(409, _no_text(d))
    meta = read_json(d / "meta.json", {})
    src = "".join(f"===== {p['n']}페이지 =====\n{p['text']}\n\n" for p in t["pages"])[:MAX_CTX]
    system = (
        "너는 논문의 논지 구조를 지도로 그리는 도구다. 출력은 JSON 배열 하나만, 다른 텍스트 금지.\n"
        '형식: [{"id":"a1","parent":"n0","label":"12자 내외 짧은 구","c":"met",'
        '"p":페이지번호,"s":"문장 시작 5~8단어","se":"문장 끝 4~6단어","note":"한 구절 설명"}]\n'
        "구조: 뿌리 노드 id는 'n0'(논문 자체, 네가 만들지 마라). 1단계는 이 논문이 실제로 세우는 "
        "논지 축 3~6개(정해진 목차를 베끼지 말고 저자의 논증 흐름을 따를 것). "
        "2단계는 각 축의 근거 2~5개. 3단계는 수치·조건이 있을 때만.\n"
        "c 분류: nov(독창성) · met(방법) · res(결과) · lim(한계) · idea(앵커 없는 구조 노드).\n"
        "규칙: label은 짧은 명사구(문장 금지) · idea 가 아닌 노드는 p/s/se 를 반드시 채운다 · "
        "s/se 는 그 페이지 원문에 나온 그대로(verbatim, 한 글자도 바꾸지 말 것) · "
        "se 는 그 문장의 진짜 마지막(문장부호 직전) · 전체 노드 35개 이하 · 한국어 label."
    )
    raw = await ask_claude_async(f"논문 제목: {meta.get('title', '')}\n\n{src}", system)
    parsed = parse_json_or_502(raw, list)
    nodes = _clean_mindmap(parsed, meta.get("title") or "논문")
    if len(nodes) < 2:
        raise HTTPException(502, f"AI가 쓸 만한 지도를 만들지 못했습니다(노드 {len(nodes)}개). "
                                 f"받은 내용 앞부분: {raw[:200]!r}")
    data = {"nodes": nodes, "ts": int(time.time())}
    write_json(d / "mindmap.json", data)
    return data


# ---------------------------------------------------------------- chat

@app.get("/api/papers/{pid}/chat")
def get_chat(pid: str):
    return read_json(paper_dir(pid) / "chat.json", [])


@app.post("/api/papers/{pid}/chat")
async def post_chat(pid: str, req: Request):
    d = paper_dir(pid)
    body = await req.json()
    q = (body.get("question") or "").strip()
    quote = (body.get("quote") or "").strip()
    if not q:
        raise HTTPException(400, "질문이 비어 있습니다")
    text = full_text(pid)[:MAX_CTX]
    hist = read_json(d / "chat.json", [])
    convo = "\n".join(f"[{'사용자' if m['role']=='user' else 'AI'}] {m['content']}" for m in hist[-8:])
    system = (
        "너는 이 논문에 대해 답하는 연구 동료다. 답은 한국어 마크다운. "
        "논문 내용에 근거해 답하고, 논문에 없는 내용은 없다고 말한 뒤 일반 지식임을 표시한다. "
        "간결하되 수식/개념은 정확히. 서두/맺음말 금지."
    )
    prompt = f"===== 논문 전문 =====\n{text}\n\n===== 이전 대화 =====\n{convo or '(없음)'}\n\n"
    if quote:
        prompt += f"===== 사용자가 인용한 구절 =====\n{quote}\n\n"
    prompt += f"===== 질문 =====\n{q}"
    answer = await ask_claude_async(prompt, system)
    hist.append({"role": "user", "content": q, "quote": quote or None, "ts": int(time.time())})
    hist.append({"role": "ai", "content": answer, "ts": int(time.time())})
    write_json(d / "chat.json", hist)
    return {"answer": answer}


# ---------------------------------------------------------------- notes

@app.get("/api/papers/{pid}/notes")
def get_notes(pid: str):
    return read_json(paper_dir(pid) / "notes.json", [])


@app.post("/api/papers/{pid}/notes")
async def add_note(pid: str, req: Request):
    d = paper_dir(pid)
    body = await req.json()
    quote = (body.get("quote") or "").strip()
    if not quote:
        raise HTTPException(400, "인용문이 비어 있습니다")
    notes = read_json(d / "notes.json", [])
    note = {"id": uuid.uuid4().hex[:8], "page": int(body.get("page") or 0),
            "quote": quote[:2000], "memo": (body.get("memo") or "")[:2000],
            "created": int(time.time())}
    notes.insert(0, note)
    write_json(d / "notes.json", notes)
    return note


@app.delete("/api/papers/{pid}/notes/{nid}")
def del_note(pid: str, nid: str):
    d = paper_dir(pid)
    notes = [n for n in read_json(d / "notes.json", []) if n["id"] != nid]
    write_json(d / "notes.json", notes)
    return {"ok": True}


# ---------------------------------------------------------------- highlights

@app.get("/api/papers/{pid}/highlights")
def get_highlights(pid: str):
    return read_json(paper_dir(pid) / "highlights.json", {"items": []})


@app.put("/api/papers/{pid}/highlights")
async def put_highlights(pid: str, req: Request):
    body = await req.json()
    if not isinstance(body.get("items"), list):
        raise HTTPException(400, "items 배열이 필요합니다")
    write_json(paper_dir(pid) / "highlights.json", {"items": body["items"]})
    return {"ok": True}


@app.post("/api/papers/{pid}/highlights/auto")
async def auto_highlights(pid: str):
    """AI가 핵심 문장을 원문 그대로 뽑는다. 좌표 매핑은 클라이언트 몫."""
    text = full_text(pid)[:MAX_CTX]
    system = (
        "너는 논문의 핵심 문장을 고르는 도구다. 논문에서 가장 중요한 문장 8~14개를 "
        "원문 그대로(verbatim, 한 글자도 바꾸지 말 것) 골라라. "
        '출력은 JSON 배열 하나만: [{"sentence": "...", "reason": "핵심인 이유(한국어 한 문장)"}]. '
        "다른 텍스트 금지."
    )
    raw = await ask_claude_async(f"===== 논문 전문 =====\n{text}", system)
    items = parse_json_or_502(raw, list)
    clean = [{"sentence": str(i.get("sentence", "")).strip(),
              "reason": str(i.get("reason", "")).strip()}
             for i in items if isinstance(i, dict) and i.get("sentence")]
    write_json(paper_dir(pid) / "auto_hl.json", clean)
    return clean


@app.get("/api/papers/{pid}/highlights/auto")
def get_auto_highlights(pid: str):
    return read_json(paper_dir(pid) / "auto_hl.json", [])


# ---------------------------------------------------------------- key points (핵심 4색)

KP_CATS = ("nov", "met", "res", "lim")   # 분홍=독창성 · 연두=방법 · 연보라=결과 · 하늘=한계


@app.get("/api/papers/{pid}/keypoints")
def get_keypoints(pid: str):
    return read_json(paper_dir(pid) / "keypoints.json", {"items": None})


@app.post("/api/papers/{pid}/keypoints")
async def make_keypoints(pid: str, req: Request):
    """논문 핵심을 세 분류(독창성/방법/결과)로 뽑는다. 문장은 verbatim 시작·끝 앵커(s,se)로
    반환하고 좌표 매핑은 클라이언트 몫(정규화 인덱스) — Achird의 3색 핵심 하이라이트."""
    d = paper_dir(pid)
    try:
        body = await req.json()
    except Exception:
        body = {}
    cached = read_json(d / "keypoints.json", None)
    if cached and cached.get("items") and not body.get("force"):
        return cached
    t = read_json(d / "text.json", None)
    if not t or not t.get("pages"):
        raise HTTPException(409, _no_text(d))
    meta = read_json(d / "meta.json", {})
    src = "".join(f"===== {p['n']}페이지 =====\n{p['text']}\n\n" for p in t["pages"])[:MAX_CTX]
    system = (
        "너는 논문의 핵심을 네 분류로 표시하는 도구다. 논문 전체에서 다음 분류의 핵심 문장을 고른다:\n"
        "- nov(독창성): 이 연구가 기존과 다른 새로운 점 — 기여·차별점·novelty 주장\n"
        "- met(방법): 핵심 방법·실험 설계 — 어떻게 했는가\n"
        "- res(결과): 핵심 결과 — 무엇을 발견했는가 (수치가 있는 문장 우선)\n"
        "- lim(한계): 저자가 명시한 한계·제약·남은 질문 (논문에 적힌 것만, 추측 금지)\n"
        "nov/met/res는 분류마다 3~6개, lim은 2~4개(명시된 한계가 없으면 lim은 생략). "
        "초록에만 몰리지 말고 본문(서론·방법·결과·결론)에 고르게 분산한다.\n"
        '출력은 JSON 배열 하나만: [{"p": 페이지번호, "c": "nov", '
        '"s": "문장 시작 5~8단어", "se": "문장 끝 4~6단어", "note": "핵심인 이유(한국어 한 구절)"}]\n'
        "규칙: s/se는 그 페이지 원문에 나온 그대로(verbatim, 한 글자도 바꾸지 말 것) · "
        "se는 그 문장의 진짜 마지막(문장부호 직전, 다음 문장·제목 미포함) · p는 문장이 실제로 있는 "
        "페이지 번호 · 원문 등장 순서대로 · 다른 텍스트 절대 금지."
    )
    raw = await ask_claude_async(f"논문 제목: {meta.get('title', '')}\n\n{src}", system)
    parsed = parse_json_or_502(raw, list)
    items = []
    for it in parsed:
        if not isinstance(it, dict) or it.get("c") not in KP_CATS or not it.get("s"):
            continue
        try:
            pg = int(it.get("p") or 0)
        except (TypeError, ValueError):
            pg = 0
        items.append({"p": pg, "c": it["c"],
                      "s": str(it.get("s", "")).strip(), "se": str(it.get("se", "")).strip(),
                      "note": str(it.get("note", "")).strip()})
    data = {"items": items, "ts": int(time.time())}
    write_json(d / "keypoints.json", data)
    return data


# ---------------------------------------------------------------- vision (그림 해설 · 스캔 전사)
# pdf.js가 이미 렌더한 페이지 캔버스를 브라우저가 이미지로 보내오고, claude CLI가 Read 도구로
# 그 이미지를 '눈으로' 읽는다 — poppler 등 서버측 PDF 렌더 의존성이 필요 없다.

TEXT_LOCK = threading.Lock()
FIG_LOCK = threading.Lock()
FORMULA_LOCK = threading.Lock()


VISION_TMP = Path(tempfile.gettempdir()) / "achird_vision"


def _ask_vision(image: UploadFile, system: str, ask: str, timeout: int = 240) -> str:
    """업로드된 페이지 이미지를 시스템 temp에 임시 저장 → Read 허용 claude 호출 → 삭제.
    OneDrive 동기화 폴더(논문 디렉토리)가 아니라 로컬 temp에 둔다 — 잦은 생성·삭제가
    OneDrive 동기화 churn을 일으키지 않도록."""
    VISION_TMP.mkdir(exist_ok=True)
    tmp = VISION_TMP / f"v_{uuid.uuid4().hex[:8]}.png"
    try:
        with open(tmp, "wb") as f:
            shutil.copyfileobj(image.file, f, length=1024 * 1024)
        prompt = f"{tmp.resolve()} 이미지를 Read 도구로 읽어라.\n\n{ask}"
        return ask_claude(prompt, system, timeout=timeout, tools="Read", cwd=str(VISION_TMP))
    finally:
        try:
            tmp.unlink()
        except OSError:
            pass


def _vision_key(n: int, region: str, hint: str):
    """(페이지, 영역, 힌트) → figures/formulas 공용 캐시 키. region("rx,ry,rw,rh", 0~1 비율)이
    float 4개 형식이 아니면 없는 것으로 취급(느슨한 검증). 정제된 region과 키를 함께 돌려준다.
    되돌리기는 _parse_figure_key — 형식을 바꾸면 반드시 짝으로 고칠 것."""
    region = region.strip()
    if region:
        parts = region.split(",")
        if len(parts) != 4:
            region = ""
        else:
            try:
                [float(p) for p in parts]
            except ValueError:
                region = ""
    key = f"{n}|{region}|{hint.strip()[:80]}" if region else f"{n}|{hint.strip()[:80]}"
    return region, key


@app.post("/api/papers/{pid}/figure")
async def explain_figure(pid: str, n: int = Form(...), hint: str = Form(""),
                         region: str = Form(""), image: UploadFile = File(...)):
    """페이지 이미지를 보고 그림·표를 해설한다. (페이지, 영역, 힌트) 단위 캐시.
    region: "rx,ry,rw,rh"(0~1 비율) — 클라이언트가 이미 그 영역만 잘라 보낸 이미지일 때 채워진다."""
    d = paper_dir(pid)
    region, key = _vision_key(n, region, hint)
    cached = read_json(d / "figures.json", {})
    if key in cached:
        return {"n": n, "markdown": cached[key], "cached": True}
    t = read_json(d / "text.json", {"pages": []})
    src = next((p["text"] for p in t.get("pages", []) if p.get("n") == n), "")
    system = (
        "너는 학술 논문의 그림·표 해설가다. 페이지 이미지를 직접 보고 한국어 마크다운으로 "
        "설명한다. 구조: 무엇을 보여주는지 한 줄 → 구성(패널·축·기호·단위) → 핵심 패턴과 "
        "수치 → 이 논문 맥락에서의 의미. 이미지에 보이지 않는 내용을 지어내지 마라. "
        "서두·맺음말·메타 발언 금지."
    )
    ask = (f"이 페이지({n}쪽)에서 '{hint.strip()}'에 해당하는 그림/표를 찾아 설명하라."
           if hint.strip() else f"이 페이지({n}쪽)의 그림/표를 설명하라.")
    if region:      # 크롭 영역 설명 — hint가 있으면 기존 hint 문구를 뒤에 이어붙인다
        prefix = f"이 이미지는 {n}쪽에서 사용자가 직접 지정한 영역만 잘라낸 것이다. 이 영역의 그림/표를 설명하라."
        ask = prefix if not hint.strip() else f"{prefix} {ask}"
    if src.strip():
        ask += f"\n\n[이 페이지의 추출 텍스트 — 캡션·본문 참고용]\n{src[:4000]}"
    out = await ask_vision_async(image, system, ask)
    with FIG_LOCK:
        cached = read_json(d / "figures.json", {})
        cached[key] = out
        write_json(d / "figures.json", cached)
    return {"n": n, "markdown": out}


def _parse_figure_key(key: str):
    """캐시 키(n|region|hint 또는 n|hint)를 (n, region, hint)로 되돌린다. region 자리가
    콤마 4개 float로 파싱되면 region으로, 아니면 나머지를 통째로 hint로 합친다(hint에 '|'가
    섞여 있어도 보존)."""
    parts = key.split("|", 2)
    n = int(parts[0])
    region = None
    if len(parts) == 3:
        rparts = parts[1].split(",")
        if len(rparts) == 4:
            try:
                region = [float(p) for p in rparts]
            except ValueError:
                region = None
    hint = (parts[2] if region is not None else "|".join(parts[1:])) or None
    return n, region, hint


@app.get("/api/papers/{pid}/figures")
def get_figures(pid: str):
    cached = read_json(paper_dir(pid) / "figures.json", {})
    items, dropped = [], 0
    for key, md in cached.items():
        try:
            n, region, hint = _parse_figure_key(key)
        except ValueError:      # 한 항목이 목록 전체를 죽여선 안 되지만, 말없이 사라져서도 안 된다
            dropped += 1
            print(f"[figures] 잘못된 캐시 키 {key!r}", file=sys.stderr)
            continue
        items.append({"key": key, "n": n, "region": region, "hint": hint, "md": md})
    items.sort(key=lambda x: x["n"])
    return {"items": items, "dropped": dropped}


@app.post("/api/papers/{pid}/formula")
async def explain_formula(pid: str, n: int = Form(...), hint: str = Form(""),
                          region: str = Form(""), image: UploadFile = File(...)):
    """페이지 이미지를 보고 수식을 LaTeX로 전사하고 한국어로 해설한다. (페이지, 영역, 힌트) 단위 캐시.
    region: "rx,ry,rw,rh"(0~1 비율) — 클라이언트가 이미 그 영역만 잘라 보낸 이미지일 때 채워진다.
    /figure와 폼필드·캐시키 구조는 동일하되, 결과가 markdown 한 덩어리가 아니라 latex/explain으로
    나뉘므로 캐시값도 문자열이 아니라 객체다."""
    d = paper_dir(pid)
    region, key = _vision_key(n, region, hint)
    cached = read_json(d / "formulas.json", {})
    if key in cached:
        return {**cached[key], "n": n, "cached": True}
    t = read_json(d / "text.json", {"pages": []})
    src = next((p["text"] for p in t.get("pages", []) if p.get("n") == n), "")
    system = (
        "너는 학술 논문의 수식 전사·해설가다. 페이지 이미지를 직접 보고 수식을 정확한 LaTeX로 "
        "옮기고 한국어로 해설한다. 이미지에 보이지 않는 내용을 지어내지 마라. "
        '출력은 JSON 하나만: {"latex": "LaTeX 수식(둘러싸는 $$ 없이 수식 본문만)", '
        '"explain": "각 기호가 무엇을 뜻하고 이 수식이 전체적으로 무엇을 하는지 한국어로 설명"}. '
        "다른 텍스트 절대 금지."
    )
    ask = (f"이 페이지({n}쪽)에서 '{hint.strip()}'에 해당하는 수식을 찾아 전사·해설하라."
           if hint.strip() else f"이 페이지({n}쪽)의 핵심 수식을 전사·해설하라.")
    if region:      # 크롭 영역 설명 — hint가 있으면 기존 hint 문구를 뒤에 이어붙인다
        prefix = f"이 이미지는 {n}쪽에서 사용자가 직접 지정한 영역만 잘라낸 것이다. 이 영역의 수식을 전사·해설하라."
        ask = prefix if not hint.strip() else f"{prefix} {ask}"
    if src.strip():
        ask += f"\n\n[이 페이지의 추출 텍스트 — 주변 문맥 참고용]\n{src[:4000]}"
    raw = await ask_vision_async(image, system, ask)
    parsed = parse_json_or_502(raw, dict)
    result = {"latex": str(parsed.get("latex", "")).strip(), "explain": str(parsed.get("explain", "")).strip()}
    with FORMULA_LOCK:
        cached = read_json(d / "formulas.json", {})
        cached[key] = result
        write_json(d / "formulas.json", cached)
    return {**result, "n": n}


@app.get("/api/papers/{pid}/formulas")
def get_formulas(pid: str):
    cached = read_json(paper_dir(pid) / "formulas.json", {})
    items, dropped = [], 0
    for key, val in cached.items():
        try:
            n, region, hint = _parse_figure_key(key)     # 캐시 키 형식이 /figure와 동일해 그대로 재사용
        except ValueError:      # 한 항목이 목록 전체를 죽여선 안 되지만, 말없이 사라져서도 안 된다
            dropped += 1
            print(f"[formulas] 잘못된 캐시 키 {key!r}", file=sys.stderr)
            continue
        items.append({"key": key, "n": n, "region": region, "hint": hint,
                      "latex": val.get("latex", ""), "explain": val.get("explain", "")})
    items.sort(key=lambda x: x["n"])
    return {"items": items, "dropped": dropped}


@app.post("/api/papers/{pid}/vision-text")
async def vision_text(pid: str, n: int = Form(...), image: UploadFile = File(...)):
    """스캔 PDF 폴백: 페이지 이미지를 전사(transcribe)해 text.json을 백필한다.
    이후 번역·요약·채팅·용어집이 이 텍스트로 동작한다(텍스트 레이어가 없어
    원문 위 하이라이트·문장 연결은 원리상 불가)."""
    d = paper_dir(pid)
    system = (
        "너는 문서 전사(transcription) 도구다. 페이지 이미지에 보이는 모든 텍스트를 읽기 "
        "순서(다단이면 왼쪽 단부터)대로 원문 그대로 전사한다. 번역·요약·해설·메타 발언 금지. "
        "그림 캡션과 표 내용도 포함하고, 표는 행 단위 텍스트로 푼다. 알아볼 수 없는 부분은 "
        "[판독불가]로 표시한다. 출력은 전사 텍스트만."
    )
    out = await ask_vision_async(image, system, f"이 페이지({n}쪽)를 전사하라.")
    with TEXT_LOCK:
        t = read_json(d / "text.json", {"pages": []})
        pages = t.get("pages", [])
        hit = next((p for p in pages if p.get("n") == n), None)
        if hit is None:
            pages.append({"n": n, "text": out, "vision": True})
            pages.sort(key=lambda p: p.get("n", 0))
        else:
            hit["text"] = out
            hit["vision"] = True
        write_json(d / "text.json", {"pages": pages})
    return {"n": n, "text": out}


# ---------------------------------------------------------------- glossary (용어집)

@app.get("/api/papers/{pid}/glossary")
def get_glossary(pid: str):
    return read_json(paper_dir(pid) / "glossary.json", {"items": None})


@app.post("/api/papers/{pid}/glossary")
async def make_glossary(pid: str, req: Request):
    """전문용어·약어를 1회 추출한다. term은 본문 표기 그대로(verbatim)라
    클라이언트가 정규화 인덱스로 첫 등장 위치 점프를 만들 수 있다."""
    d = paper_dir(pid)
    try:
        body = await req.json()
    except Exception:
        body = {}
    cached = read_json(d / "glossary.json", None)
    if cached and cached.get("items") and not body.get("force"):
        return cached
    text = full_text(pid)[:MAX_CTX]
    meta = read_json(d / "meta.json", {})
    system = (
        "너는 논문 용어집을 만드는 도구다. 이 분야 비전공 대학원생이 걸려 넘어질 전문용어· "
        "약어 15~30개를 골라라. 약어는 반드시 포함하고 풀네임을 def에 적는다.\n"
        '출력은 JSON 배열 하나만: [{"term": "본문 표기 그대로(verbatim)", '
        '"ko": "한국어 번역", "def": "한 문장 정의(약어면 풀네임 포함)"}]\n'
        "규칙: term은 본문에 실제로 나온 문자열 그대로(한 글자도 바꾸지 말 것) · "
        "일반 상식 단어 제외 · 중요도 순 · 다른 텍스트 절대 금지."
    )
    raw = await ask_claude_async(f"논문 제목: {meta.get('title', '')}\n\n===== 논문 전문 =====\n{text}", system)
    parsed = parse_json_or_502(raw, list)
    items = [{"term": str(i.get("term", "")).strip(), "ko": str(i.get("ko", "")).strip(),
              "def": str(i.get("def", "")).strip()}
             for i in parsed if isinstance(i, dict) and i.get("term")]
    data = {"items": items, "ts": int(time.time())}
    write_json(d / "glossary.json", data)
    return data


@app.put("/api/papers/{pid}/glossary")
async def put_glossary(pid: str, req: Request):
    """용어집 직접 수정 저장 — 역어(ko)가 번역 용어 고정 사전으로 쓰이므로 사용자가 다듬을 수 있다."""
    body = await req.json()
    items = body.get("items")
    if not isinstance(items, list):
        raise HTTPException(400, "items 배열이 필요합니다")
    clean = [{"term": str(i.get("term", "")).strip()[:200], "ko": str(i.get("ko", "")).strip()[:200],
              "def": str(i.get("def", "")).strip()[:500]}
             for i in items if isinstance(i, dict) and str(i.get("term", "")).strip()][:200]
    data = {"items": clean, "ts": int(time.time())}
    write_json(paper_dir(pid) / "glossary.json", data)
    return data


# ---------------------------------------------------------------- library search (서재 전체 검색)
# 로컬 개인 서재(수십 편)라 색인 DB 없이 브루트포스 + mtime 캐시로 충분하다 — 색인 유지보수
# 버그 원천 제거. 원문(text.json)과 번역(translation.json)을 함께 검색한다.

_SEARCH_CACHE: dict = {}   # pid → {"tk": text mtime, "rk": trans mtime, "pages": [...], "tpages": [...]}


def _norm_search(t: str) -> str:
    """검색용 정규화: soft hyphen 제거 · 행바꿈 하이픈 결합(atten-\\ntion→attention) · 공백 압축.
    스니펫도 이 문자열에서 그대로 자르므로 위치 매핑이 필요 없다."""
    t = t.replace("­", "")
    t = re.sub(r"-\s*\n\s*", "", t)
    t = re.sub(r"\s+", " ", t)
    return t.strip()


def _mtime(p: Path) -> float:
    try:
        return p.stat().st_mtime
    except OSError:
        return 0.0


def _vault_note_path(pid: str, d: Path):
    """이 논문의 내보낸 Obsidian 노트 경로. 볼트 미설정·노트 없음·남의 노트면 None."""
    try:
        cfg = load_config()
        vault = _rehome(cfg.get("obsidian_vault_path", ""))
        if not vault or not Path(vault).is_dir():
            return None
        sub = str(cfg.get("obsidian_subfolder", "Achird")).strip() or "Achird"
        fname = safe_filename(read_json(d / "meta.json", {}).get("title", pid), pid)
        base = Path(vault) / sub
        for p in (base / f"{fname}.md", base / f"{fname} ({pid}).md"):
            if not p.is_file():
                continue
            with open(p, encoding="utf-8", errors="replace") as f:
                head = f.read(2000)     # 번역 포함 노트는 수백 KB — 소유 판정에 전체 읽기는 낭비
            if not _note_owned_by_other(head, pid):
                return p
    except OSError:
        pass
    return None


def _vault_memo_text(p: Path) -> str:
    """노트의 사용자 영역 — 관리블록(AUTO_END) 뒤의 "내 메모"만. 관리블록 안 내용은
    achird 데이터의 사본이라 다시 색인하면 같은 문장이 두 번 걸린다."""
    try:
        txt = p.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return ""
    i = txt.rfind(AUTO_END)
    memo = txt[i + len(AUTO_END):] if i >= 0 else ""
    return re.sub(r"^\s*#+\s*내 메모\s*", "", memo.strip()).strip()


def _search_entry(pid: str, d: Path) -> dict:
    """페이지별 (n, 원형, 소문자) 캐시. 원문·번역에 더해 내가 남긴 것(노트·하이라이트)과
    Obsidian 노트의 "내 메모"도 담는다 — "그 메모 어느 논문에 남겼더라"가 검색되지 않으면
    노트는 쌓일수록 못 찾는 데이터가 된다. 어느 소스든 mtime이 바뀌면 재구축한다."""
    vp = _vault_note_path(pid, d)
    key = (_mtime(d / "text.json"), _mtime(d / "translation.json"),
           _mtime(d / "notes.json"), _mtime(d / "highlights.json"),
           _mtime(vp) if vp else 0.0)
    e = _SEARCH_CACHE.get(pid)
    if e and e["key"] == key:
        return e
    tk, rk, nk, hk, vk = key
    pages, tpages, npages, hpages, vpages = [], [], [], [], []
    if vk and vp:
        s = _norm_search(re.sub(r"[#*`|>_]+", " ", _vault_memo_text(vp)))
        if s:
            vpages.append((0, s, s.lower()))
    if tk:
        for p in read_json(d / "text.json", {"pages": []}).get("pages", []):
            s = _norm_search(str(p.get("text", "")))
            if s:
                pages.append((int(p.get("n", 0)), s, s.lower()))
    if rk:
        for n, mdtext in read_json(d / "translation.json", {"pages": {}}).get("pages", {}).items():
            s = _norm_search(re.sub(r"[#*`|>_]+", " ", str(mdtext)))   # 마크다운 기호는 스니펫 잡음
            if s:
                tpages.append((int(n), s, s.lower()))
    if nk:
        for note in read_json(d / "notes.json", []):
            # 인용문과 메모를 한 덩어리로 — 둘 중 어느 쪽에 걸려도 같은 노트를 찾게 된다
            s = _norm_search(f"{note.get('quote', '')} {note.get('memo', '')}")
            if s:
                npages.append((int(note.get("page") or 0), s, s.lower()))
    if hk:
        for h in read_json(d / "highlights.json", {"items": []}).get("items", []):
            s = _norm_search(f"{h.get('text', '')} {h.get('reason', '')}")
            if s:
                hpages.append((int(h.get("page") or 0), s, s.lower()))
    e = {"key": key, "pages": pages, "tpages": tpages, "npages": npages, "hpages": hpages,
         "vpages": vpages}
    _SEARCH_CACHE[pid] = e
    return e


def _snip(orig: str, low: str, i: int, qlen: int) -> str:
    """매치 주변 ~160자 스니펫. low는 orig.lower()라 인덱스가 (희귀 유니코드 제외) 일치 — 클램프로 안전."""
    a = max(0, i - 70)
    b = min(len(orig), i + qlen + 90)
    return ("…" if a > 0 else "") + orig[a:b].strip() + ("…" if b < len(orig) else "")


@app.get("/api/search")
def search_library(q: str = ""):
    q = q.strip()
    if len(q) < 2:
        return {"q": q, "papers": []}
    ql = _norm_search(q).lower()
    out = []
    for d in sorted(LIB.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        e = _search_entry(d.name, d)
        title_hit = ql in str(meta.get("title", "")).lower()
        hits, total = [], 0
        # 소스마다 자리를 따로 준다 — 한 통에 담으면 분량이 압도적인 원문이 슬롯을 다 먹어
        # 노트·강조는 걸려도 화면에 안 나온다(이 기능을 넣은 이유가 사라진다).
        # 하이라이트는 내가 그은 것과 AI가 뽑은 것이 한 파일에 섞여 있다 — "내 강조"라 쓰면 사실과 다르다
        for src, plist, cap in (("원문", e["pages"], 3), ("번역", e["tpages"], 3),
                                ("내 노트", e["npages"], 2), ("강조", e["hpages"], 2),
                                ("Obsidian 메모", e.get("vpages") or [], 2)):
            shown = 0
            for n, orig, low in plist:
                c = low.count(ql)
                if not c:
                    continue
                total += c
                if shown < cap:
                    shown += 1
                    i = low.find(ql)
                    hits.append({"n": n, "src": src, "count": c, "snip": _snip(orig, low, i, len(ql))})
        if total or title_hit:
            out.append({"id": d.name, "title": meta.get("title", ""), "title_hit": title_hit,
                        "total": total, "hits": hits})
    out.sort(key=lambda p: (p["title_hit"], p["total"]), reverse=True)
    return {"q": q, "papers": out[:8]}


# ---------------------------------------------------------------- library graph (인용 그래프)
# static/app.js의 buildRefLinks(참고문헌 ↔ 서재 논문을 doi/정규화제목으로 매칭)를 서버측에서
# 재현해 서재 전체의 인용 그래프를 만든다. AI 호출 없는 순수 매칭이라 빠르다.

_ALNUM_RE = re.compile(r"[^a-z0-9]")


def _alnum_sig(t) -> str:
    """DOI/제목 매칭용 정규화: 소문자화 후 영숫자만 남긴다(app.js buildRefLinks의 alnum과 동일)."""
    return _ALNUM_RE.sub("", str(t or "").lower())


def _ref_pair(r) -> tuple:
    """참고문헌 항목 → (원문, 검증된 DOI). 검증 전에는 문자열이거나 oa가 없는 dict라
    DOI 자리가 빈다 — 그때는 종전대로 본문 부분일치로 떨어진다(옛 호출부·옛 refs.json 호환)."""
    if isinstance(r, dict):
        return str(r.get("text") or ""), str((r.get("oa") or {}).get("doi") or "").lower().strip()
    return str(r or ""), ""


def _graph_edges(papers: list) -> list:
    """papers: [{"id","title","doi","refs":[참고문헌(문자열 또는 {text,oa}), ...]}, ...].
    매칭은 세 단계로 내려간다 — ① 검증된 DOI 완전일치(가장 확실) ② 참고문헌 본문에 상대 DOI가
    박혀 있는지 ③ 정규화 제목 부분일치. 같은 (from,to) 쌍은 먼저 찾은 것 하나만 남긴다
    (그래프이지 멀티그래프가 아니므로). AI 호출 없는 순수 함수 — 서재 규모에서 충분히 빠르다."""
    cands = []
    for p in papers:
        sig = _alnum_sig(p.get("title"))
        cands.append({"id": p["id"], "doi": str(p.get("doi") or "").lower().strip(),
                      "sig": sig if len(sig) >= 15 else ""})
    edges, seen = [], set()
    for p in papers:
        others = [c for c in cands if c["id"] != p["id"] and (c["doi"] or c["sig"])]
        if not others:
            continue
        for ref in p.get("refs") or []:
            ref_text, ref_doi = _ref_pair(ref)
            why = "doi"
            hit = next((c for c in others if c["doi"] and ref_doi and c["doi"] == ref_doi), None)
            if hit:
                why = "doi_verified"
            if not hit:
                ref_low = ref_text.lower()
                hit = next((c for c in others if c["doi"] and c["doi"] in ref_low), None)
            if not hit:
                ref_sig = _alnum_sig(ref_text)
                hit = next((c for c in others if c["sig"] and c["sig"] in ref_sig), None)
                why = "title"
            if hit and (p["id"], hit["id"]) not in seen:
                seen.add((p["id"], hit["id"]))
                edges.append({"from": p["id"], "to": hit["id"], "why": why})
    return edges


# ---- 노트 개념 그래프: 노트·강조 메모의 [[개념]] 표기를 노드로 세운다.
# 인용 그래프는 논문끼리만 이어서, 서로 인용하지 않지만 같은 개념을 다루는 논문이 남남으로 남았다.

_WIKI_RE = re.compile(r"\[\[([^\[\]|]{1,80})(?:\|[^\[\]]{0,80})?\]\]")


def _wikilinks(text) -> list:
    """[[개념]] 표기를 뽑는다. [[대상|별칭]] 형태는 앞쪽(실제 대상)만 쓴다.
    같은 글 안의 중복은 접고, 대소문자·공백만 다른 것은 같은 개념으로 본다."""
    out, seen = [], set()
    for m in _WIKI_RE.finditer(str(text or "")):
        name = " ".join(m.group(1).split())
        key = name.lower()
        if name and key not in seen:
            seen.add(key)
            out.append(name)
    return out


def _concept_links(pid: str, d: Path) -> list:
    """이 논문의 노트·강조에 적힌 개념 이름들(중복 포함 — 세면 그 논문의 언급 횟수)."""
    names = []
    for nt in read_json(d / "notes.json", []):
        names += _wikilinks(nt.get("memo")) + _wikilinks(nt.get("quote"))
    for h in read_json(d / "highlights.json", {"items": []}).get("items", []):
        names += _wikilinks(h.get("reason"))
    return names


@app.get("/api/library/graph")
def library_graph(concepts: int = 0):
    """서재 전체 노드(id/title/year/status/refs개수)와, 참고문헌 매칭으로 만든 엣지를 반환한다.
    refs.json이 없는 논문도 노드로는 나오되(refs:0), 매칭할 참고문헌이 없으니 엣지의 출발점은
    되지 못한다 — _graph_edges가 빈 refs 리스트를 받으면 자연히 건너뛴다.

    concepts=1이면 노트·강조의 [[개념]] 표기를 노드로 함께 세운다. 개념 노드 id는 'c:소문자이름'
    이라 논문 id(8자리 hex)와 절대 겹치지 않는다."""
    papers, concept_edges, cnames, ccount = [], [], {}, {}
    for d in sorted(LIB.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        refs = read_json(d / "refs.json", None)
        ref_items = (refs or {}).get("items") or []
        papers.append({
            "id": d.name, "title": meta.get("title", ""), "year": meta.get("year"),
            "status": meta.get("status"), "doi": meta.get("doi"),
            "refs": [r for r in ref_items if isinstance(r, dict)],
        })
        if concepts:
            for name in _concept_links(d.name, d):
                key = "c:" + name.lower()
                cnames.setdefault(key, name)
                ccount[key] = ccount.get(key, 0) + 1
                if not any(e["from"] == d.name and e["to"] == key for e in concept_edges):
                    concept_edges.append({"from": d.name, "to": key, "why": "wikilink"})
    nodes = [{"id": p["id"], "kind": "paper", "title": p["title"], "year": p["year"],
              "status": p["status"], "refs": len(p["refs"]),
              "verified": sum(1 for r in p["refs"] if r.get("oa"))} for p in papers]
    nodes += [{"id": k, "kind": "concept", "title": cnames[k], "year": None,
               "status": None, "refs": ccount[k], "verified": 0}
              for k in sorted(cnames, key=lambda k: (-ccount[k], k))]
    return {"nodes": nodes, "edges": _graph_edges(papers) + concept_edges,
            "concepts": len(cnames)}


# ---------------------------------------------------------------- library AI (서재 질문 · 논문 비교)

CMP_LOCK = threading.Lock()


def _find_passages(entry: dict, ql: str, per_source: int = 2) -> list:
    """검색 캐시에서 키워드 주변 넓은 구절을 뽑는다 — 서재 질문의 근거 컨텍스트용."""
    out = []
    for src, plist in (("원문", entry["pages"]), ("번역", entry["tpages"])):
        cnt = 0
        for n, orig, low in plist:
            i = low.find(ql)
            if i == -1:
                continue
            a, b = max(0, i - 150), min(len(orig), i + 250)
            out.append((n, src, ("…" if a else "") + orig[a:b].strip() + ("…" if b < len(orig) else "")))
            cnt += 1
            if cnt >= per_source:
                break
    return out


@app.post("/api/library/ask")
async def library_ask(req: Request):
    """서재 전체에 질문. 키워드(haiku) → 서재 검색으로 근거 구절 수집 → sonnet이 출처 인용 답변."""
    body = await req.json()
    q = (body.get("question") or "").strip()
    if not q:
        raise HTTPException(400, "질문이 비어 있습니다")
    papers = []
    for d in sorted(LIB.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if meta:
            papers.append((d, meta))
    if not papers:
        raise HTTPException(409, "서재가 비어 있습니다")
    # 1) 영어 논문에서 찾을 검색 키워드 (질문이 한국어여도 원문을 때릴 수 있게)
    kw_sys = ("질문에 답할 근거를 영어 논문에서 찾기 위한 검색어를 만든다. 영어 키워드 2~4개와 "
              "질문의 핵심 한국어 단어 1~2개(번역본 검색용). "
              '출력은 JSON 배열 하나만: ["keyword", ...]. 다른 텍스트 금지.')
    try:
        kws = extract_json(await ask_claude_async(f"질문: {q}", kw_sys, timeout=90, model="haiku"))
        kws = [str(k).strip() for k in kws if str(k).strip()] if isinstance(kws, list) else []
    except (ValueError, json.JSONDecodeError, HTTPException):
        kws = []
    keys, seen = [], set()
    for k in kws + [q]:
        nk = _norm_search(k).lower()
        if len(nk) >= 2 and nk not in seen:
            seen.add(nk)
            keys.append(nk)
    # 2) 논문별 요약 + 근거 구절 조립
    blocks = ["===== 서재 목록 ====="]
    out_papers = []
    for n_, (d, meta) in enumerate(papers, 1):
        blocks.append(f"[{n_}] {meta.get('title', '')}")
        out_papers.append({"n": n_, "id": d.name, "title": meta.get("title", "")})
    for n_, (d, meta) in enumerate(papers, 1):
        sec = [f"===== [{n_}] {meta.get('title', '')} ====="]
        smp = d / "summary.md"
        if smp.exists():
            sec.append("[요약]\n" + smp.read_text(encoding="utf-8")[:1500])
        entry = _search_entry(d.name, d)
        hits = []
        for k in keys[:5]:
            hits += _find_passages(entry, k)
            if len(hits) >= 6:
                break
        dedup, seenp = [], set()
        for n2, src, snip in hits:
            hkey = (n2, snip[:60])
            if hkey not in seenp:
                seenp.add(hkey)
                dedup.append(f"({src} p.{n2}) {snip}")
        if dedup:
            sec.append("[관련 구절]\n" + "\n".join(dedup[:6]))
        blocks.append("\n".join(sec))
    system = (
        "너는 개인 논문 서재의 연구 사서다. 서재 논문들의 요약과 검색된 근거 구절이 주어진다. "
        "질문에 한국어 마크다운으로 답한다. 규칙: 근거가 있는 주장 끝에 [번호·p쪽] 인용(예: [2·p.5]) · "
        "주어진 자료로 답할 수 없으면 그렇다고 밝히고, 일반 지식으로 보충할 땐 그렇게 표시 · "
        "마지막 줄에 어느 논문 몇 쪽을 읽으면 되는지 한 줄 안내 · 서두·맺음말 금지."
    )
    prompt = "\n\n".join(blocks)[:MAX_CTX] + f"\n\n===== 질문 =====\n{q}"
    answer = await ask_claude_async(prompt, system, timeout=300)
    return {"answer": answer, "papers": out_papers}


@app.post("/api/library/compare")
async def library_compare(req: Request):
    """선택한 2~3편 전문을 나란히 읽고 비교표+차이 분석. 같은 조합은 캐시."""
    body = await req.json()
    ids = [i for i in (body.get("ids") or []) if isinstance(i, str)]
    if not 2 <= len(ids) <= 3:
        raise HTTPException(400, "비교할 논문을 2~3편 선택하세요")
    key = "|".join(sorted(ids))
    cache = read_json(LIB / "_compare.json", {})
    if key in cache and not body.get("force"):
        return {"markdown": cache[key]["md"], "papers": cache[key].get("papers", []), "cached": True}
    parts, papers = [], []
    budget = (MAX_CTX - 5000) // len(ids)
    for k, pid in enumerate(ids, 1):
        d = paper_dir(pid)
        meta = read_json(d / "meta.json", {})
        t = read_json(d / "text.json", None)
        if not t or not t.get("pages"):
            raise HTTPException(409, f"'{meta.get('title', pid)}' — 본문이 아직 없습니다. 논문을 먼저 한 번 열어주세요.")
        full = "\n\n".join(p["text"] for p in t["pages"])[:budget]
        papers.append({"n": k, "id": pid, "title": meta.get("title", "")})
        parts.append(f"===== [{k}] {meta.get('title', '')} =====\n{full}")
    system = (
        "너는 논문 비교 전문 연구조교다. 주어진 논문들을 비교해 한국어 마크다운으로 답한다. 구조:\n"
        "`## 한눈 비교` — 마크다운 표. 행=항목(문제의식/방법/데이터·실험/핵심 결과(수치 포함)/한계), "
        "열=논문([1] 짧은제목 형식).\n"
        "`## 차이의 핵심` — 불릿 3~5개. 방법론 차이가 결과 차이로 이어지는 지점 중심.\n"
        "`## 종합` — 2~4문장. 어떤 목적이면 어느 논문을 봐야 하는지.\n"
        "전문용어는 한국어(원어) 병기 · 논문에 없는 내용을 지어내지 않기 · 서두·맺음말 금지."
    )
    raw = await ask_claude_async("\n\n".join(parts), system, timeout=300)
    with CMP_LOCK:
        cache = read_json(LIB / "_compare.json", {})
        cache[key] = {"md": raw, "papers": papers, "ts": int(time.time())}
        if len(cache) > 20:      # 오래된 조합부터 정리
            for old in sorted(cache, key=lambda x: cache[x].get("ts", 0))[:len(cache) - 20]:
                del cache[old]
        write_json(LIB / "_compare.json", cache)
    return {"markdown": raw, "papers": papers}


def _ids_hash(ids: list) -> str:
    """정렬된 id 집합의 안정적 해시 — 조합 단위 캐시 파일명에 쓴다(_compare.json과 같은 자리, LIB 바로 아래)."""
    return hashlib.sha256("|".join(sorted(ids)).encode("utf-8")).hexdigest()[:16]


TABLE_COLUMNS = ["대상·문제", "방법", "데이터·실험대상", "주요 결과(수치 포함)", "한계"]
TABLE_LOCK = threading.Lock()


@app.post("/api/library/table")
async def library_table(req: Request):
    """선택한 2~8편에서 리뷰논문용 근거표(대상·문제/방법/데이터·실험대상/주요 결과/한계)를 뽑는다.
    논문별로 기존 요약(summary.md)이 있으면 그것을, 없으면 본문 앞부분을 근거로 준다(compare와
    같은 입력 수집 패턴이지만 전문이 아니라 앞부분만 — 표 칸은 짧게 채우면 되므로 예산을 아낀다).
    같은 조합은 캐시(refresh로 무시)."""
    body = await req.json()
    ids = [i for i in (body.get("ids") or []) if isinstance(i, str)]
    if not 2 <= len(ids) <= 8:
        raise HTTPException(400, "표를 만들 논문을 2~8편 선택하세요")
    cache_path = LIB / f"_table_{_ids_hash(ids)}.json"
    cached = read_json(cache_path, None)
    if cached and not body.get("refresh"):
        return {**cached, "key": _ids_hash(ids), "cached": True}
    budget = (MAX_CTX - 5000) // len(ids)
    parts, papers = [], []
    for k, pid in enumerate(ids, 1):
        d = paper_dir(pid)
        meta = read_json(d / "meta.json", {})
        smp = d / "summary.md"
        if smp.exists():
            src = smp.read_text(encoding="utf-8")[:budget]
        else:
            t = read_json(d / "text.json", None)
            if not t or not t.get("pages"):
                raise HTTPException(409, f"'{meta.get('title', pid)}' — 본문이 아직 없습니다. 논문을 먼저 한 번 열어주세요.")
            src = "\n\n".join(p["text"] for p in t["pages"][:8])[:budget]
        papers.append({"n": k, "id": pid, "title": meta.get("title", ""), "year": meta.get("year")})
        parts.append(f"===== [{k}] {meta.get('title', '')} =====\n{src}")
    system = (
        "너는 리뷰 논문 작성을 돕는 연구 조교다. 주어진 논문들 각각에서 다음 5개 항목의 근거를 뽑아 "
        "표로 만든다: 대상·문제 / 방법 / 데이터·실험대상 / 주요 결과(수치 포함) / 한계.\n"
        "규칙: 칸마다 1~2문장으로 짧게 · 수치는 원문 그대로 보존 · 논문에 없는 내용을 지어내지 않기 · "
        "정보가 없으면 그 칸에 '명시 안 됨'이라고 쓴다(비워두지 않기) · 5개 칸을 이 순서대로 빠짐없이.\n"
        '출력은 JSON 배열 하나만: [{"n": 논문번호, "cells": ["대상·문제", "방법", "데이터·실험대상", '
        '"주요 결과", "한계"]}] (cells는 반드시 5개 원소, 이 순서대로). 다른 텍스트 절대 금지.'
    )
    raw = await ask_claude_async("\n\n".join(parts), system, timeout=300)
    parsed = parse_json_or_502(raw, list)
    by_n = {}
    for it in parsed:
        if not isinstance(it, dict):
            continue
        try:
            n_ = int(it.get("n"))
        except (TypeError, ValueError):
            continue
        cells = it.get("cells")
        if not isinstance(cells, list):
            continue
        cells = [str(c).strip() for c in cells][:len(TABLE_COLUMNS)]
        cells += [""] * (len(TABLE_COLUMNS) - len(cells))
        by_n[n_] = cells
    rows = [{"pid": p["id"], "title": p["title"], "year": p["year"],
             "cells": by_n.get(p["n"], [""] * len(TABLE_COLUMNS))} for p in papers]
    data = {"columns": TABLE_COLUMNS, "rows": rows, "ts": int(time.time())}
    with TABLE_LOCK:
        write_json(cache_path, data)
    return {**data, "key": _ids_hash(ids)}


@app.get("/api/library/tables")
def list_tables():
    """만들어 둔 근거표 목록. 캐시(_table_*.json)는 있는데 꺼낼 길이 없어 사장돼 있던 데이터다 —
    같은 조합을 그대로 다시 고르지 않으면 2분짜리 생성을 다시 물어야 했다."""
    out = []
    for p in LIB.glob("_table_*.json"):
        d = read_json(p, None)
        if not d or not d.get("rows"):
            continue
        out.append({
            "key": p.stem[len("_table_"):],
            "ts": d.get("ts", 0),
            "papers": [{"pid": r.get("pid", ""), "title": r.get("title", ""), "year": r.get("year")}
                       for r in d["rows"]],
        })
    out.sort(key=lambda x: x["ts"], reverse=True)
    return {"tables": out}


@app.get("/api/library/table/{key}")
def get_table(key: str):
    """저장된 근거표 1건. key 는 _ids_hash 출력이므로 16자리 hex 외에는 받지 않는다
    (경로가 파일명에 그대로 들어가므로 traversal 차단이 필요하다)."""
    if not re.fullmatch(r"[0-9a-f]{16}", key):
        raise HTTPException(400, "잘못된 표 식별자입니다")
    d = read_json(LIB / f"_table_{key}.json", None)
    if not d:
        raise HTTPException(404, "저장된 표를 찾을 수 없습니다")
    return {**d, "key": key, "cached": True}


# ---------------------------------------------------------------- metadata (저자·연도·저널·DOI)

# ---------------------------------------------------------------- 근거 보드 (인용 카트)
# 하이라이트·노트·핵심 4색은 논문별 사이드카에만 있어서, 리뷰논문을 쓸 때 "내가 결과라고
# 표시해 둔 문장 전부"를 보려면 편마다 열어야 했다. 여기서 서재 전체를 한 목록으로 세운다.

def _cite_short(meta: dict) -> str:
    """'Kim et al. (2021)' 형태의 짧은 출처. 근거 항목마다 붙어 그대로 인용문에 들어간다."""
    authors = [a for a in (meta.get("authors") or []) if str(a).strip()]
    year = _cite_year(meta)
    if authors:
        last, _ini = _split_author(authors[0])
        who = (last or str(authors[0]).strip()) + (" et al." if len(authors) > 1 else "")
    else:
        who = str(meta.get("title") or "").strip()[:40] or "출처 미상"
    return f"{who} ({year})" if year else who


def _ev_key(pid: str, kind: str, seed) -> str:
    """근거 항목의 안정적인 식별자. 초안의 각 절이 이 키로 근거를 붙들고 있으므로,
    같은 하이라이트·노트·핵심문장은 다시 불러도 같은 키가 나와야 한다.
    하이라이트·노트는 자기 id를, 핵심 4색은 id가 없어 시작 앵커 해시를 쓴다."""
    h = hashlib.sha1(str(seed).encode("utf-8")).hexdigest()[:10]
    return f"{pid}:{kind}:{h}"


def _kp_text(entry: dict, it: dict) -> str:
    """keypoints는 문장의 시작·끝 앵커(s/se)만 저장한다 — 페이지 원문에서 문장을 되살린다.
    저장된 쪽 번호를 먼저 보되 어긋나면 전 페이지를 훑는다(AI가 쪽을 틀리는 경우가 있다)."""
    s = str(it.get("s") or "").strip()
    if not s:
        return ""
    se = str(it.get("se") or "").strip()
    want = int(it.get("p") or 0)
    pages = entry["pages"]
    for _n, orig, _low in [p for p in pages if p[0] == want] + [p for p in pages if p[0] != want]:
        i = orig.find(s)
        if i < 0:
            continue
        if se:
            j = orig.find(se, i)
            if j >= 0:
                return orig[i:j + len(se)]
        return orig[i:i + 300].strip()
    return s          # 본문에서 못 찾으면 시작 앵커라도 — 빈 줄보다 낫다


@app.get("/api/library/evidence")
def library_evidence():
    """서재 전체의 근거를 한 목록으로. 내가 그은 하이라이트 · 내가 쓴 노트 · AI 핵심 4색을
    출처 문자열과 함께 돌려준다. 필터·정렬은 클라이언트가 한다(수백 건 규모라 충분)."""
    papers, items = [], []
    for d in sorted(LIB.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        pid, cite, n0 = d.name, _cite_short(meta), len(items)
        ck = str(meta.get("citekey") or "").strip()   # citekey 인용 모드(config.cite_citekey)용
        for h in read_json(d / "highlights.json", {"items": []}).get("items", []):
            text = str(h.get("text") or "").strip()
            if text:
                items.append({"key": _ev_key(pid, "hl", h.get("id") or text),
                              "pid": pid, "kind": "hl", "src": h.get("source") or "user",
                              "cat": None, "page": int(h.get("page") or 0), "text": text,
                              "memo": str(h.get("reason") or "").strip(), "cite": cite,
                              "citekey": ck})
        for nt in read_json(d / "notes.json", []):
            quote = str(nt.get("quote") or "").strip()
            if quote:
                items.append({"key": _ev_key(pid, "note", nt.get("id") or quote),
                              "pid": pid, "kind": "note", "src": "user", "cat": None,
                              "page": int(nt.get("page") or 0), "text": quote,
                              "memo": str(nt.get("memo") or "").strip(), "cite": cite,
                              "citekey": ck})
        kps = read_json(d / "keypoints.json", {}).get("items") or []
        if kps:
            entry = _search_entry(pid, d)
            for k in kps:
                text = _kp_text(entry, k)
                if text:
                    items.append({"key": _ev_key(pid, "kp", k.get("s") or text),
                                  "pid": pid, "kind": "kp", "src": "ai", "cat": k.get("c"),
                                  "page": int(k.get("p") or 0), "text": text,
                                  "memo": str(k.get("note") or "").strip(), "cite": cite,
                                  "citekey": ck})
        for a in read_json(d / "analysis.json", {}).get("items") or []:
            body = str(a.get("body") or "").strip()
            if body:
                items.append({"key": _ev_key(pid, "an", a.get("key") or body),
                              "pid": pid, "kind": "an", "src": "ai", "cat": a.get("key"),
                              "page": 0, "text": body,
                              "memo": str(a.get("title") or ""), "cite": cite,
                              "citekey": ck})
        if len(items) > n0:
            papers.append({"id": pid, "title": meta.get("title") or pid, "cite": cite,
                           "year": _cite_year(meta), "count": len(items) - n0})
    return {"papers": papers, "items": items}


# ---------------------------------------------------------------- 초안 (집필 탭)
# 절 단위 초안과 "이 절에 담은 근거"를 한 파일에 둔다. Achird는 산문 편집기가 아니다 —
# 서식·문단 편집은 워드프로세서가 이긴다. 여기가 소유하는 건 근거와 그 출처뿐이고,
# 초안은 겹침 검사·인용 삽입을 반복하기 위한 작업 사본이다.

DRAFT = LIB / "_draft.json"
DRAFT_MAX_TEXT = 60000
DRAFT_MAX_SECTIONS = 30


def _clean_draft(body: dict) -> dict:
    """저장 전 정규화(순수 함수). 절 id는 클라이언트가 만들지만 비었거나 겹쳐 들어오면 여기서
    바로잡는다 — id가 겹치면 '근거 담기'가 엉뚱한 절에 붙는다."""
    secs, seen = [], set()
    for i, s in enumerate(body.get("sections") or []):
        if not isinstance(s, dict) or len(secs) >= DRAFT_MAX_SECTIONS:
            continue
        sid = str(s.get("id") or "").strip()[:32] or f"s{i}"
        while sid in seen:
            sid += "_"
        seen.add(sid)
        ev, evseen = [], set()
        for k in (s.get("ev") or [])[:500]:
            k = str(k).strip()[:120]
            if k and k not in evseen:
                evseen.add(k)
                ev.append(k)
        secs.append({"id": sid, "name": str(s.get("name") or "").strip()[:80] or "이름 없는 절",
                     "text": str(s.get("text") or "")[:DRAFT_MAX_TEXT], "ev": ev})
    return {"title": str(body.get("title") or "").strip()[:200], "sections": secs}


DRAFT_LOCK = threading.Lock()   # put_draft(이벤트 루프)와 _prune_draft_refs(threadpool)가 교차한다


def _prune_draft_refs(pid: str) -> int:
    """논문을 지우면 초안이 붙들고 있던 그 논문 근거 키도 같이 지운다. 키는 순수 포인터라
    (본문은 지워진 논문 폴더에 있었다) 남겨봐야 '못 찾은 근거' 숫자만 부풀린다.
    근거표 캐시(_table_*.json)는 반대로 남긴다 — 거기엔 추출된 본문이 들어 있어 지우면 작업이 날아간다."""
    with DRAFT_LOCK:
        d = read_json(DRAFT, None)
        if not d or not d.get("sections"):
            return 0
        pre, dropped = f"{pid}:", 0
        for s in d["sections"]:
            ev = s.get("ev") or []
            keep = [k for k in ev if not str(k).startswith(pre)]
            dropped += len(ev) - len(keep)
            s["ev"] = keep
        if dropped:
            write_json(DRAFT, d)
        return dropped


@app.get("/api/draft")
def get_draft():
    return read_json(DRAFT, None) or {"title": "", "sections": [], "ts": 0}


@app.put("/api/draft")
async def put_draft(req: Request):
    """초안 전체를 통째로 저장한다. 절이 수십 개, 본문이 수십 KB 규모라 부분 갱신 프로토콜을
    만들 값어치가 없다 — 클라이언트가 디바운스해서 보낸다."""
    body = await req.json()
    if not isinstance(body, dict):
        raise HTTPException(400, "초안 객체가 필요합니다")
    data = {**_clean_draft(body), "ts": int(time.time())}
    with DRAFT_LOCK:        # 삭제의 근거 정리(_prune_draft_refs, threadpool)와 저장이 겹치면 한쪽이 날아간다
        write_json(DRAFT, data)
    return data


# ---------------------------------------------------------------- 초안 중복 자기점검
# 요약하다 원문을 그대로 옮긴 곳을 스스로 찾는 장치. AI도 네트워크도 쓰지 않는 순수 문자열 비교라
# 오탐·환각이 없고 즉시 끝난다. 한국어 초안은 번역문과, 영어 초안은 원문과 걸린다.

OVERLAP_K = 8          # 겹침으로 볼 최소 연속 어절. 8이면 흔한 상투구는 안 걸리고 옮겨쓴 문장은 걸린다
OVERLAP_MAX = 40000    # 초안 입력 상한(자) — 한 섹션씩 검사하는 용도
OVERLAP_TOP = 60       # 화면에 돌려줄 상위 구간 수 (긴 것부터). 전체 개수는 spots 로 따로 알린다


def _merge_spans(spans: list) -> list:
    """겹치거나 맞닿은 [a,b) 구간을 합친다."""
    out: list = []
    for a, b in sorted(spans):
        if out and a <= out[-1][1]:
            out[-1][1] = max(out[-1][1], b)
        else:
            out.append([a, b])
    return [(a, b) for a, b in out]


def _find_overlaps(draft: str, k: int = OVERLAP_K) -> dict:
    """초안과 서재(원문·번역)의 k-어절 이상 겹침. 큰 색인을 세우지 않고 초안 쪽 지문만
    사전에 담아 서재를 한 번 훑는다 — 메모리가 서재 크기가 아니라 초안 길이에만 비례한다."""
    # 대조는 소문자로, 화면에 돌려줄 발췌는 초안에 쓴 대소문자 그대로 — 같은 토큰화라 색인이 맞는다
    words = _norm_search(draft).split()
    dw = [w.lower() for w in words]
    if len(dw) < k:
        return {"words": len(dw), "matched": 0, "hits": []}
    want: dict = {}
    for i in range(len(dw) - k + 1):
        want.setdefault(" ".join(dw[i:i + k]), []).append(i)
    runs: dict = {}                     # (pid, 소스, 쪽) → 초안 어절 구간들
    titles: dict = {}
    for d in sorted(LIB.iterdir()):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        titles[d.name] = meta.get("title") or d.name
        e = _search_entry(d.name, d)
        for src, plist in (("원문", e["pages"]), ("번역", e["tpages"])):
            for n, _orig, low in plist:
                lw = low.split()
                for j in range(len(lw) - k + 1):
                    for i in want.get(" ".join(lw[j:j + k]), ()):
                        runs.setdefault((d.name, src, n), []).append((i, i + k))
    hits = []
    for (pid, src, page), spans in runs.items():
        for a, b in _merge_spans(spans):
            hits.append({"pid": pid, "title": titles.get(pid, pid), "src": src, "page": page,
                         "start": a, "words": b - a, "text": " ".join(words[a:b])})
    hits.sort(key=lambda h: (h["words"], -h["start"]), reverse=True)
    covered = _merge_spans([(h["start"], h["start"] + h["words"]) for h in hits])
    # spots 는 자르기 전의 진짜 개수다 — hits 만 세면 60에서 멈춰 "60곳"이라 보고하게 된다
    return {"words": len(dw), "matched": sum(b - a for a, b in covered),
            "spots": len(hits), "hits": hits[:OVERLAP_TOP]}


@app.post("/api/library/overlap")
async def library_overlap(req: Request):
    """초안을 붙여넣으면 서재 원문·번역과 겹치는 구간을 돌려준다. 겹침이 곧 표절은 아니다 —
    인용부호를 빠뜨렸거나 요약이 아직 원문에 붙어 있는 곳을 스스로 보라는 용도다."""
    try:
        body = await req.json()
    except Exception as e:      # 본문이 필수다 — 빈 dict 로 넘기면 아래 길이 검증이 실패를 자기 원인으로 착각한다
        raise HTTPException(400, f"요청 본문을 읽지 못했습니다: {type(e).__name__}: {e}")
    text = str(body.get("text") or "")[:OVERLAP_MAX]
    try:
        k = int(body.get("k") or OVERLAP_K)
    except (TypeError, ValueError):
        k = OVERLAP_K
    k = max(4, min(20, k))
    if len(_norm_search(text).split()) < k:
        raise HTTPException(400, f"{k}어절 이상 붙여넣어야 검사할 수 있습니다")
    # 서재 전체를 어절 단위로 훑는 CPU 작업이라 이벤트 루프에서 돌리면 안 된다 —
    # 실측(논문 4편·6000어절): 루프에서 돌 때 동시에 들어온 가벼운 GET 이 63ms → 1389ms 로 밀렸다.
    # 진행 중이던 번역·요약이 그 시간만큼 통째로 멈춘다.
    return {**await asyncio.to_thread(_find_overlaps, text, k), "k": k}


# ---------------------------------------------------------------- 초안 채점 (4축 + 가상 피어리뷰)
# 겹침 검사가 "옮겨 쓴 곳"을 잡는다면 여기는 "논증이 서 있는지"를 본다.
# 문장을 고쳐 쓰지는 않는다 — Achird는 산문을 소유하지 않는다. 지적까지가 여기 몫이고
# 고치는 건 사람과 워드프로세서 몫이다. 네 축 중 셋(근거·인용·용어)은 AI 없이 결정론으로 끝나고,
# 구조·피어리뷰만 모델을 부른다. AI가 죽어도 결정론 축은 그대로 나온다.

REVIEW_MAX = 20000

_CITE_AY_RE = re.compile(r"\(([^()]{2,80}?),\s*(1[89]\d{2}|20\d{2})(?:\s*,\s*p\.?\s*\d+)?\)")
_CITE_NUM_RE = re.compile(r"\[\s*\d{1,3}(?:\s*[,;–\-]\s*\d{1,3})*\s*\]")
# pandoc 식 [@citekey] / [@citekey, p.4] — citekey 인용 모드(config.cite_citekey)가 만드는 형태
_CITE_KEY_RE = re.compile(r"\[@([A-Za-z0-9_.:\-]{2,80})(?:\s*,\s*p{1,2}\.?\s*(\d{1,4}))?\]")
_SENT_RE = re.compile(r"[^.!?。？！\n]+[.!?。？！]|[^\n]+")
_ABBREV_RE = re.compile(r"\b(?:et\s+al|e\.g|i\.e|cf|vs|Fig|Eq|Ref|No|pp?|approx|etc)\.", re.I)
_DECIMAL_RE = re.compile(r"(?<=\d)\.(?=\d)")


def _cite_marks(text) -> dict:
    """본문의 인용 표기를 센다. (저자, 연도)·[n]·[@citekey] 세 꼴 — 초안은 섞어 쓴다."""
    t = str(text or "")
    ay, num = len(_CITE_AY_RE.findall(t)), len(_CITE_NUM_RE.findall(t))
    ck = len(_CITE_KEY_RE.findall(t))
    return {"authoryear": ay, "numeric": num, "citekey": ck, "total": ay + num + ck}


def _sentences(text) -> list:
    """문장 단위로 자른다. 마침표가 없는 줄(제목·목록 항목)도 한 문장으로 센다 —
    인용 밀도의 분모라서 빠뜨리면 밀도가 실제보다 높게 나온다.

    'et al.'·'e.g.'·소수점의 마침표는 문장 끝이 아니다. 안 막으면 인용 표기가 많은 문단일수록
    문장 수가 부풀어 밀도가 낮게 나온다 — 인용을 잘 단 글이 벌점을 받는 꼴이 된다."""
    t = str(text or "")
    masked = _DECIMAL_RE.sub("\x00", _ABBREV_RE.sub(lambda m: m.group(0).replace(".", "\x00"), t))
    return [s.replace("\x00", ".").strip() for s in _SENT_RE.findall(masked) if len(s.strip()) > 1]


def _cite_parts(cite) -> tuple:
    """'Kim et al. (2021)' → ('kim', '2021'). 성과 연도만 남긴다 — 본문 표기가
    APA든 ACS든 번호식이든 이 둘은 살아남는다."""
    s = str(cite or "").strip()
    m = re.match(r"(.+?)\s*\((\d{4})\)\s*$", s)
    who, year = (m.group(1), m.group(2)) if m else (s, "")
    who = who.replace("et al.", " ").replace(",", " ").strip().lower()
    return (who.split()[0] if who.split() else ""), year


def _uncited_evidence(text, evidence: list) -> list:
    """절에 담아놓고 본문에는 안 쓴 근거. 초안 점검에서 가장 흔한 구멍이다.
    성과 연도가 둘 다 본문에 있으면 인용된 것으로 본다 — 표기 형태가 제각각이라
    출처 문자열 완전일치로 보면 전부 미인용으로 잡힌다."""
    low = str(text or "").lower()
    out = []
    for e in evidence:
        if not isinstance(e, dict):
            continue
        cite = str(e.get("cite") or "").strip()
        if not cite:
            continue
        ck = str(e.get("citekey") or "").strip()
        if ck and f"@{ck.lower()}" in low:      # citekey 인용 모드로 쓴 본문
            continue
        who, year = _cite_parts(cite)
        hit = bool(who) and who in low and (not year or year in low)
        if not hit:
            out.append({"cite": cite, "text": str(e.get("text") or "")[:200]})
    return out


def _term_conflicts(text, canon: dict, alts: dict) -> list:
    """대표 역어를 고정해 둔 용어인데 본문이 다른 역어를 쓰고 있는 곳.
    canon: {소문자 원어: 대표 역어} · alts: {소문자 원어: [서재에서 쓰이는 역어들]}"""
    t = str(text or "")
    out = []
    for key, want in sorted(canon.items()):
        if not want:
            continue
        used = [k for k in alts.get(key, []) if k and k != want and k in t]
        if used:
            out.append({"term": key, "canon": want, "used": used, "fixed": want in t})
    return out


# ---- 인용 무결성: 겹침 검사의 역방향.
# 겹침 검사는 "본문이 원문을 옮겼는가"를 보고, 여기는 "본문이 가리킨 자리에 근거가 있는가"를 본다.
# 서재에 없는 출처를 인용하는 건 정상이다(PDF 없이 인용하는 논문이 늘 있다) — 오류가 아니라
# 정보로만 표시한다. 진짜 오류는 쪽수를 넘긴 인용과, 담은 근거가 없는 쪽을 가리킨 인용이다.

_INLINE_CITE_RE = re.compile(
    r"\(([^()]{2,80}?),\s*(1[89]\d{2}|20\d{2})(?:\s*,\s*p{1,2}\.?\s*(\d{1,4}))?\)")


def _parse_inline_cites(text) -> list:
    """본문의 (저자, 연도[, p.N]) 와 [@citekey[, p.N]] 를 뽑는다. 번호식 [n] 은 대상이
    아니다 — 번호는 원고 밖의 목록을 가리켜서 서재만으로는 맞는지 알 수 없다."""
    t = str(text or "")
    out = []
    for m in _INLINE_CITE_RE.finditer(t):
        who = m.group(1).replace("et al.", " ").replace(",", " ").strip().lower()
        out.append({"raw": m.group(0), "who": who.split()[0] if who.split() else "",
                    "year": m.group(2), "page": int(m.group(3)) if m.group(3) else 0})
    for m in _CITE_KEY_RE.finditer(t):
        out.append({"raw": m.group(0), "who": "", "year": "", "citekey": m.group(1),
                    "page": int(m.group(2)) if m.group(2) else 0})
    return out


def _lib_cite_index() -> list:
    """[{pid, who, year, pages, title}] — 초안의 (저자, 연도) 를 서재 논문으로 되짚는 색인."""
    out = []
    for d in sorted(LIB.iterdir()):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        who, year = _cite_parts(_cite_short(meta))
        out.append({"pid": d.name, "who": who, "year": year,
                    "citekey": str(meta.get("citekey") or "").strip(),
                    "pages": int(meta.get("pages") or 0), "title": meta.get("title") or d.name})
    return out


def _check_inline_cites(text, index: list, ev_pages: set) -> list:
    """본문 인용마다 판정 하나. ev_pages 는 이 절이 담은 근거의 (pid, 쪽) 집합 —
    비어 있으면(옛 클라이언트가 쪽을 안 보낸 경우) 근거 대조는 건너뛴다."""
    out = []
    for c in _parse_inline_cites(text):
        if c.get("citekey"):
            hit = next((p for p in index if p.get("citekey") and p["citekey"] == c["citekey"]), None)
        else:
            hit = next((p for p in index if p["who"] and p["who"] == c["who"] and p["year"] == c["year"]), None)
        if not hit:
            out.append({**c, "kind": "outside", "title": ""})
            continue
        if c["page"] and hit["pages"] and c["page"] > hit["pages"]:
            out.append({**c, "kind": "page_range", "title": hit["title"], "pages": hit["pages"]})
        elif c["page"] and ev_pages and (hit["pid"], c["page"]) not in ev_pages:
            out.append({**c, "kind": "no_evidence", "title": hit["title"]})
    return out


def _review_checks(text: str, evidence: list, canon: dict, alts: dict, index: list = None) -> dict:
    """AI를 쓰지 않는 네 축(근거·인용·용어·인용 무결성) + 분량. 순수 함수라 회귀 테스트로 지킨다."""
    marks, sents = _cite_marks(text), _sentences(text)
    uncited = _uncited_evidence(text, evidence)
    ev_pages = {(e.get("pid"), int(e.get("page") or 0)) for e in evidence
                if isinstance(e, dict) and e.get("pid") and e.get("page")}
    inline = _check_inline_cites(text, index or [], ev_pages)
    return {
        "chars": len(text), "sentences": len(sents),
        "cites": marks, "density": round(marks["total"] / max(1, len(sents)), 2),
        "evidence_total": len([e for e in evidence if isinstance(e, dict)]),
        "uncited": uncited, "terms": _term_conflicts(text, canon, alts),
        "inline": inline,
        "inline_bad": sum(1 for i in inline if i["kind"] != "outside"),
    }


@app.post("/api/draft/review")
async def draft_review(req: Request):
    """절 하나를 네 축으로 점검한다. 결정론 축은 즉시, 구조·피어리뷰는 모델 한 번.
    모델이 실패해도 ai:null 로 돌려주고 결정론 결과는 그대로 살린다 — 셋은 이미 끝나 있다."""
    try:
        body = await req.json()
    except Exception as e:      # 본문이 필수다 — 빈 dict 로 넘기면 아래 길이 검증이 실패를 자기 원인으로 착각한다
        raise HTTPException(400, f"요청 본문을 읽지 못했습니다: {type(e).__name__}: {e}")
    name = str(body.get("name") or "이 절").strip()[:80]
    text = str(body.get("text") or "")[:REVIEW_MAX]
    if len(text.strip()) < 100:
        raise HTTPException(400, "100자 이상 써야 점검할 수 있습니다")
    evidence = [e for e in (body.get("evidence") or [])[:200] if isinstance(e, dict)]
    gl = library_glossary()
    canon = {t["key"]: t["canon"] for t in gl["terms"] if t["canon"]}
    alts = {t["key"]: [k["ko"] for k in t["kos"]] for t in gl["terms"]}
    checks = _review_checks(text, evidence, canon, alts, _lib_cite_index())

    ev_lines = "\n".join(
        f"- {e.get('cite','')}: {str(e.get('text',''))[:300]}" for e in evidence[:40]) or "(담은 근거 없음)"
    system = (
        "학술 원고의 한 절을 심사하는 도구다. 출력은 JSON 객체 하나만, 다른 텍스트 금지. "
        '형식: {"structure":[{"where":"문제가 있는 자리(원문 일부 인용)","issue":"무엇이 문제인지",'
        '"fix":"어떻게 손볼지"}], "reviewers":[{"who":"역할","verdict":"accept|minor|major|reject",'
        '"points":["지적 1","지적 2"]}]}. '
        "reviewers 는 정확히 3명: '엄격한 심사자'(주장이 근거보다 앞서는 곳), "
        "'방법론 심사자'(근거가 주장을 실제로 받치는지), '분야 밖 독자'(용어·연결이 따라가지는지). "
        "structure 는 최대 6건. 고쳐 쓴 문장을 내놓지 마라 — 어디가 왜 문제인지만 적는다. "
        "칭찬·총평·맺음말 금지. 한국어."
    )
    prompt = (f"===== 절 이름 =====\n{name}\n\n===== 절 본문 =====\n{text}\n\n"
              f"===== 글쓴이가 이 절에 담아둔 근거 =====\n{ev_lines}")
    ai, ai_error = None, None
    try:
        raw = await ask_claude_async(prompt, system)
        data = parse_json_or_502(raw, dict)
        ai = {"structure": [s for s in (data.get("structure") or [])[:6] if isinstance(s, dict)],
              "reviewers": [r for r in (data.get("reviewers") or [])[:3] if isinstance(r, dict)]}
    except HTTPException as e:
        ai_error = e.detail          # 502 를 올리지 않고 담아만 둔다 — 결정론 축 셋은 이미 끝나 있다
    return {"name": name, "checks": checks, "ai": ai, "ai_error": ai_error}


# ---------------------------------------------------------------- 서재 통합 용어집
# 용어집이 논문별로 따로라 같은 원어가 논문마다 다르게 번역된다. 리뷰논문에서는 그게 그대로
# 문단마다 다른 용어로 드러난다. 여기서 충돌을 보여주고 대표 역어를 하나로 고정한다.

LIB_GLOSS = LIB / "_glossary.json"      # {"terms": {소문자 원어: 대표 역어}, "ts": ...}


def _canon_terms() -> dict:
    t = read_json(LIB_GLOSS, {}).get("terms")
    return t if isinstance(t, dict) else {}


@app.get("/api/library/glossary")
def library_glossary():
    """논문별 용어집을 원어 기준으로 합친다. 역어가 둘 이상이면 conflict=True."""
    canon = _canon_terms()
    merged: dict = {}
    for d in sorted(LIB.iterdir()):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", None)
        if not meta:
            continue
        for i in read_json(d / "glossary.json", {}).get("items") or []:
            term, ko = str(i.get("term") or "").strip(), str(i.get("ko") or "").strip()
            if not term:
                continue
            g = merged.setdefault(term.lower(), {"key": term.lower(), "term": term, "kos": {},
                                                 "papers": [], "def": ""})
            if ko:
                g["kos"].setdefault(ko, []).append(d.name)
            if d.name not in g["papers"]:
                g["papers"].append(d.name)
            if not g["def"]:
                g["def"] = str(i.get("def") or "").strip()
    out = []
    for g in merged.values():
        kos = [{"ko": k, "papers": v} for k, v in sorted(g["kos"].items(), key=lambda x: -len(x[1]))]
        out.append({"key": g["key"], "term": g["term"], "def": g["def"], "kos": kos,
                    "papers": len(g["papers"]), "conflict": len(kos) > 1,
                    "canon": canon.get(g["key"], "")})
    # 충돌 먼저, 그 다음 여러 논문에 걸친 용어 먼저 — 손볼 값어치 순
    out.sort(key=lambda g: (g["conflict"], g["papers"]), reverse=True)
    return {"terms": out, "canon_count": sum(1 for g in out if g["canon"])}


@app.put("/api/library/glossary")
async def put_library_glossary(req: Request):
    """대표 역어를 저장한다. 빈 문자열은 고정 해제. 이후 모든 페이지 번역이 이 역어를 우선한다."""
    body = await req.json()
    terms = body.get("terms")
    if not isinstance(terms, dict):
        raise HTTPException(400, "terms 객체가 필요합니다")
    clean = {str(k).strip().lower()[:200]: str(v).strip()[:200]
             for k, v in list(terms.items())[:2000] if str(k).strip() and str(v).strip()}
    write_json(LIB_GLOSS, {"terms": clean, "ts": int(time.time())})
    return {"terms": clean, "count": len(clean)}


# ---------------------------------------------------------------- 읽을 논문 대기열
# 심화 검색에서 좋은 논문을 찾아도 담아둘 데가 없어 브라우저 탭으로 흘려보내던 것을 붙잡는다.
# PDF는 저작권 때문에 자동으로 못 가져온다 — 서지정보만 담고, 나중에 그 PDF를 올리면 매칭된다.

QUEUE = LIB / "_queue.json"


@app.get("/api/queue")
def get_queue():
    index = _lib_index()
    # 대기열 논문의 PDF 를 Zotero 로 받아오는 경로가 흔하다 — DOI 로 매칭해 "도착"을 알리고
    # 그 자리에서 가져오게 한다. Zotero 미사용이면 조용히 빈 매핑(배지와 같은 원칙).
    try:
        zot = {str(s.get("doi") or "").strip().lower().replace("https://doi.org/", ""): s["att_key"]
               for s in scan_zotero_cached(load_config())
               if s.get("doi") and s.get("resolved") and s.get("att_key")}
    except Exception:                # noqa: BLE001
        zot = {}
    imported = _lib_zotero_att_keys() if zot else set()
    items = []
    for it in read_json(QUEUE, []):
        lib = _lib_hit(str(it.get("doi") or ""), str(it.get("title") or ""), index)
        ak = zot.get(str(it.get("doi") or "").strip().lower())
        items.append({**it, "in_library": lib,
                      "in_zotero": ak if (ak and not lib and ak not in imported) else None})
    return {"items": items, "arrived": sum(1 for i in items if i["in_library"])}


@app.post("/api/queue")
async def add_queue(req: Request):
    """대기열에 한 건 추가. doi가 같으면 덮어쓰지 않고 그대로 둔다(중복 담기 방지).

    read-modify-write 사이에 await 가 없어 같은 이벤트 루프의 다른 요청이 끼어들 수 없다
    (실측: 동시 POST 10건 전부 보존). 이 함수를 sync def 로 바꾸면 스레드풀에서 진짜로
    병렬 실행돼 유실이 생긴다 — 바꿀 거면 락을 함께 넣을 것."""
    body = await req.json()
    title = str(body.get("title") or "").strip()[:400]
    doi = str(body.get("doi") or "").strip().lower().replace("https://doi.org/", "")[:200]
    if not title and not doi:
        raise HTTPException(400, "제목이나 DOI 중 하나는 있어야 합니다")
    items = read_json(QUEUE, [])
    if doi and any(str(i.get("doi") or "").lower() == doi for i in items):
        raise HTTPException(409, "이미 대기열에 있습니다")
    if not doi and any(_alnum_sig(i.get("title")) == _alnum_sig(title) for i in items):
        raise HTTPException(409, "이미 대기열에 있습니다")
    item = {"id": uuid.uuid4().hex[:8], "title": title, "doi": doi,
            "url": str(body.get("url") or "").strip()[:500],
            "venue": str(body.get("venue") or "").strip()[:200],
            "year": body.get("year") if isinstance(body.get("year"), int) else None,
            "note": str(body.get("note") or "").strip()[:500],
            "added": int(time.time())}
    items.insert(0, item)
    write_json(QUEUE, items[:500])
    return item


@app.delete("/api/queue/{qid}")
def del_queue(qid: str):
    items = read_json(QUEUE, [])
    left = [i for i in items if i.get("id") != qid]
    if len(left) == len(items):
        raise HTTPException(404, "대기열에 없는 항목입니다")
    write_json(QUEUE, left)
    return {"ok": True, "count": len(left)}


@app.post("/api/papers/{pid}/metadata")
async def extract_metadata(pid: str):
    """논문 앞부분에서 서지정보 추출(haiku). 서가 표시·Obsidian frontmatter·인용 연결의 기반.
    결과가 비어도 meta_ai 스탬프를 남겨 재시도 루프를 막되, 한 필드도 못 건지면
    meta_ai_error 로 원인을 함께 돌려준다(AI 호출 자체의 실패는 502 로 올라간다)."""
    d = paper_dir(pid)
    t = read_json(d / "text.json", None)
    if not t or not t.get("pages"):
        raise HTTPException(409, _no_text(d))
    head = "\n".join(p["text"] for p in t["pages"][:2])[:6000]
    system = (
        "논문 첫 부분 텍스트에서 서지정보를 추출하는 도구다. "
        '출력은 JSON 하나만: {"authors": ["표기된 그대로의 저자명", ...], "year": 2017, '
        '"venue": "학회/저널명", "doi": "10.xxxx/..."} '
        "모르는 필드는 null. 저자는 본문 표기 순서대로 전부. 다른 텍스트 절대 금지."
    )
    # AI 호출 자체의 실패(502/504)는 올려보낸다 — 삼키면 meta_ai 스탬프가 남아 재시도가 영구히 닫힌다
    raw = await ask_claude_async(f"===== 논문 앞부분 =====\n{head}", system, timeout=120, model="haiku")
    err = None
    try:
        parsed = extract_json(raw)
        if not isinstance(parsed, dict):
            raise ValueError(f"object expected, got {type(parsed).__name__}")
    except (ValueError, json.JSONDecodeError) as e:      # 파싱 실패는 흡수 — 정규식 DOI 폴백이 남았다
        parsed = {}
        err = f"AI 응답을 해석하지 못했습니다({type(e).__name__}). 받은 내용 앞부분: {raw[:200]!r}"
    authors = [str(a).strip() for a in (parsed.get("authors") or []) if str(a).strip()][:20]
    year = parsed.get("year")
    try:
        year = int(year) if year and 1900 <= int(year) <= 2100 else None
    except (TypeError, ValueError):
        year = None
    venue = str(parsed.get("venue") or "").strip()[:200] or None
    doi = str(parsed.get("doi") or "").strip()[:200] or None
    if not doi:                                  # AI가 놓치면 정규식 폴백
        m = re.search(r"10\.\d{4,9}/[-._;()/:a-zA-Z0-9]+", head)
        doi = m.group(0).rstrip(".,;") if m else None
    with META_LOCK:
        meta = read_json(d / "meta.json", {})
        if authors:
            meta["authors"] = authors
        if year:
            meta["year"] = year
        if venue:
            meta["venue"] = venue
        if doi:
            meta["doi"] = doi
        meta["meta_ai"] = int(time.time())
        write_json(d / "meta.json", meta)
    if err and not (authors or year or venue or doi):     # 한 필드도 못 건졌으면 200 이어도 원인을 말해야 한다
        return {**meta, "meta_ai_error": err}
    return meta


# ---------------------------------------------------------------- suggested questions (추천 질문)

@app.get("/api/papers/{pid}/questions")
def get_questions(pid: str):
    return read_json(paper_dir(pid) / "questions.json", {"items": None})


@app.post("/api/papers/{pid}/questions")
async def make_questions(pid: str, req: Request):
    """채팅 진입장벽을 낮추는 추천 질문 3개. 논문당 1회 생성·캐시(force로 재생성)."""
    d = paper_dir(pid)
    try:
        body = await req.json()
    except Exception:
        body = {}
    cached = read_json(d / "questions.json", None)
    if cached and cached.get("items") and not body.get("force"):
        return cached
    text = full_text(pid)[:MAX_CTX]
    meta = read_json(d / "meta.json", {})
    system = (
        "너는 논문 읽기를 돕는 조교다. 이 논문을 읽는 대학원생이 AI에게 물어보면 좋을 질문 "
        "3개를 한국어로 만든다. 구성: ① 핵심 개념·방법의 이해를 확인하는 질문 ② 결과·주장의 "
        "근거를 파고드는 질문 ③ 한계·응용·후속 연구로 잇는 질문. 각 질문은 이 논문 내용에 "
        "구체적으로 근거해야 하며(일반론 금지) 한 문장으로 짧게. "
        '출력은 JSON 배열 하나만: ["질문1", "질문2", "질문3"]. 다른 텍스트 절대 금지.'
    )
    raw = await ask_claude_async(f"논문 제목: {meta.get('title', '')}\n\n===== 논문 전문 =====\n{text}", system)
    parsed = parse_json_or_502(raw, list)
    items = [str(x).strip() for x in parsed if str(x).strip()][:4]
    data = {"items": items, "ts": int(time.time())}
    write_json(d / "questions.json", data)
    return data


# ---------------------------------------------------------------- full translation

TRANS_LOCK = threading.Lock()
ALIGN_MARKER = "===ALIGN==="


def split_translation(out: str):
    """번역 응답을 (마크다운, 앵커쌍[]) 으로 분리한다. 번역 호출이 문장 정렬 앵커를
    함께 내므로 정렬 전용 AI 호출이 필요 없다(근본적 속도 해결). 마커·JSON이 없으면 []."""
    idx = out.find(ALIGN_MARKER)
    if idx == -1:
        return out.strip(), []
    md = out[:idx].strip()
    try:
        parsed = extract_json(out[idx + len(ALIGN_MARKER):])
        pairs = ([_anchor(p) for p in parsed if isinstance(p, dict) and p.get("s") and p.get("t")]
                 if isinstance(parsed, list) else [])
    except (ValueError, json.JSONDecodeError):
        pairs = []
    return md, pairs


def _anchor(p: dict) -> dict:
    """앵커쌍 정제: 시작(s,t)은 필수, 끝(se,te)은 있으면 담고 없으면 빈 문자열."""
    return {"s": str(p.get("s", "")).strip(), "se": str(p.get("se", "")).strip(),
            "t": str(p.get("t", "")).strip(), "te": str(p.get("te", "")).strip()}


@app.get("/api/papers/{pid}/translation")
def get_translation(pid: str):
    return read_json(paper_dir(pid) / "translation.json", {"pages": {}})


@app.post("/api/papers/{pid}/translation/page")
async def translate_page(pid: str, req: Request):
    """페이지 1장 번역. 클라이언트가 페이지 루프를 돌며 호출한다 — 긴 논문도
    호출당 출력이 작아 타임아웃·출력한도에 안전하고, 페이지 단위로 캐시된다."""
    body = await req.json()
    return await _translate_one(pid, int(body.get("n") or 0), bool(body.get("force")))


async def _translate_one(pid: str, n: int, force: bool = False) -> dict:
    """번역 한 장의 알맹이. 엔드포인트와 백그라운드 잡이 함께 쓴다 — 잡이 HTTP 를 거쳐
    자기 서버를 다시 부르면 세마포어 밖에서 도는 호출이 생겨 동시성 상한이 무너진다."""
    d = paper_dir(pid)
    t = read_json(d / "text.json", None)
    if not t or not t.get("pages"):
        raise HTTPException(409, _no_text(d))
    pages = t["pages"]
    if n < 1 or n > len(pages):
        raise HTTPException(400, "잘못된 페이지 번호")
    cached = read_json(d / "translation.json", {"pages": {}})
    if str(n) in cached["pages"] and not force:
        return {"n": n, "markdown": cached["pages"][str(n)], "cached": True}

    src = pages[n - 1]["text"]
    if not src.strip():
        out, pairs = "", []
    else:
        meta = read_json(d / "meta.json", {})
        prev_tail = pages[n - 2]["text"][-600:] if n >= 2 else ""
        next_head = pages[n]["text"][:600] if n < len(pages) else ""
        system = (
            "너는 학술 논문 전문 번역가다. 주어진 원문 페이지 전체를 자연스러운 한국어 "
            "마크다운으로 번역한다. 규칙: 문단 구조 유지 · 섹션 제목은 `##` 헤딩 · "
            "수식·기호·변수·수치는 원문 그대로 · 표는 마크다운 표로 재구성 · "
            "참고문헌 목록(서지)은 번역하지 않고 원문 그대로 옮김 · 핵심 전문용어는 첫 "
            "등장에만 '한국어(원어)' 병기 · 머리글/페이지 번호 같은 잡음은 버림. 서두·역주 금지.\n"
            "[절대 규칙] 너는 오직 번역만 출력한다. 페이지에 서술 본문이 없고 그림·캡션·표·메타데이터뿐이어도 "
            "사용자에게 말을 걸거나 상황을 설명하지 마라 — '죄송', '여쭙니다', '확인차', '없어 보입니다', "
            "'제시된 상태', '다음 페이지' 같은 메타 발언을 절대 쓰지 마라. 페이지에 있는 번역 가능한 "
            "텍스트(그림 캡션·표 포함)만 번역하고, 번역할 텍스트가 전혀 없으면 설명 없이 정렬 마커와 "
            "빈 배열 `[]`만 출력한다.\n"
            "[페이지 경계 문장] 경계에 걸친 문장은 더 많이 포함된 페이지에서만 온전히 번역한다. "
            "이 페이지 첫머리가 이전 페이지에서 시작된 문장의 뒷부분이면: 그 문장이 이 페이지에 더 "
            "많으면 이전 페이지 조각(맥락)까지 합쳐 완결 문장으로 번역하고, 이전 페이지에 더 많으면 "
            "이 조각은 건너뛴다. 이 페이지 끝이 다음 페이지로 이어지는 미완 문장이면: 이 페이지에 더 "
            "많으면 다음 페이지 시작 조각까지 합쳐 완결 문장으로 번역하고, 다음 페이지에 더 많으면 "
            "이 미완 조각은 남기지 않는다. 목표: 경계 문장이 양쪽에 미완·중복으로 나오지 않게.\n"
            "[각주] 페이지 하단 각주(footnote)가 있으면 본문에 섞지 말고, 번역 맨 끝(정렬 마커 앞)에 "
            "`## 각주` 헤딩으로 모아 번역한다(번호·기호 유지). 각주가 없으면 이 섹션을 만들지 않는다.\n"
            "번역문(각주 포함)을 먼저 출력하고, 그 뒤에 정확히 다음 형식으로 문장 정렬을 덧붙인다:\n"
            f"{ALIGN_MARKER}\n"
            '[{"s": "원문 문장 시작 5~8단어", "se": "원문 문장 끝 4~6단어", '
            '"t": "번역 문장 시작 5~8단어", "te": "번역 문장 끝 4~6단어"}]\n'
            "정렬 규칙: 본문·그림 캡션의 각 문장마다 원문(s,se)·번역(t,te) 앵커 · 모두 그대로(verbatim) · "
            "s/t는 문장 시작, se/te는 그 문장의 진짜 마지막(문장부호 직전, 다음 문장·제목·각주 절대 미포함) · "
            "t/te는 위 번역문에 실제로 나온 그대로 · 원문 등장 순서대로 · 각 앵커는 유일하게 찾을 만큼 "
            "길게 · 마크다운 기호(#,*,|,`) 제외 · 제목·표·수식·참고문헌·각주만인 조각은 건너뜀."
        )
        # 용어 고정 사전: 용어집의 역어를 주입해 페이지별 번역의 용어 일관성을 지킨다.
        # 서재 통합 용어집에서 고정한 대표 역어가 논문별 역어를 덮는다 — 그게 통합의 목적이다.
        gl = read_json(d / "glossary.json", {}) or {}
        canon = _canon_terms()
        pairs_gl = {}
        for i in (gl.get("items") or []):
            t, ko = str(i.get("term") or "").strip(), str(i.get("ko") or "").strip()
            if t and ko:
                pairs_gl[t] = canon.get(t.lower(), ko)
        terms = [f"- {t} → {ko}" for t, ko in pairs_gl.items()][:40]
        term_block = ("===== 용어 대역 (아래 용어는 반드시 이 역어로 통일) =====\n"
                      + "\n".join(terms) + "\n\n") if terms else ""
        prompt = (f"논문 제목: {meta.get('title', '')}\n\n{term_block}"
                  f"===== 이전 페이지 끝부분 (맥락 참고용, 번역하지 말 것) =====\n{prev_tail or '(없음)'}\n\n"
                  f"===== {n}페이지 원문 =====\n{src}\n\n"
                  f"===== 다음 페이지 시작부분 (맥락 참고용, 경계 문장 판단만) =====\n{next_head or '(없음)'}")
        raw = await ask_claude_async(prompt, system, timeout=300)
        out, pairs = split_translation(raw)

    with TRANS_LOCK:
        cached = read_json(d / "translation.json", {"pages": {}})
        cached["pages"][str(n)] = out
        write_json(d / "translation.json", cached)
    with ALIGN_LOCK:
        acached = read_json(d / "align.json", {"pages": {}})
        acached["pages"][str(n)] = pairs
        write_json(d / "align.json", acached)
    return {"n": n, "markdown": out, "pairs": pairs}


# ---------------------------------------------------------------- 백그라운드 작업 (서버측 번역)
# 여태 전체 번역은 클라이언트가 페이지 루프를 돌렸다. 리더를 닫거나 홈으로 나가면 거기서 끝난다 —
# 40쪽짜리를 걸어두고 다른 논문을 읽을 수가 없었다. 여기서는 서버가 돌리고 화면은 진행률만 본다.
# 세마포어(_CLAUDE_SEM, 상한 4)는 그대로 공유한다. 잡 워커를 2로 묶어 대화형 호출이 굶지 않게 한다.

JOBS: dict = {}
JOBS_LOCK = threading.Lock()
JOB_WORKERS = 2
JOB_KEEP = 20                  # 끝난 잡을 목록에 남겨두는 수 (진행률을 놓쳐도 결과는 보이게)


def _pending_pages(pid: str) -> list:
    """아직 번역이 없는 쪽 번호. text.json 이 없으면 빈 목록 — 본문 추출이 먼저다."""
    d = paper_dir(pid)
    t = read_json(d / "text.json", None)
    if not t or not t.get("pages"):
        return []
    have = set(read_json(d / "translation.json", {"pages": {}}).get("pages", {}))
    return [p["n"] for p in t["pages"] if str(p["n"]) not in have]


def _job_view(j: dict) -> dict:
    """클라이언트에 보내는 모양 — 내부 stop 플래그나 태스크 핸들은 빼고."""
    return {k: j[k] for k in ("id", "kind", "label", "total", "done", "failed",
                              "state", "started", "errs")}


async def _run_translate_job(job: dict, units: list):
    """units: [(pid, n), ...]. 워커 JOB_WORKERS개가 같은 목록을 나눠 먹는다."""
    idx = 0
    lock = asyncio.Lock()

    async def worker():
        nonlocal idx
        while not job["stop"]:
            async with lock:
                if idx >= len(units):
                    return
                pid, n = units[idx]
                idx += 1
            try:
                await _translate_one(pid, n)
            except HTTPException as e:
                job["failed"] += 1
                job["errs"][f"{pid}:{n}"] = str(e.detail)[:120]
            except Exception as e:                      # noqa: BLE001 — 잡 하나가 서버를 끌 수 없다
                job["failed"] += 1
                job["errs"][f"{pid}:{n}"] = f"{type(e).__name__}: {e}"[:120]
            job["done"] += 1

    try:
        await asyncio.gather(*[worker() for _ in range(JOB_WORKERS)])
        job["state"] = "stopped" if job["stop"] else "done"
    except asyncio.CancelledError:
        job["state"] = "stopped"
        raise
    finally:
        job["task"] = None


@app.get("/api/jobs")
def list_jobs():
    with JOBS_LOCK:
        return {"items": [_job_view(j) for j in
                          sorted(JOBS.values(), key=lambda x: -x["started"])]}


@app.post("/api/jobs/translate")
async def start_translate_job(req: Request):
    """ids 가 비면 서재 전체. 이미 번역된 쪽은 건너뛴다 — 재실행이 안전해야 마음 놓고 누른다."""
    try:
        body = await req.json()
    except Exception:
        body = {}
    dirs = _pick_dirs(body.get("ids"))
    # 중복 실행 방지 — 진행 중 번역 잡이 물고 있는 논문은 제외한다. 같은 쪽을 두 잡이
    # 동시에 번역하면 비용이 두 배로 나가고 결과는 어차피 한쪽이 덮는다.
    with JOBS_LOCK:
        busy = {p for j in JOBS.values()
                if j["kind"] == "translate" and j["state"] == "running"
                for p in j.get("pids", [])}
    skipped_busy = [d for d in dirs if d.name in busy]
    dirs = [d for d in dirs if d.name not in busy]
    units, papers = [], []
    for d in dirs:
        pend = _pending_pages(d.name)
        if pend:
            papers.append(read_json(d / "meta.json", {}).get("title") or d.name)
            units += [(d.name, n) for n in pend]
    if not units:
        if skipped_busy:
            raise HTTPException(409, "이미 같은 번역 작업이 실행 중입니다 — 백그라운드 작업에서 진행률을 확인하세요")
        raise HTTPException(400, "번역할 쪽이 없습니다 (모두 번역돼 있거나 본문 추출 전입니다)")
    label = papers[0] if len(papers) == 1 else f"{len(papers)}편"
    job = {"id": uuid.uuid4().hex[:8], "kind": "translate", "label": label,
           "total": len(units), "done": 0, "failed": 0, "state": "running",
           "started": int(time.time()), "errs": {}, "stop": False, "task": None,
           "pids": sorted({p for p, _ in units})}
    with JOBS_LOCK:
        for old in sorted(JOBS.values(), key=lambda x: x["started"])[:-JOB_KEEP]:
            if old["state"] != "running":
                JOBS.pop(old["id"], None)
        JOBS[job["id"]] = job
    job["task"] = asyncio.create_task(_run_translate_job(job, units))
    return _job_view(job)


# ---------------------------------------------------------------- 자동 준비 체인
# 논문이 서재에 들어오면 "준비 완료" 상태(서지·요약·핵심4색·용어집·추천질문·출처검증)까지
# 백그라운드에서 만든다. 번역은 제외 — 안 읽을 논문 번역비가 제일 아깝다.
# 멱등의 근거는 prep 기록이 아니라 사이드카 실물이다: 있으면 건너뛰고, 지우면 다시 만든다.

PREP_STEPS = ("text", "meta", "summary", "keypoints", "glossary", "questions", "refs")
PREP_LABEL = {"text": "본문 추출", "meta": "서지", "summary": "요약", "keypoints": "핵심 4색",
              "glossary": "용어집", "questions": "추천 질문", "refs": "출처 검증"}
MAIN_LOOP = None


@app.on_event("startup")
async def _capture_loop():
    global MAIN_LOOP
    MAIN_LOOP = asyncio.get_running_loop()


class _NoBody:
    """Request 대역 — 체인이 엔드포인트 함수를 직접 부를 때 쓴다(옵션 body 없음)."""
    async def json(self):
        return {}


def _prep_status(d: Path) -> dict:
    meta = read_json(d / "meta.json", {})
    return {
        "text": (d / "text.json").exists(),
        "meta": bool(meta.get("authors") or meta.get("doi") or meta.get("meta_ai")),
        "summary": (d / "summary.md").exists(),
        "keypoints": bool(read_json(d / "keypoints.json", {}).get("items")),
        "glossary": read_json(d / "glossary.json", {}).get("items") is not None,
        "questions": read_json(d / "questions.json", {}).get("items") is not None,
        "refs": bool(read_json(d / "refs.json", {}).get("verified_ts")),
    }


def _prep_pending(pid: str) -> list:
    st = _prep_status(paper_dir(pid))
    return [s for s in PREP_STEPS if not st[s]]


def _extract_text_pypdf(pid: str) -> bool:
    """서버측 본문 추출(pypdf, 선택 의존성). 품질이 pdf.js보다 낮아 provisional 로 표시하고
    첫 열람 때 pdf.js 추출본이 덮어쓴다(클라이언트가 text_provisional 을 보고 재추출)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return False
    d = paper_dir(pid)
    with TEXT_LOCK:
        if (d / "text.json").exists():
            return True
    try:
        reader = PdfReader(str(d / "paper.pdf"))
        pages = []
        for i, pg in enumerate(reader.pages, 1):
            try:
                txt = pg.extract_text() or ""
            except Exception:       # noqa: BLE001 — 한 쪽 실패로 전체를 버리지 않는다
                txt = ""
            pages.append({"n": i, "text": txt})
    except Exception:               # noqa: BLE001 — 깨진 PDF 는 뷰어(pdf.js) 경로로
        return False
    if not any(p["text"].strip() for p in pages):
        return False                # 스캔 PDF — vision 폴백은 뷰어에서만 가능하다
    with TEXT_LOCK:
        if not (d / "text.json").exists():
            write_json(d / "text.json", {"pages": pages, "provisional": True})
    with META_LOCK:
        meta = read_json(d / "meta.json", {})
        meta["pages"] = len(pages)
        write_json(d / "meta.json", meta)
    return True


async def _prep_step(pid: str, name: str):
    d = paper_dir(pid)
    if name == "meta":
        await extract_metadata(pid)
    elif name == "summary":
        await make_summary(pid)
    elif name == "keypoints":
        await make_keypoints(pid, _NoBody())
    elif name == "glossary":
        await make_glossary(pid, _NoBody())
    elif name == "questions":
        await make_questions(pid, _NoBody())
    elif name == "refs":
        data = read_json(d / "refs.json", {})
        if not (data.get("items") or []):
            data = await parse_refs(pid)
        if not (data.get("items") or []):
            # 참고문헌이 아예 없는 논문(에세이·일부 리뷰) — 검증할 것도 없다. 스탬프를 안 찍으면
            # 이 단계가 영원히 '남은 단계'로 남아 준비를 돌릴 때마다 AI 파싱을 반복한다.
            write_json(d / "refs.json", {**data, "items": [], "verified_ts": int(time.time())})
            return
        await verify_refs(pid, _NoBody())


async def _run_prep_job(job: dict, papers: list):
    """papers: [(pid, [남은 단계])]. 워커가 논문 단위로 물고 단계는 순차 — 단계 간 의존
    (본문→나머지) 때문에 한 논문을 두 워커가 쪼개면 안 된다."""
    idx = 0
    lock = asyncio.Lock()

    async def worker():
        nonlocal idx
        while not job["stop"]:
            async with lock:
                if idx >= len(papers):
                    return
                pid, steps = papers[idx]
                idx += 1
            if "text" in steps:
                ok = await asyncio.to_thread(_extract_text_pypdf, pid)
                job["done"] += 1
                if not ok:
                    job["failed"] += 1      # 실패는 본문 추출 하나 — 나머지는 건너뜀이지 실패가 아니다
                    job["errs"][f"{pid}:text"] = "본문 추출 불가(pypdf 미설치·스캔 PDF) — 첫 열람 후 자동 재개"
                    job["done"] += len(steps) - 1
                    continue
                steps = [s for s in steps if s != "text"]
            for name in steps:
                if job["stop"]:
                    return
                try:
                    await _prep_step(pid, name)
                except HTTPException as e:
                    job["failed"] += 1
                    job["errs"][f"{pid}:{name}"] = f"{PREP_LABEL[name]}: {str(e.detail)[:100]}"
                except Exception as e:          # noqa: BLE001 — 한 단계가 체인을 끌 수 없다
                    job["failed"] += 1
                    job["errs"][f"{pid}:{name}"] = f"{PREP_LABEL[name]}: {type(e).__name__}: {e}"[:120]
                job["done"] += 1

    try:
        await asyncio.gather(*[worker() for _ in range(JOB_WORKERS)])
        job["state"] = "stopped" if job["stop"] else "done"
    except asyncio.CancelledError:
        job["state"] = "stopped"
        raise
    finally:
        job["task"] = None


def _start_prep_job(pids: list) -> dict | None:
    """준비할 게 있는 논문만 골라 잡을 등록한다. 이미 다른 준비 잡이 물고 있는 논문은 제외.
    이벤트 루프 밖(sync 엔드포인트 = 스레드풀)에서도 안전하게 시작한다."""
    with JOBS_LOCK:
        busy_pids = {p for j in JOBS.values() if j["kind"] == "prep" and j["state"] == "running"
                     for p in j.get("pids", [])}
    papers = []
    for pid in pids:
        if pid in busy_pids or not (LIB / pid / "meta.json").exists():
            continue
        pending = _prep_pending(pid)
        if pending:
            papers.append((pid, pending))
    if not papers:
        return None
    first_title = read_json(LIB / papers[0][0] / "meta.json", {}).get("title") or papers[0][0]
    label = first_title if len(papers) == 1 else f"{len(papers)}편 준비"
    job = {"id": uuid.uuid4().hex[:8], "kind": "prep", "label": label,
           "total": sum(len(s) for _, s in papers), "done": 0, "failed": 0,
           "state": "running", "started": int(time.time()), "errs": {}, "stop": False,
           "task": None, "pids": [p for p, _ in papers]}
    with JOBS_LOCK:
        for old in sorted(JOBS.values(), key=lambda x: x["started"])[:-JOB_KEEP]:
            if old["state"] != "running":
                JOBS.pop(old["id"], None)
        JOBS[job["id"]] = job
    coro = _run_prep_job(job, papers)
    try:
        job["task"] = asyncio.get_running_loop().create_task(coro)
    except RuntimeError:                       # 스레드풀(sync 라우트)에서 호출된 경우
        if MAIN_LOOP is None:
            JOBS.pop(job["id"], None)
            return None
        job["task"] = asyncio.run_coroutine_threadsafe(coro, MAIN_LOOP)
    return _job_view(job)


def _auto_prep(pids: list):
    """유입 트리거용 — 설정이 꺼져 있으면 조용히 아무것도 안 한다."""
    try:
        if load_config().get("auto_prep") and pids:
            _start_prep_job(pids)
    except Exception:                          # noqa: BLE001 — 준비 실패가 유입 자체를 막으면 안 된다
        pass


@app.post("/api/prep")
async def start_prep(req: Request):
    """수동 준비(소급). ids 가 비면 서재 전체. 자동 토글과 무관하게 항상 동작한다."""
    try:
        body = await req.json()
    except Exception:
        body = {}
    dirs = _pick_dirs(body.get("ids"))
    j = _start_prep_job([d.name for d in dirs])
    if not j:
        return {"started": False, "reason": "준비할 단계가 남은 논문이 없습니다"}
    return {"started": True, **j}


@app.delete("/api/jobs/{jid}")
def stop_job(jid: str):
    """실행 중이면 중지, 끝난(완료·중지) 잡이면 목록에서 제거 — 버튼 하나가 상태 따라 다르게 동작한다."""
    with JOBS_LOCK:
        job = JOBS.get(jid)
        if not job:
            raise HTTPException(404, "그런 작업이 없습니다")
        if job["state"] != "running":
            JOBS.pop(jid, None)
            return {"removed": True}
    job["stop"] = True              # 진행 중인 페이지는 마저 끝낸다 — 중간에 끊으면 그 쪽이 날아간다
    return _job_view(job)


# ---------------------------------------------------------------- sentence alignment

ALIGN_LOCK = threading.Lock()

# 앵커 방식: 문장 전체를 재출력하면 출력이 비대해 haiku 생성이 느려진다(배치 무의미).
# 각 문장의 '앞부분 앵커'만 뽑으면 출력이 1/4~1/5로 줄어 빨라지고 배치도 유효해진다.
# 프론트는 s/t 앵커로 원문·번역에서 문장 시작점을 찾고 다음 앵커 직전까지를 문장 범위로 삼는다.
_ANCHOR_RULES = (
    "각 문장쌍은 문장 전체가 아니라 시작·끝 '앵커'만 출력한다 — 원문 문장의 처음 5~8단어(s)와 "
    "마지막 4~6단어(se), 번역 문장의 처음 5~8단어(t)와 마지막 4~6단어(te). 규칙: 네 앵커 모두 "
    "원문/번역문에 나온 그대로(verbatim, 한 글자도 바꾸지 말 것) · se/te는 그 문장의 진짜 마지막 "
    "부분(문장부호 직전까지, 다음 문장·제목·각주는 절대 포함하지 말 것) · 각 앵커는 그 페이지에서 "
    "유일하게 찾을 만큼 길게 · 원문 등장 순서대로 · 마크다운 기호(#,*,|,`) 제외 · "
    "제목·표·수식·참고문헌·각주만 있는 조각은 건너뜀."
)
ALIGN_SYS_SINGLE = (
    "너는 원문(영어 논문)과 그 한국어 번역을 문장 단위로 정렬하는 도구다. " + _ANCHOR_RULES +
    ' 출력은 JSON 배열 하나만: [{"s":"원문 시작","se":"원문 끝","t":"번역 시작","te":"번역 끝"}]. 다른 텍스트 금지.'
)


@app.get("/api/papers/{pid}/align")
def get_align(pid: str):
    return read_json(paper_dir(pid) / "align.json", {"pages": {}})


@app.post("/api/papers/{pid}/align/page")
async def align_page(pid: str, req: Request):
    """원문 페이지와 그 한국어 번역을 문장 단위로 정렬한다(haiku — 단순 매칭이라 빠름).
    출력 [{src, tgt}]는 원문 강조↔번역 하이라이트 양방향에 공유된다. 페이지 단위 캐시."""
    d = paper_dir(pid)
    body = await req.json()
    n = int(body.get("n") or 0)
    t = read_json(d / "text.json", None)
    if not t or not t.get("pages"):
        raise HTTPException(409, _no_text(d))
    pages = t["pages"]
    if n < 1 or n > len(pages):
        raise HTTPException(400, "잘못된 페이지 번호")
    trans = read_json(d / "translation.json", {"pages": {}})
    tgt = trans["pages"].get(str(n))
    if tgt is None:
        raise HTTPException(409, "이 페이지를 먼저 번역하세요")
    cached = read_json(d / "align.json", {"pages": {}})
    if str(n) in cached["pages"] and not body.get("force"):
        return {"n": n, "pairs": cached["pages"][str(n)], "cached": True}

    src = pages[n - 1]["text"]
    if not src.strip() or not tgt.strip():
        pairs = []
    else:
        prompt = f"===== 원문 =====\n{src[:20000]}\n\n===== 한국어 번역 =====\n{tgt[:20000]}"
        # 수식·표 밀집 페이지는 haiku 정렬이 180초를 상습 초과(2026-07-24 실측: 15쪽 중 3쪽 504) → 기본 300초로 완화
        raw = await ask_claude_async(prompt, ALIGN_SYS_SINGLE, timeout=300, model="haiku")
        # 표·수식·참고문헌만인 쪽은 정렬할 문장이 없어 빈 배열이 정상이다 → empty_ok
        parsed = parse_json_or_502(raw, list, "정렬 응답", empty_ok=True)
        pairs = [_anchor(p) for p in parsed if isinstance(p, dict) and p.get("s") and p.get("t")]

    with ALIGN_LOCK:
        cached = read_json(d / "align.json", {"pages": {}})
        cached["pages"][str(n)] = pairs
        write_json(d / "align.json", cached)
    return {"n": n, "pairs": pairs}


# ---------------------------------------------------------------- references

@app.get("/api/papers/{pid}/refs")
def get_refs(pid: str):
    return read_json(paper_dir(pid) / "refs.json", {"items": []})


@app.post("/api/papers/{pid}/refs")
async def parse_refs(pid: str):
    """참고문헌 섹션을 AI로 구조화한다 (클라이언트 정규식 파싱이 실패한 비정형 대응)."""
    text = full_text(pid)
    tail = text[-40_000:]
    system = (
        "논문 텍스트 끝부분에서 참고문헌 목록을 추출하는 도구다. "
        '출력은 JSON 하나만: {"style": "numeric|authoryear", '
        '"items": [{"n": "1", "text": "서지 전체"}]}. '
        "numeric이면 n은 번호, authoryear면 n은 'Smith2019' 꼴 키. 다른 텍스트 금지."
    )
    raw = await ask_claude_async(f"===== 논문 끝부분 =====\n{tail}", system)
    data = parse_json_or_502(raw, dict)
    if not isinstance(data.get("items"), list):
        raise HTTPException(502, f"참고문헌 목록(items 배열)이 없습니다. 받은 내용 앞부분: {raw[:200]!r}")
    write_json(paper_dir(pid) / "refs.json", data)
    return data


# ---------------------------------------------------------------- citations (BibTeX/APA/ACS)
# 저장된 메타데이터(authors/year/venue/doi/title)만으로 문자열을 조립한다 — AI 호출 없음, 순수함수라
# --self-check으로 회귀를 잡는다. 누락 필드는 조용히 생략(None·빈 {} 금지)가 전체 규칙.

_KEY_STOPWORDS = {"a", "an", "the", "of", "on", "in", "for", "and", "to", "with", "using", "toward", "via"}


def _cite_year(meta: dict) -> str:
    """연도를 4자리 문자열로. 없거나 범위를 벗어나면 빈 문자열 — extract_metadata와 같은 1900~2100 검증."""
    try:
        y = int(meta.get("year"))
    except (TypeError, ValueError):
        return ""
    return str(y) if 1900 <= y <= 2100 else ""


def _bibtex_key(meta: dict) -> str:
    """BibTeX 인용 키: Zotero(Better BibTeX) citekey 가 있으면 그것이 정본이다 — 사용자의
    기존 .bib·원고와 키가 일치해야 한다. 없으면 첫 저자 성 + 연도 + 제목 첫 의미어(불용어
    제외)를 소문자 ASCII로 정제해 조합하고, 아무것도 없으면 논문 id — 항상 결정론적."""
    ck = re.sub(r"[^A-Za-z0-9_.:\-]", "", str(meta.get("citekey") or "").strip())
    if ck:
        return ck
    authors = meta.get("authors") or []
    surname = ""
    if authors:
        first = str(authors[0]).strip()
        last = first.split(",")[0].strip() if "," in first else (first.split()[-1] if first.split() else "")
        surname = re.sub(r"[^A-Za-z]", "", last).lower()
    year = _cite_year(meta)
    word = ""
    for w in re.findall(r"[A-Za-z]+", str(meta.get("title") or "")):
        if len(w) > 2 and w.lower() not in _KEY_STOPWORDS:
            word = w.lower()
            break
    key = f"{surname}{year}{word}"
    return key or re.sub(r"[^A-Za-z0-9]", "", str(meta.get("id") or "")) or "paper"


def _split_author(name: str):
    """"Firstname Lastname" 또는 "Lastname, Firstname" → (성, 이니셜들)."""
    name = str(name).strip()
    if not name:
        return "", ""
    if "," in name:
        last, given = (p.strip() for p in name.split(",", 1))
    else:
        parts = name.split()
        last, given = (parts[-1], " ".join(parts[:-1])) if len(parts) > 1 else (parts[0], "")
    initials = "".join(f"{w[0].upper()}." for w in re.findall(r"[A-Za-z']+", given))
    return last, initials


def _authors_apa(authors: list) -> str:
    names = [f"{last}, {ini}" if ini else last for last, ini in (_split_author(a) for a in authors) if last]
    if not names:
        return ""
    if len(names) == 1:
        return names[0]
    return ", ".join(names[:-1]) + f", & {names[-1]}"


def _authors_acs(authors: list) -> str:
    names = [f"{last}, {ini}" if ini else last for last, ini in (_split_author(a) for a in authors) if last]
    return "; ".join(names)


def _cite_strings(meta: dict) -> dict:
    """저장된 메타데이터로 BibTeX/ACS/APA 문자열을 만드는 순수 함수(AI 호출 없음).
    누락 필드는 그 필드만 조용히 빠진다 — None이나 빈 {}가 출력에 섞이지 않는다."""
    authors = [str(a).strip() for a in (meta.get("authors") or []) if str(a).strip()]
    year = _cite_year(meta)
    venue = str(meta.get("venue") or "").strip()
    doi = str(meta.get("doi") or "").strip()
    title = str(meta.get("title") or "").strip()
    key = _bibtex_key(meta)

    lines = []
    if authors:
        lines.append("  author = {" + " and ".join(authors) + "}")
    if title:
        lines.append("  title = {" + title + "}")
    if venue:
        lines.append("  journal = {" + venue + "}")
    if year:
        lines.append("  year = {" + year + "}")
    if doi:
        lines.append("  doi = {" + doi + "}")
    body = ",\n".join(lines)
    bibtex = "@article{" + key + ",\n" + body + "\n}" if body else "@article{" + key + "}"

    apa = " ".join(b for b in (
        _authors_apa(authors), f"({year})." if year else "", f"{title}." if title else "",
        f"{venue}." if venue else "", f"https://doi.org/{doi}" if doi else "") if b)

    acs_tail = " ".join(x for x in (venue, year) if x)
    acs = " ".join(b for b in (
        _authors_acs(authors), f"{title}." if title else "",
        f"{acs_tail}." if acs_tail else "", f"https://doi.org/{doi}" if doi else "") if b)

    return {"bibtex": bibtex, "acs": acs, "apa": apa}


@app.get("/api/papers/{pid}/cite")
def get_cite(pid: str):
    meta = read_json(paper_dir(pid) / "meta.json", {})
    return _cite_strings(meta)


# ---------------------------------------------------------------- 심화 검색 (OpenAlex)
#
# README가 "외부 검색 API 필요"라며 미지원으로 뒀던 기능. OpenAlex는 키·가입이 필요 없고
# (mailto를 붙이면 polite pool로 더 안정적) 라이선스가 CC0라 그대로 쓸 수 있다.
# HTTP는 stdlib urllib — requests/httpx를 끌어오면 run.bat 의존성이 늘어난다.
# 라우트를 async가 아닌 sync def로 두는 이유: FastAPI가 sync 라우트를 스레드풀에서 돌리므로
# 블로킹 urlopen이 이벤트 루프(= 번역·요약 동시 진행)를 멈추지 않는다.

OPENALEX = "https://api.openalex.org"
OA_UA = {"User-Agent": "achird-local (local research reader; mailto:achird@localhost)"}
OA_FIELDS = "id,doi,display_name,publication_year,cited_by_count,primary_location,authorships"
OA_TIMEOUT = 20


def _http_json(url: str, tries: int = 3) -> dict:
    """외부 JSON GET + 429/5xx 백오프 재시도.

    참고문헌 40건을 연달아 물으면 OpenAlex·Crossref 둘 다 간헐적으로 429를 낸다. 한 번 튕겼다고
    '미확인'으로 적으면 같은 목록을 다시 돌릴 때마다 결과가 달라진다 — 검증이 검증이 아니게 된다.
    404는 즉시 올려보낸다(없는 것과 못 물어본 것은 다르다)."""
    delay = 0.8
    for i in range(tries):
        try:
            req = urllib.request.Request(url, headers=OA_UA)
            with urllib.request.urlopen(req, timeout=OA_TIMEOUT) as r:
                return json.loads(r.read().decode("utf-8"))
        except urllib.error.HTTPError as e:
            if e.code not in (429, 500, 502, 503, 504) or i == tries - 1:
                raise
        except urllib.error.URLError:
            if i == tries - 1:
                raise
        time.sleep(delay)
        delay *= 2
    raise urllib.error.URLError("retries exhausted")      # 도달 불가 — 위 루프가 항상 반환/예외


def _oa_get(path: str) -> dict:
    return _http_json(f"{OPENALEX}/{path}")


def _oa_slim(w: dict, lib_index: dict) -> dict:
    """OpenAlex work → UI가 쓰는 최소 필드. lib_index로 '이미 서재에 있음'을 붙인다."""
    loc = w.get("primary_location") or {}
    src = loc.get("source") or {}
    doi = (w.get("doi") or "").replace("https://doi.org/", "").lower()
    title = w.get("display_name") or ""
    return {
        "oa_id": (w.get("id") or "").rsplit("/", 1)[-1],
        "title": title,
        "year": w.get("publication_year"),
        "doi": doi,
        "venue": src.get("display_name") or "",
        "authors": [a.get("author", {}).get("display_name", "")
                    for a in (w.get("authorships") or [])[:3]],
        "cited_by_count": w.get("cited_by_count") or 0,
        "url": w.get("id") or "",
        "in_library": _lib_hit(doi, title, lib_index),
    }


def _oa_clean(works: list, index: list) -> list:
    """제목 없는 레코드는 버린다 — OpenAlex에는 display_name이 빈 항목이 섞여 있고
    (실측: 인용수 5만짜리 레코드도 제목이 비어 있었다) 목록에 뜨면 클릭할 수 없는 빈 줄이 된다."""
    out = []
    for w in works:
        s = _oa_slim(w, index)
        if s["title"].strip():
            out.append(s)
    return out


def _lib_index() -> list:
    """서재 논문의 (pid, doi, 제목 시그니처). 관련 논문이 이미 서재에 있는지 표시하는 데 쓴다."""
    out = []
    for d in sorted(LIB.iterdir()):
        if not d.is_dir():
            continue
        m = read_json(d / "meta.json", None)
        if not m:
            continue
        sig = _alnum_sig(m.get("title"))
        out.append({"pid": m.get("id", d.name),
                    "doi": str(m.get("doi") or "").lower().strip(),
                    "sig": sig if len(sig) >= 15 else ""})
    return out


def _lib_hit(doi: str, title: str, index: list):
    """doi 완전일치 우선, 없으면 정규화 제목 일치. _graph_edges와 같은 매칭 규칙."""
    if doi:
        hit = next((c for c in index if c["doi"] and c["doi"] == doi), None)
        if hit:
            return hit["pid"]
    sig = _alnum_sig(title)
    if len(sig) < 15:
        return None
    hit = next((c for c in index if c["sig"] and (c["sig"] == sig or c["sig"] in sig or sig in c["sig"])), None)
    return hit["pid"] if hit else None


def _oa_resolve(meta: dict) -> dict:
    """서재 논문 → OpenAlex work. doi가 있으면 그것으로, 없으면 제목 검색 후
    정규화 제목이 실제로 겹치는지 확인한다(검색은 항상 무언가를 돌려주므로 확인이 필수)."""
    doi = str(meta.get("doi") or "").strip().lower().replace("https://doi.org/", "")
    if doi:
        try:
            return _oa_get(f"works/https://doi.org/{doi}")
        except urllib.error.HTTPError as e:
            if e.code != 404:
                raise
    title = str(meta.get("title") or "").strip()
    if not title:
        raise HTTPException(404, "DOI도 제목도 없어 외부 검색을 할 수 없습니다. 인용 탭에서 서지정보를 먼저 추출하세요.")
    q = urllib.parse.quote(title[:250])
    res = _oa_get(f"works?search={q}&per-page=5&select={OA_FIELDS}")
    want = _alnum_sig(title)
    for w in res.get("results") or []:
        got = _alnum_sig(w.get("display_name"))
        if got and want and (got in want or want in got):
            return _oa_get(f"works/{(w.get('id') or '').rsplit('/', 1)[-1]}")
    raise HTTPException(404, "OpenAlex에서 이 논문을 찾지 못했습니다. 제목이 정확한지 확인해주세요.")


# ---- 참고문헌 검증: 참고문헌 한 줄 한 줄을 OpenAlex에 물어 DOI·서지를 붙인다.
# 지금까지 refs.json에 있던 건 AI가 뽑은 '문자열'뿐이라, 서지 팝업이 링크가 아니었고
# 인용 그래프도 제목 부분일치에 기대야 했다. 못 찾은 항목은 비워 둔다 — 지어내는 것보다 낫다.

_DOI_RE = re.compile(r"\b10\.\d{4,9}/[^\s\"'<>]+")
REF_VERIFY_MAX = 80        # 한 번에 조회할 참고문헌 상한 (그 뒤 항목은 손대지 않고 남긴다)
REF_VERIFY_SLEEP = 0.12    # OpenAlex polite pool 배려 (~8 req/s)


def _ref_doi(text) -> str:
    """참고문헌 문자열에서 DOI를 뽑는다. 문장 끝 마침표·닫는 괄호는 DOI의 일부가 아니라
    서지 문장의 일부라 떼어낸다 — 안 떼면 조회가 통째로 404가 난다."""
    m = _DOI_RE.search(str(text or ""))
    if not m:
        return ""
    return m.group(0).rstrip(".,;:)]}>”'\"").lower()


_ARXIV_RE = re.compile(r"arxiv[:\s]*\s*(\d{4}\.\d{4,5})(?:v\d+)?", re.I)


def _arxiv_doi(text) -> str:
    """서지 문장에 arXiv 번호만 적힌 경우의 DOI. arXiv는 모든 프리프린트에 10.48550/arXiv.<id>
    DOI를 소급 부여했고 OpenAlex가 그대로 색인한다(실측: Adam·Bahdanau·Layer Norm 모두 해석됨).
    Crossref에 등재되지 않는 ICLR·NIPS·arXiv 인용이 참고문헌의 큰 몫이라 이 폴백이 없으면
    ML 계열 논문에서 절반 넘게 '미확인'으로 떨어진다."""
    m = _ARXIV_RE.search(str(text or ""))
    return f"10.48550/arxiv.{m.group(1)}" if m else ""


def _ref_title_ok(ref_text, work_title) -> bool:
    """OpenAlex 검색은 무엇을 넣든 무언가를 돌려준다 — 돌려준 제목이 참고문헌 문자열 안에
    실제로 들어 있을 때만 같은 논문으로 본다(_lib_hit·_graph_edges와 같은 정규화 포함 규칙).
    15자 미만 제목은 우연히 포함될 수 있어 버린다."""
    got, want = _alnum_sig(work_title), _alnum_sig(ref_text)
    return bool(got) and len(got) >= 15 and got in want


CROSSREF = "https://api.crossref.org"


def _cr_lookup(text: str) -> dict:
    """서지 문장 → {doi,title,year,venue} (또는 None).

    Crossref의 query.bibliographic은 인용 문자열을 해석하라고 만든 것이라 이 일에 OpenAlex의
    전문 검색보다 훨씬 정확하다 — 실측: Yoshida 2016(Science)을 OpenAlex search는 상위 10위
    안에 못 올렸고 Crossref는 1위로 올렸다. 돌려준 제목이 서지 문장 안에 실제로 있을 때만
    채택한다(검색은 언제나 무언가를 돌려주므로 이 확인이 유일한 방어선이다)."""
    q = urllib.parse.quote(re.sub(r"\s+", " ", str(text)).strip()[:400])
    if not q:
        return None
    url = f"{CROSSREF}/works?query.bibliographic={q}&rows=3&select=DOI,title,issued,container-title"
    items = (_http_json(url).get("message") or {}).get("items") or []
    for it in items:
        title = ((it.get("title") or [""]) or [""])[0]
        if not _ref_title_ok(text, title):
            continue
        parts = ((it.get("issued") or {}).get("date-parts") or [[None]])[0]
        return {"doi": str(it.get("DOI") or "").lower(), "title": title,
                "year": parts[0] if parts and isinstance(parts[0], int) else None,
                "venue": ((it.get("container-title") or [""]) or [""])[0]}
    return None


def _verify_one(text: str, index: list) -> dict:
    """참고문헌 한 건 → 검증된 서지(또는 None).
    ① 서지 문장에 DOI가 박혀 있으면 그것으로 ② 없으면 Crossref가 인용 문자열을 DOI로 풀고
    ③ 그 DOI로 OpenAlex work를 받아 피인용 수·서재 매칭까지 붙인다.

    DOI가 정해진 뒤의 OpenAlex 조회는 '덤'이다. 없든(404) 막히든(429) 이미 확정된 DOI를
    버리지 않는다 — 실측: 이 폴백이 없을 때 40건짜리 목록에서 39건이 '미확인'으로 떨어졌는데,
    같은 참고문헌을 하나씩 물으면 대부분 Crossref가 정확히 찾아냈다."""
    doi, cr = _ref_doi(text), None
    if not doi:
        # Crossref를 arXiv보다 먼저 본다 — 같은 논문이면 학술지 게재본 쪽이 인용하기 좋다
        cr = _cr_lookup(text)
        doi = (cr["doi"] if cr else "") or _arxiv_doi(text)
    if not doi:
        return None
    try:
        return _oa_slim(_oa_get(f"works/https://doi.org/{urllib.parse.quote(doi)}?select={OA_FIELDS}"), index)
    except (urllib.error.URLError, json.JSONDecodeError, KeyError, IndexError):
        pass
    title = cr["title"] if cr else ""
    return {"oa_id": "", "title": title, "year": cr["year"] if cr else None, "doi": doi,
            "venue": cr["venue"] if cr else "", "authors": [], "cited_by_count": 0,
            "url": f"https://doi.org/{doi}", "in_library": _lib_hit(doi, title, index)}


def _verify_refs(items: list) -> list:
    """참고문헌 목록에 oa 필드를 붙여 돌려준다(원본 필드는 보존).
    한 건이 네트워크로 실패해도 나머지를 계속한다 — 50건짜리 목록이 1건 때문에 통째로
    죽으면 다시 돌릴 때 성공했던 49건도 같이 다시 조회된다."""
    index = _lib_index()
    out = []
    for i, it in enumerate(items):
        rec = dict(it)
        if i >= REF_VERIFY_MAX:
            out.append(rec)
            continue
        try:
            rec["oa"] = _verify_one(str(rec.get("text") or ""), index)
        except (urllib.error.URLError, json.JSONDecodeError, KeyError, ValueError, IndexError) as e:
            rec["oa"] = None    # 없는 것과 못 물어본 것은 다르다 — 이유를 남겨 재시도할 수 있게
            rec["oa_error"] = f"{type(e).__name__}: {e}"[:120]
        out.append(rec)
        time.sleep(REF_VERIFY_SLEEP)
    return out


@app.post("/api/papers/{pid}/refs/verify")
async def verify_refs(pid: str, req: Request):
    """참고문헌을 OpenAlex와 대조해 DOI·연도·저널을 붙인다. 못 찾은 건 oa:null로 남아
    화면에 '미확인'으로 뜬다. 결과는 refs.json에 그대로 저장돼 인용 그래프도 같이 정확해진다."""
    d = paper_dir(pid)
    data = read_json(d / "refs.json", None) or {}
    items = data.get("items") or []
    if not items:
        # 클라이언트 정규식 파싱본(parseRefsClient)은 서버에 저장되지 않는다 — 같이 받아 저장한다.
        # 안 그러면 AI 재분석을 안 돌린 논문에서 "먼저 정리하세요"만 나온다(이미 목록은 보이는데도).
        try:
            body = await req.json()
        except Exception:
            body = {}
        sent = [{"n": str(i.get("n") or ""), "text": str(i.get("text") or "")}
                for i in (body.get("items") or [])[:400]
                if isinstance(i, dict) and str(i.get("text") or "").strip()]
        if sent:
            data = {"style": str(body.get("style") or "numeric"), "items": sent}
            items = sent
    if not items:
        raise HTTPException(400, "먼저 참고문헌을 정리해주세요 — 인용 탭의 'AI로 재분석'")
    items = await asyncio.to_thread(_verify_refs, items)
    data = {**data, "items": items, "verified_ts": int(time.time())}
    write_json(d / "refs.json", data)
    return {**data, "verified": sum(1 for i in items if i.get("oa")), "total": len(items),
            "errors": sum(1 for i in items if i.get("oa_error")),   # 못 찾은 건과 조회 실패를 화면이 구분하도록
            "checked": min(len(items), REF_VERIFY_MAX)}


# ---- 주제어 검색: 심화 검색이 '이 논문과 이어진 것'만 보여주는 데 비해 여기는 처음부터 주제로 찾는다.
# 결과는 캐시하지 않는다 — 질의가 매번 다르고, 캐시 키가 곧 질의 문자열이라 값어치가 없다.

@app.get("/api/search/works")
def search_works(q: str = "", year_from: int = 0, page: int = 1):
    q = str(q or "").strip()
    if len(q) < 2:
        raise HTTPException(400, "검색어를 2자 이상 입력해주세요")
    path = (f"works?search={urllib.parse.quote(q[:250])}&per-page=20"
            f"&page={max(1, min(20, int(page or 1)))}&select={OA_FIELDS}")
    if 1900 <= int(year_from or 0) <= 2100:
        path += f"&filter=from_publication_date:{int(year_from)}-01-01"
    try:
        r = _oa_get(path)
    except urllib.error.HTTPError as e:      # URLError 서브클래스 — 먼저 걸러야 상태코드가 안 사라진다
        raise HTTPException(502, f"OpenAlex 오류 HTTP {e.code} {e.reason}")
    except urllib.error.URLError as e:
        raise HTTPException(502, f"OpenAlex에 연결하지 못했습니다 — {e.reason}")
    except (json.JSONDecodeError, KeyError) as e:
        raise HTTPException(502, f"OpenAlex 응답을 해석하지 못했습니다 — {e}")
    return {"q": q, "page": page, "total": (r.get("meta") or {}).get("count", 0),
            "results": _oa_clean(r.get("results") or [], _lib_index())}


@app.get("/api/papers/{pid}/related")
def related_works(pid: str, refresh: int = 0):
    """심화 검색: 이 논문을 인용한 논문 + OpenAlex가 고른 관련 논문.
    네트워크·외부 rate limit이 있으므로 related.json에 캐시하고 refresh=1로만 다시 받는다."""
    d = paper_dir(pid)
    cache = d / "related.json"
    if not refresh:
        c = read_json(cache, None)
        if c:
            return {**c, "cached": True}
    meta = read_json(d / "meta.json", {})
    try:
        work = _oa_resolve(meta)
        wid = (work.get("id") or "").rsplit("/", 1)[-1]
        index = _lib_index()
        cited = _oa_get(f"works?filter=cites:{wid}&sort=cited_by_count:desc&per-page=12&select={OA_FIELDS}")
        rel_ids = [r.rsplit("/", 1)[-1] for r in (work.get("related_works") or [])[:12]]
        related = []
        if rel_ids:
            r = _oa_get(f"works?filter=openalex_id:{'|'.join(rel_ids)}&per-page=12&select={OA_FIELDS}")
            related = _oa_clean(r.get("results") or [], index)
    except HTTPException:
        raise
    except urllib.error.HTTPError as e:      # URLError 서브클래스 — 먼저 걸러야 상태코드가 안 사라진다
        raise HTTPException(502, f"OpenAlex 오류 HTTP {e.code} {e.reason}")
    except urllib.error.URLError as e:
        raise HTTPException(502, f"OpenAlex에 연결하지 못했습니다 — {e.reason}")
    except (json.JSONDecodeError, KeyError) as e:
        raise HTTPException(502, f"OpenAlex 응답을 해석하지 못했습니다 — {e}")
    data = {
        "work": {"oa_id": wid, "title": work.get("display_name") or "",
                 "year": work.get("publication_year"),
                 "cited_by_count": work.get("cited_by_count") or 0,
                 "url": work.get("id") or ""},
        "cited_by": _oa_clean(cited.get("results") or [], index),
        "cited_by_total": (cited.get("meta") or {}).get("count", 0),
        "related": related,
        "ts": int(time.time()),
    }
    write_json(cache, data)
    return data


def _pick_dirs(raw_ids) -> list:
    """ids 가 비었으면 서재 전체(최근순), 있으면 그 논문들만.
    보낸 id 가 하나도 안 맞으면 조용히 전체를 돌려주지 않고 400 — '이 1편 달라'는 요청에
    서재 40편을 주는 건 틀린 답이다(실측: citations 는 전체를, bib 는 0건을 돌려주고 있었다)."""
    given = [i for i in (raw_ids or []) if isinstance(i, str)]
    if not given:
        return sorted((d for d in LIB.iterdir() if d.is_dir()),
                      key=lambda p: p.stat().st_mtime, reverse=True)
    dirs = [LIB / i for i in given if ID_RE.match(i) and (LIB / i).is_dir()]
    if not dirs:
        raise HTTPException(400, "요청한 논문을 서재에서 찾지 못했습니다")
    return dirs


@app.post("/api/export/bib")
async def export_bib(req: Request):
    """선택 논문(ids 미지정/빈 배열이면 서재 전체)의 BibTeX을 이어붙인 평문 .bib로 내보낸다.
    id가 잘못됐거나 메타데이터가 아예 없는 논문은 조용히 건너뛴다 — 벌크 작업이라
    한 편 때문에 전체가 실패하면 안 된다(선택 2~3편이 전제인 compare와는 다른 성격)."""
    try:
        body = await req.json()
    except Exception:
        body = {}
    dirs = _pick_dirs(body.get("ids"))
    entries = []
    for d in dirs:
        meta = read_json(d / "meta.json", {})
        if not (meta.get("title") or meta.get("authors")):
            continue
        entries.append(_cite_strings(meta)["bibtex"])
    text = ("\n\n".join(entries) + "\n") if entries else ""
    return PlainTextResponse(text, media_type="text/plain; charset=utf-8")


CITE_STYLES = ("acs", "apa", "bibtex")


def _numbered_refs(metas: list, style: str) -> list:
    """선택 논문을 1..N 번호로 매긴다(순수 함수). 본문에는 [n]을 쓰고 목록은 이 순서로 나열한다."""
    out = []
    for i, m in enumerate(metas, 1):
        c = _cite_strings(m)
        out.append({"n": i, "pid": m.get("id", ""), "title": str(m.get("title") or "").strip(),
                    "short": _cite_short(m), "cite": c[style]})
    return out


@app.post("/api/export/citations")
async def export_citations(req: Request):
    """번호식 참고문헌 목록. 번호 순서는 넘겨받은 ids 순서 그대로다 — 본문에서 처음 인용한
    순서가 무엇인지는 서버가 알 수 없고, 그 순서를 정하는 건 글쓴이 몫이다.
    ids가 비면 서재 전체를 최근순으로 매긴다."""
    try:
        body = await req.json()
    except Exception:
        body = {}
    style = str(body.get("style") or "acs").lower()
    if style not in CITE_STYLES:
        raise HTTPException(400, f"style은 {'/'.join(CITE_STYLES)} 중 하나여야 합니다")
    dirs = _pick_dirs(body.get("ids"))
    metas, skipped = [], 0
    for d in dirs:
        meta = read_json(d / "meta.json", {})
        if not (meta.get("title") or meta.get("authors")):
            skipped += 1                 # 서지정보가 아예 없는 논문은 번호를 낭비하지 않는다
            continue
        metas.append({**meta, "id": meta.get("id", d.name)})
    items = _numbered_refs(metas, style)
    sep = "\n\n" if style == "bibtex" else "\n"
    text = sep.join((it["cite"] if style == "bibtex" else f"({it['n']}) {it['cite']}")
                    for it in items)
    return {"style": style, "items": items, "skipped": skipped,
            "text": text + ("\n" if text else "")}


# ---------------------------------------------------------------- .docx 내보내기
# Markdown 내보내기는 이미 있다. 학과 제출본은 대개 Word 라 한 단계가 더 필요했다.
# python-docx 는 선택 의존성이다 — 없으면 이 버튼만 안 되고 앱은 그대로 돈다.

def _docx_paragraphs(doc, text: str):
    """빈 줄로 나뉜 덩어리를 문단으로. Markdown 서식은 풀지 않는다 —
    Achird 는 산문을 소유하지 않으므로 여기서 굵게·기울임을 흉내 낼 이유가 없다."""
    for chunk in re.split(r"\n\s*\n", str(text or "").strip()):
        chunk = chunk.strip()
        if chunk:
            doc.add_paragraph(chunk)


def _build_draft_docx(draft: dict, ev_index: dict, metas: list, style: str) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(draft.get("title") or "제목 없는 원고", level=0)
    for s in draft.get("sections") or []:
        doc.add_heading(s.get("name") or "이름 없는 절", level=1)
        _docx_paragraphs(doc, s.get("text"))

    picked = [(s, [ev_index[k] for k in (s.get("ev") or []) if k in ev_index])
              for s in draft.get("sections") or []]
    if any(items for _s, items in picked):
        doc.add_page_break()
        doc.add_heading("담은 근거", level=1)
        for s, items in picked:
            if not items:
                continue
            doc.add_heading(s.get("name") or "", level=2)
            for it in items:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{it['cite']}{f', p.{it['page']}' if it.get('page') else ''} — ").bold = True
                p.add_run(" ".join(str(it.get("text") or "").split()))
                if it.get("memo"):
                    doc.add_paragraph(" ".join(str(it["memo"]).split()), style="List Bullet 2")
    if metas:
        doc.add_page_break()
        doc.add_heading("참고문헌", level=1)
        # _numbered_refs 는 dict 를 돌려준다({n,pid,title,short,cite}) — 그대로 넘기면
        # 문단에 키 이름이 찍힌다. 화면에 나갈 문자열은 cite 하나뿐이다.
        for r in _numbered_refs(metas, style):
            doc.add_paragraph(f"[{r['n']}] {r['cite']}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/api/export/docx")
async def export_docx(req: Request):
    """저장된 초안을 그대로 .docx 로. 본문 + 절별 담은 근거 + 인용한 논문의 번호 참고문헌."""
    try:
        body = await req.json()
    except Exception:
        body = {}
    style = body.get("style") if body.get("style") in ("acs", "apa") else "acs"
    draft = read_json(DRAFT, None)
    if not draft or not (draft.get("sections") or []):
        raise HTTPException(400, "내보낼 초안이 없습니다")
    ev_index = {i["key"]: i for i in library_evidence()["items"]}
    used = {k.split(":", 1)[0] for s in draft["sections"] for k in (s.get("ev") or [])}
    metas = [m for m in (read_json(LIB / pid / "meta.json", None) for pid in sorted(used)) if m]
    try:
        blob = _build_draft_docx(draft, ev_index, metas, style)
    except ImportError:
        raise HTTPException(501, "python-docx 가 설치돼 있지 않습니다. run.bat 을 다시 실행하면 설치됩니다.")
    name = re.sub(r'[\\/:*?"<>|]', "", draft.get("title") or "achird-draft")[:60] or "achird-draft"
    return Response(blob, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition":
                             f"attachment; filename*=UTF-8''{urllib.parse.quote(name)}.docx"})


# ---------------------------------------------------------------- 연구 트렌드 (OpenAlex topics)
# 카라멜랩은 관심 분야를 사람이 고르게 한다. Achird 는 고를 필요가 없다 — 서재가 곧 관심사다.
# 각 논문의 OpenAlex topic 을 모아 상위 몇 개를 뽑고, 그 주제의 최근 논문을 피인용 순으로 본다.

TRENDS = LIB / "_trends.json"
TREND_TOPICS = 3


def _paper_topics(d: Path) -> list:
    """논문 하나의 OpenAlex topic. 편당 한 번만 물어보고 topics.json 에 캐시한다."""
    cache = read_json(d / "topics.json", None)
    if cache is not None:
        return cache.get("items") or []
    meta = read_json(d / "meta.json", {})
    try:
        w = _oa_resolve(meta)
        items = [{"id": (t.get("id") or "").rsplit("/", 1)[-1], "name": t.get("display_name") or ""}
                 for t in (w.get("topics") or [])[:3] if t.get("id")]
    except (HTTPException, json.JSONDecodeError, KeyError):
        items = []                       # 이 논문을 해석 못 함 — 빈 목록으로 캐시, 매번 다시 묻지 않는다
    # URLError·HTTPError 는 여기서 잡지 않는다 — 캐시하면 '못 물어본 것'이 영구 오답으로 굳는다
    write_json(d / "topics.json", {"items": items, "ts": int(time.time())})
    return items


def _top_topics(counts: dict, n: int = TREND_TOPICS) -> list:
    """(id, name, 논문 수) 상위 n개. 동률이면 이름순 — 같은 서재면 늘 같은 결과가 나와야 한다."""
    return sorted(counts.values(), key=lambda t: (-t["papers"], t["name"]))[:n]


@app.get("/api/library/trends")
def library_trends(days: int = 90, refresh: int = 0):
    days = max(7, min(730, int(days or 90)))
    if not refresh:
        c = read_json(TRENDS, None)
        if c and c.get("days") == days and int(time.time()) - int(c.get("ts") or 0) < 86400:
            return {**c, "cached": True}
    counts = {}
    # 주제 수집(_paper_topics)도 같은 try 안에 둔다 — 네트워크 실패를 '주제 0편'으로 오인해
    # 사용자에게 엉뚱한 숙제(서지정보 추출)를 시키던 자리다
    try:
        for d in sorted(LIB.iterdir()):
            if not d.is_dir() or not (d / "meta.json").exists():
                continue
            for t in _paper_topics(d):
                e = counts.setdefault(t["id"], {"id": t["id"], "name": t["name"], "papers": 0})
                e["papers"] += 1
        tops = _top_topics(counts)
        if not tops:
            raise HTTPException(404, "주제를 뽑은 논문이 0편입니다 — DOI 나 제목이 있는 논문이 필요합니다")
        since = time.strftime("%Y-%m-%d", time.gmtime(time.time() - days * 86400))
        ids = "|".join(t["id"] for t in tops)
        r = _oa_get(f"works?filter=topics.id:{ids},from_publication_date:{since}"
                    f"&sort=cited_by_count:desc&per-page=20&select={OA_FIELDS}")
    except urllib.error.HTTPError as e:       # URLError 서브클래스 — 먼저 걸러야 상태코드가 안 사라진다
        raise HTTPException(502, f"OpenAlex 오류 HTTP {e.code} {e.reason}")
    except urllib.error.URLError as e:
        raise HTTPException(502, f"OpenAlex에 연결하지 못했습니다 — {e.reason}")
    except (json.JSONDecodeError, KeyError) as e:
        raise HTTPException(502, f"OpenAlex 응답을 해석하지 못했습니다 — {e}")
    data = {"days": days, "since": since, "topics": tops, "ts": int(time.time()),
            "total": (r.get("meta") or {}).get("count", 0),
            "results": _oa_clean(r.get("results") or [], _lib_index())}
    write_json(TRENDS, data)
    return data


# ---------------------------------------------------------------- selection actions

@app.post("/api/ai/selection")
async def ai_selection(req: Request):
    body = await req.json()
    action = body.get("action")
    sel = (body.get("text") or "").strip()
    ctx = (body.get("context") or "").strip()[:6000]
    title = (body.get("title") or "").strip()
    if not sel:
        raise HTTPException(400, "선택된 텍스트가 없습니다")
    if action == "translate":
        system = (
            "너는 학술 논문 전문 번역가다. 주어진 선택 구절을 자연스러운 한국어로 번역한다. "
            "전문용어는 '한국어(원어)' 병기. 수식 기호는 그대로 둔다. "
            "번역문만 출력한다 — 서두, 설명, 따옴표 금지."
        )
    elif action == "explain":
        system = (
            "너는 논문을 옆에서 설명해주는 튜터다. 선택된 구절(수식·도표 캡션·이론·주장)을 "
            "한국어 마크다운으로 설명한다. 수식이면 각 항의 의미를, 개념이면 직관을 먼저. "
            "주변 맥락을 참고하되 3~8문장 정도로 간결하게. 서두/맺음말 금지."
        )
    else:
        raise HTTPException(400, f"알 수 없는 action: {action!r}")
    prompt = f"논문 제목: {title}\n\n===== 주변 맥락 =====\n{ctx}\n\n===== 선택 구절 =====\n{sel}"
    return {"result": await ask_claude_async(prompt, system, timeout=180)}


# ---------------------------------------------------------------- config

CONFIG = ROOT / "config.json"
CONFIG_DEFAULTS = {
    "obsidian_vault_path": "",
    "obsidian_subfolder": "Achird",
    "zotero_data_dir": str(Path.home() / "Zotero"),
    "export_translation": True,
    "export_chat": True,
    "cite_citekey": False,   # 쓰기 탭 인용을 (저자, 연도, p.N) 대신 [@citekey, p.N] 로
    "auto_prep": True,       # 논문 유입 즉시 백그라운드 준비 체인(번역 제외) 자동 실행
}


def load_config() -> dict:
    return {**CONFIG_DEFAULTS, **read_json(CONFIG, {})}


def _rehome(p: str, home: Path = None) -> str:
    """설정에 저장된 절대경로를 이 PC의 홈 기준으로 되짚는다(순수 함수).

    config.json은 OneDrive로 두 대의 PC가 함께 쓰는데 사용자명이 다르면
    C:\\Users\\A\\... 는 B의 PC에 존재하지 않는다. 홈 바로 아래 사용자 폴더 이름만
    갈아끼워 같은 꼬리 경로를 이 PC 홈에서 찾아본다.

    저장값은 건드리지 않고 쓰는 순간에만 되짚는다 — 되짚은 값을 되쓰면 이번엔 반대편
    PC가 깨진다. 실제로 존재할 때만 바꾸고, 없으면 원래 값을 그대로 돌려줘 오류 메시지가
    사용자가 설정한 경로를 가리키게 한다."""
    s = str(p or "").strip()
    if not s:
        return s
    raw = Path(s)
    if raw.is_dir():
        return s
    home = home or Path.home()
    try:
        tail = raw.relative_to(home.parent)          # <사용자명>\나머지
    except ValueError:
        return s
    if len(tail.parts) < 2:                          # 사용자 폴더 자체는 되짚을 꼬리가 없다
        return s
    cand = home.joinpath(*tail.parts[1:])
    return str(cand) if cand.is_dir() else s


@app.get("/api/config")
def get_config():
    return load_config()


@app.put("/api/config")
async def put_config(req: Request):
    cfg = load_config()
    body = await req.json()
    for k in CONFIG_DEFAULTS:
        if k in body:
            if not isinstance(body[k], type(CONFIG_DEFAULTS[k])):
                raise HTTPException(400, f"{k}: {type(CONFIG_DEFAULTS[k]).__name__} 타입이어야 합니다")
            cfg[k] = body[k]
    for key in ("obsidian_vault_path", "zotero_data_dir"):     # 신뢰경계: 실재하는 폴더만
        if key not in body:
            continue          # 안 건드린 값 때문에 저장이 막히면 안 된다 — Zotero를 안 쓰는데
                              # 낡은 Zotero 경로가 남아 있다고 볼트 경로를 못 바꾸는 일이 생긴다
        v = str(cfg.get(key, "")).strip()
        if v and not Path(_rehome(v)).is_dir():   # 반대편 PC의 경로도 되짚어 보고 판정
            raise HTTPException(400, f"경로가 존재하는 폴더가 아닙니다: {v}")
    write_json(CONFIG, cfg)
    return cfg


# ---------------------------------------------------------------- Obsidian export

AUTO_START = "<!-- achird:auto:start -->"
AUTO_END = "<!-- achird:auto:end -->"
_OLD_START = "<!-- moonlight:auto:start -->"   # 구 이름(Moonlight)으로 내보낸 노트도 우리 것으로 인식
_OLD_END = "<!-- moonlight:auto:end -->"


_WIN_RESERVED = {"CON", "PRN", "AUX", "NUL",
                 *(f"COM{i}" for i in range(1, 10)), *(f"LPT{i}" for i in range(1, 10))}


def safe_filename(title: str, pid: str) -> str:
    """파일명 안전화: 금지문자·제어문자 제거, 공백 정리, 120자, Windows 예약명 회피, 빈 결과는 pid."""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", title or "")
    name = re.sub(r"\s+", " ", name).strip().strip(".")
    name = name[:80].strip()      # Windows MAX_PATH 260: 볼트경로+attachments/+'-'+pid+'.pdf' 여유 확보
    if not name:
        return pid
    if name.upper() in _WIN_RESERVED:      # CON/NUL 등은 콘솔·널 디바이스 I/O로 행·데이터유실 유발
        return name + "_"
    return name


def _yaml_str(s) -> str:
    s = str(s).replace("\r", " ").replace("\n", " ")   # 개행은 frontmatter를 여러 줄로 찢는다
    return '"' + s.replace("\\", "\\\\").replace('"', '\\"') + '"'


def _atomic_write_text(path: Path, text: str) -> None:
    tmp = path.with_suffix(path.suffix + ".tmp")     # OneDrive 동기화 중 부분파일 방지
    tmp.write_text(text, encoding="utf-8")
    tmp.replace(path)


def merge_note(path: Path, managed: str) -> str:
    """AUTO_END 마커 뒤(사용자 영역)를 보존하고 앞을 managed로 교체. 신규는 '내 메모' 스텁 추가."""
    if path.exists():
        old = path.read_text(encoding="utf-8", errors="replace")
        for end in (AUTO_END, _OLD_END):
            i = old.rfind(end)     # 마지막 마커=관리블록 끝. 본문(번역·대화)에 마커 리터럴 섞여도 안전
            if i != -1:
                return managed + old[i + len(end):]
    return managed + "\n\n## 내 메모\n\n여기에 쓴 내용은 재내보내기해도 보존됩니다.\n"


def _note_owned_by_other(head: str, pid: str) -> bool:
    """기존 노트 앞부분이 사용자 소유(마커 없음)이거나 다른 논문 소유면 True → 파일명 충돌 회피."""
    owner = re.search(r"^paper_id:\s*(\S+)", head, re.M)
    ours = any(s in head for s in (AUTO_START, _OLD_START, "source: achird", "source: moonlight"))
    return (not ours) or (owner is not None and owner.group(1) != pid)


def copy_pdf_to_vault(pid: str, out_dir: Path, fname: str) -> str:
    """PDF를 볼트 attachments/ 로 원자적 복사(이미 있으면 스킵). 노트 폴더 기준 상대경로 반환."""
    rel = f"attachments/{fname}-{pid}.pdf"
    src = paper_dir(pid) / "paper.pdf"
    dst = out_dir / "attachments" / f"{fname}-{pid}.pdf"
    if src.exists() and not dst.exists():
        dst.parent.mkdir(parents=True, exist_ok=True)
        tmp = dst.with_suffix(".pdf.tmp")
        shutil.copy2(src, tmp)
        tmp.replace(dst)
    return rel


def _hl_lines(d: Path) -> str:
    items = read_json(d / "highlights.json", {"items": []}).get("items", [])
    if not items:
        return ""
    out = ["## 핵심 하이라이트", ""]
    for h in items:
        txt = re.sub(r"\s+", " ", str(h.get("text", ""))).strip().replace("==", "")
        if not txt:
            continue
        tag = {"ai": "AI", "zotero": "Zotero"}.get(h.get("source"), "내 강조")
        page = f"{h['page']}쪽 · " if h.get("page") else ""
        out.append(f"- =={txt}== — {page}{tag}")
        if h.get("reason"):
            out.append(f"  - {h['reason']}")
    out.append("")
    return "\n".join(out)


def _note_lines(d: Path) -> str:
    notes = read_json(d / "notes.json", [])
    if not notes:
        return ""
    out = ["## 내 노트", ""]
    for n in notes:
        page = f"**{n['page']}쪽** — " if n.get("page") else ""
        quote = str(n.get("quote", "")).strip().replace("\n", " ")
        out.append(f"- {page}> {quote}")
        memo = str(n.get("memo", "")).strip()
        out.append(f"  - {memo}" if memo else "  - _(메모 없음)_")
    out.append("")
    return "\n".join(out)


def _refs_lines(d: Path) -> str:
    refs = read_json(d / "refs.json", {"items": []}).get("items", [])
    if not refs:
        return ""
    out = ["## 참고문헌", ""]
    for r in refs:
        out.append(f"- [{r.get('n', '')}] {str(r.get('text', '')).strip()}")
    out.append("")
    return "\n".join(out)


def _translation_lines(d: Path) -> str:
    tr = read_json(d / "translation.json", {"pages": {}}).get("pages", {})
    if not tr:
        return ""
    out = ["## 전체 번역", ""]
    for n in sorted(tr, key=lambda x: int(x)):
        body = str(tr[n]).strip()
        if body:
            out.append(f"### {n}쪽\n\n{body}\n")
    out.append("")
    return "\n".join(out)


def _chat_lines(d: Path) -> str:
    hist = read_json(d / "chat.json", [])
    if not hist:
        return ""
    out = ["## AI 대화", ""]
    for m in hist:
        content = str(m.get("content", "")).strip()
        if m.get("role") == "user":
            out.append(f"**질문:** {content}")
            if m.get("quote"):
                out.append(f"> {str(m['quote']).strip()}")
        else:
            out.append(f"\n**답변:** {content}\n")
    out.append("")
    return "\n".join(out)


def note_markdown(pid: str, cfg: dict, pdf_rel: str) -> str:
    """frontmatter + AUTO 관리블록 텍스트를 만든다 (사용자 영역은 merge_note가 이어붙임)."""
    d = paper_dir(pid)
    meta = read_json(d / "meta.json", {})
    title = meta.get("title") or pid
    citekey = (meta.get("citekey") or "").strip()
    added = time.strftime("%Y-%m-%d", time.localtime(meta.get("added") or time.time()))

    # @citekey 를 alias 로 — vault 어디서든 [[@citekey]] 위키링크가 이 노트로 해석된다
    # (초안의 [@citekey] 인용, 다른 노트의 관련 링크가 전부 살아난다)
    aliases = [_yaml_str(title)] + ([_yaml_str("@" + citekey)] if citekey else [])
    fm = ["---",
          f"title: {_yaml_str(title)}",
          f"aliases: [{', '.join(aliases)}]",
          f"paper_id: {pid}",
          "source: achird"]
    if citekey:
        fm.append(f"zotero: {_yaml_str('@' + citekey)}")
    if meta.get("authors"):
        fm.append("authors: [" + ", ".join(_yaml_str(a) for a in meta["authors"][:20]) + "]")
    if meta.get("year"):
        fm.append(f"year: {int(meta['year'])}")
    if meta.get("venue"):
        fm.append(f"venue: {_yaml_str(meta['venue'])}")
    if meta.get("doi"):
        fm.append(f"doi: {_yaml_str(meta['doi'])}")
    fm += [f"added: {added}",
           f"pages: {meta.get('pages', 0)}"]
    if pdf_rel:  # export_pdf=false 면 볼트에 PDF 사본이 없다 — 원본은 Zotero/Achird 링크로
        fm.append(f"pdf_vault: {_yaml_str(cfg.get('obsidian_subfolder', 'Achird') + '/' + pdf_rel)}")
    fm += [f"pdf_url: {_yaml_str(f'{URL}/api/papers/{pid}/pdf')}",
           "tags: [paper]",
           "---"]

    rel_link = f"관련: [[@{citekey}]] · " if citekey else ""
    # 삼각 순환: 노트(Obsidian) → 리더(Achird) → 서지관리자(Zotero). 링크가 죽는 경우도
    # 자명하다 — Achird 링크는 서버가 꺼져 있으면, Zotero 링크는 Zotero 가 없으면 안 열린다.
    links = rel_link + (f"[📄 PDF 열기](<{pdf_rel}>) · " if pdf_rel else "")
    links += f"[Achird에서 열기]({URL}/#paper={pid})"
    if meta.get("zotero_key"):
        links += f" · [Zotero에서 열기](zotero://select/library/items/{meta['zotero_key']})"
    body = [AUTO_START,
            f"> [!info] Achird 자동 생성 · {added} 갱신. 이 블록은 재내보내기 시 덮어써집니다.",
            "> 개인 메모는 맨 아래 \"## 내 메모\"에 작성하세요 (보존됩니다).",
            "",
            f"# {title}",
            links,
            ""]

    if meta.get("abstract"):        # Zotero 가져오기가 흡수한 초록 — vault 검색에 걸리게 한다
        body += ["## 초록", "", str(meta["abstract"]).strip(), ""]
    summary = (d / "summary.md")
    if summary.exists():
        body += ["## 요약", "", summary.read_text(encoding="utf-8").strip(), ""]
    for section in (_hl_lines(d), _note_lines(d), _refs_lines(d)):
        if section:
            body.append(section)
    if cfg.get("export_translation"):
        t = _translation_lines(d)
        if t:
            body.append(t)
    if cfg.get("export_chat"):
        c = _chat_lines(d)
        if c:
            body.append(c)
    body.append(AUTO_END)
    return "\n".join(fm) + "\n" + "\n".join(body)


# ---- SRS 복습 카드: 내 강조·노트를 Haiku 가 Q/A 카드로 만들어 vault 의 flashcards/ 에 쓴다.
# obsidian-spaced-repetition 플러그인 형식(질문 / ? / 답). 노트가 아니라 별도 파일인 이유:
# 플러그인이 복습 성적을 카드 옆에 <!--SR:...--> 주석으로 박는데, 노트 관리블록 안이면
# 재내보내기마다 스케줄이 날아간다. 카드 파일은 내용이 바뀔 때만 다시 쓴다(스케줄 보존).

def _flash_sig(d: Path) -> list:
    return [_mtime(d / "highlights.json"), _mtime(d / "notes.json")]


def make_flashcards(pid: str) -> dict:
    """{"cards": [...], "changed": bool}. 강조·노트가 없으면 카드 0장 — 카드는 '내가 표시한
    것'에서만 나온다. 캐시(flashcards.json)는 소스 mtime 서명으로 무효화한다."""
    d = paper_dir(pid)
    sig = _flash_sig(d)
    cached = read_json(d / "flashcards.json", None)
    if cached and cached.get("sig") == sig:
        return {"cards": cached.get("cards") or [], "changed": False}
    blocks = [b for b in (_hl_lines(d), _note_lines(d)) if b]
    if not blocks:
        return {"cards": [], "changed": False}
    meta = read_json(d / "meta.json", {})
    system = (
        "너는 논문 독자의 하이라이트·노트로 복습 플래시카드를 만드는 도구다. "
        "제공된 발췌문에 있는 사실만 쓰고, 논문에 대한 일반 지식이나 추측을 보태지 마라. "
        "질문은 능동 회상을 유도하게(정의·수치·이유·비교), 답은 1~3문장으로 짧게. 한국어로. "
        '출력은 JSON 배열 하나만: [{"q": "질문", "a": "답"}] — 5~8장. 다른 텍스트 절대 금지.'
    )
    raw = ask_claude(f"논문 제목: {meta.get('title', '')}\n\n" + "\n\n".join(blocks),
                     system, model="haiku")
    try:
        parsed = extract_json(raw)
    except (ValueError, json.JSONDecodeError):
        raise HTTPException(502, f"카드 생성 응답을 해석하지 못했습니다: {raw[:120]!r}")
    cards = [{"q": str(c.get("q", "")).strip(), "a": str(c.get("a", "")).strip()}
             for c in (parsed if isinstance(parsed, list) else [])
             if isinstance(c, dict) and str(c.get("q", "")).strip() and str(c.get("a", "")).strip()][:12]
    write_json(d / "flashcards.json", {"cards": cards, "sig": sig, "ts": int(time.time())})
    return {"cards": cards, "changed": True}


def write_flashcard_file(pid: str, out_dir: Path) -> Path | None:
    """카드를 vault 의 flashcards/@citekey.md 로. 카드가 안 바뀌었고 파일이 있으면 안 만진다 —
    파일 안의 <!--SR:...--> 복습 스케줄을 보존하는 조건이다."""
    r = make_flashcards(pid)
    if not r["cards"]:
        return None
    d = paper_dir(pid)
    meta = read_json(d / "meta.json", {})
    citekey = str(meta.get("citekey") or "").strip()
    fname = f"@{citekey}.md" if citekey else f"{pid}.md"
    path = out_dir / "flashcards" / fname
    if path.exists() and not r["changed"]:
        return path
    note_name = safe_filename(meta.get("title", pid), pid)
    lines = ["---", f"paper_id: {pid}", "tags: [flashcards]", "---", "",
             f"출처: [[{note_name}]]",
             "", "> [!info] Achird 자동 생성. 강조·노트가 바뀌면 다시 생성되며 그때 복습 스케줄이 초기화됩니다.", ""]
    for c in r["cards"]:
        lines += [c["q"], "?", c["a"], ""]
    path.parent.mkdir(parents=True, exist_ok=True)
    _atomic_write_text(path, "\n".join(lines))
    return path


# 증분 내보내기 상태 — {pid: 마지막 내보내기 시각}. meta.json 에 스탬프를 쓰면 그 쓰기가
# meta.json 의 mtime 을 올려 "항상 변경됨"이 되므로 논문 폴더 밖 별도 파일에 둔다.
EXPORT_STAMP = LIB / "_obsidian_export.json"

# note_markdown 이 실제로 읽는 파일들 — 이 중 하나라도 스탬프보다 새로우면 "변경됨"
_EXPORT_SRC = ("meta.json", "summary.md", "highlights.json", "notes.json",
               "refs.json", "translation.json", "chat.json")


def _export_mtime(d: Path) -> float:
    return max((_mtime(d / f) for f in _EXPORT_SRC), default=0.0)


def export_obsidian_one(pid: str, cfg: dict) -> Path:
    vault = _rehome(cfg.get("obsidian_vault_path", ""))
    if not vault or not Path(vault).is_dir():
        raise HTTPException(400, "먼저 Obsidian 볼트 경로를 설정하세요 (내보내기 탭)")
    d = paper_dir(pid)
    meta = read_json(d / "meta.json", {})
    sub = str(cfg.get("obsidian_subfolder", "Achird")).strip() or "Achird"
    out_dir = Path(vault) / sub
    out_dir.mkdir(parents=True, exist_ok=True)
    fname = safe_filename(meta.get("title", pid), pid)
    path = out_dir / f"{fname}.md"
    if path.exists():                       # 사용자 소유 또는 다른 논문의 노트면 덮지 않고 pid 접미사로 회피
        head = path.read_text(encoding="utf-8", errors="replace")[:2000]
        if _note_owned_by_other(head, pid):
            path = out_dir / f"{fname} ({pid}).md"
    pdf_rel = copy_pdf_to_vault(pid, out_dir, fname) if cfg.get("export_pdf", False) else ""
    _atomic_write_text(path, merge_note(path, note_markdown(pid, cfg, pdf_rel)))
    try:                        # 카드 실패(AI 오류 등)가 노트 내보내기를 막으면 안 된다
        write_flashcard_file(pid, out_dir)
    except Exception as e:      # noqa: BLE001
        print(f"[flashcards] {pid}: {type(e).__name__}: {e}", file=sys.stderr)
    stamps = read_json(EXPORT_STAMP, {})
    stamps[pid] = time.time()
    write_json(EXPORT_STAMP, stamps)
    return path


def write_vault_bib(cfg: dict) -> int:
    """서재 전체 BibTeX 을 볼트의 노트 폴더에 achird.bib 로 유지한다. 내보내기 때마다 통째로
    재생성 — 서재 수백 편 규모에서 즉시 끝나고, 증분 관리보다 항상 정확하다. citekey 인용
    모드 + Obsidian citation 플러그인/pandoc 이 이 파일 하나로 완성된다."""
    vault = _rehome(cfg.get("obsidian_vault_path", ""))
    if not vault or not Path(vault).is_dir():
        return 0
    sub = str(cfg.get("obsidian_subfolder", "Achird")).strip() or "Achird"
    out_dir = Path(vault) / sub
    out_dir.mkdir(parents=True, exist_ok=True)
    entries = []
    for d in sorted(LIB.iterdir()):
        if not d.is_dir():
            continue
        meta = read_json(d / "meta.json", {})
        if meta.get("title") or meta.get("authors"):
            entries.append(_cite_strings(meta)["bibtex"])
    _atomic_write_text(out_dir / "achird.bib", "\n\n".join(entries) + ("\n" if entries else ""))
    return len(entries)


@app.post("/api/papers/{pid}/export/obsidian")
def export_one(pid: str):
    cfg = load_config()
    path = export_obsidian_one(pid, cfg)
    write_vault_bib(cfg)
    return {"path": str(path), "filename": path.name}


@app.post("/api/export/obsidian")
async def export_all(req: Request):
    """전체 또는 변경분만. changed_only 면 마지막 내보내기 이후 사이드카가 바뀐 논문만 민다 —
    내보내기가 습관이 되려면 '전부 다시'가 아니라 '바뀐 것만'이 기본 동작이어야 한다."""
    try:
        body = await req.json()
    except Exception:
        body = {}
    changed_only = bool(body.get("changed_only"))
    cfg = load_config()
    stamps = read_json(EXPORT_STAMP, {})
    files, failed, skipped = [], [], 0
    for d in sorted(LIB.iterdir()):
        if d.is_dir() and (d / "meta.json").exists():
            if changed_only and _export_mtime(d) <= float(stamps.get(d.name) or 0):
                skipped += 1
                continue
            try:
                files.append(export_obsidian_one(d.name, cfg).name)
            except HTTPException:
                raise                        # 볼트 미설정 등은 전체 중단
            except Exception as e:           # noqa: BLE001 — 개별 논문 실패는 건너뜀
                # 8자리 pid 로는 어느 논문인지 모르고, str(e) 가 빈 예외면 이유가 통째로 빈다
                failed.append(f"{read_json(d / 'meta.json', {}).get('title') or d.name}: "
                              f"{type(e).__name__}: {e}")
    bib = write_vault_bib(cfg)
    return {"count": len(files), "files": files, "failed": failed, "skipped": skipped, "bib": bib}


# ---------------------------------------------------------------- my summary (내가 표시한 것만 요약)

@app.get("/api/papers/{pid}/mysummary")
def get_mysummary(pid: str):
    return read_json(paper_dir(pid) / "mysummary.json", {})


@app.post("/api/papers/{pid}/mysummary")
async def make_mysummary(pid: str):
    """독자가 직접 남긴 하이라이트·노트'만'을 근거로 개인화 다이제스트를 만든다. 두 사이드카를
    합쳐 하나의 발췌문 블록으로 주고(_hl_lines/_note_lines — Obsidian 내보내기와 같은 포맷터를
    재사용), AI에게는 이 발췌문 밖 사실을 보태지 말라고 못박는다."""
    d = paper_dir(pid)
    hl_items = read_json(d / "highlights.json", {"items": []}).get("items", [])
    notes = read_json(d / "notes.json", [])
    if not hl_items and not notes:
        raise HTTPException(400, "하이라이트나 노트가 아직 없습니다. 먼저 논문을 읽으며 표시해주세요.")
    meta = read_json(d / "meta.json", {})
    blocks = [b for b in (_hl_lines(d), _note_lines(d)) if b]
    system = (
        "너는 독자가 직접 표시한 하이라이트와 노트만으로 개인화 다이제스트를 만드는 조교다. "
        "아래 제공된 발췌문(하이라이트·노트)에 없는 사실은 절대 추가하지 마라 — 이 논문에 대한 "
        "일반 지식이나 추측을 보태지 말고 오직 주어진 발췌문을 요약·연결·정리만 한다. 독자가 "
        "무엇을 중요하게 여겼는지(강조점)를 그대로 살려라. 한국어 마크다운으로 답하되, 제목(`#`)은 "
        "독자가 표시한 내용을 관통하는 주제로 직접 짓는다(논문 제목을 그대로 쓰지 말 것). "
        "구조: 제목 → 짧은 개요 한두 문장 → 표시한 내용을 주제별로 묶은 소단락들. "
        "서두·맺음말·메타 발언 금지."
    )
    prompt = f"논문 제목: {meta.get('title', '')}\n\n" + "\n\n".join(blocks)
    md = await ask_claude_async(prompt, system)
    data = {"markdown": md, "highlights": len(hl_items), "notes": len(notes), "ts": int(time.time())}
    write_json(d / "mysummary.json", data)
    return data


# ---------------------------------------------------------------- Zotero import

def zotero_db_copy(cfg: dict) -> Path:
    """zotero.sqlite(+wal/shm)를 임시 폴더로 복사 — Zotero 실행 중 락을 회피한다."""
    data_dir = Path(_rehome(cfg.get("zotero_data_dir", "")))
    src = data_dir / "zotero.sqlite"
    if not src.exists():
        raise HTTPException(400, f"zotero.sqlite를 찾을 수 없습니다: {src}")
    tmpdir = Path(tempfile.mkdtemp(prefix="ml_zot_"))
    for suf in ("", "-wal", "-shm"):
        s = data_dir / f"zotero.sqlite{suf}"
        if s.exists():
            shutil.copy2(s, tmpdir / s.name)
    return tmpdir / "zotero.sqlite"


def _zot_field(q, item_id, field_id):
    if not field_id:
        return None
    r = q.execute("SELECT v.value FROM itemData d JOIN itemDataValues v ON v.valueID=d.valueID "
                  "WHERE d.itemID=? AND d.fieldID=?", (item_id, field_id)).fetchone()
    return r[0] if r else None


def _zot_creators(q, item_id):
    """"Lastname, Firstname" 목록 — _split_author 가 그대로 해석하는 형태라
    meta.authors 에 넣으면 인용 문자열(_cite_short/_cite_strings)이 바로 정확해진다."""
    rows = q.execute("SELECT c.lastName, c.firstName FROM itemCreators ic "
                     "JOIN creators c ON c.creatorID=ic.creatorID "
                     "WHERE ic.itemID=? ORDER BY ic.orderIndex", (item_id,)).fetchall()
    names = []
    for r in rows:
        last, first = (r[0] or "").strip(), (r[1] or "").strip()
        name = f"{last}, {first}" if last and first else (last or first)
        if name:
            names.append(name)
    return names


def _zot_year(date_val) -> int | None:
    """Zotero date 필드값("2020-05-01", "May 2020", 멀티파트 문자열)에서 연도만.
    형식이 제각각이라 4자리 연도 첫 등장을 취하고 extract_metadata 와 같은 범위 검증을 건다."""
    m = re.search(r"\b(1[89]\d{2}|20\d{2}|2100)\b", str(date_val or ""))
    return int(m.group(0)) if m else None


def resolve_pdf(data_dir: Path, att_key, link_mode, path, base):
    """첨부 path → 실제 PDF 절대경로. 해석 불가면 None."""
    if not path:
        return None
    if path.startswith("storage:"):                     # imported: dataDir/storage/<key>/file
        p = data_dir / "storage" / att_key / path[len("storage:"):]
        return p if p.exists() else None
    if link_mode == 2:                                  # linked_file
        if path.startswith("attachments:"):
            if not base:
                return None                             # baseAttachmentPath 미설정 → 미해석
            p = Path(base) / path[len("attachments:"):]
        else:
            p = Path(path)
        return p if p.exists() else None
    return None


def scan_zotero(cfg: dict) -> list:
    dbpath = zotero_db_copy(cfg)
    con = None
    try:
        con = sqlite3.connect(f"file:{dbpath}?mode=ro", uri=True)
        con.row_factory = sqlite3.Row
        q = con.cursor()
        deleted = {r[0] for r in q.execute("SELECT itemID FROM deletedItems")}
        field_ids = {r["fieldName"]: r["fieldID"] for r in
                     q.execute("SELECT fieldID, fieldName FROM fields")}
        title_fid, ck_fid = field_ids.get("title"), field_ids.get("citationKey")
        # 서지 전체 흡수 — Zotero 가 이미 검증해 둔 값을 그대로 쓰면 AI 재추출이 필요 없다
        doi_fid = field_ids.get("DOI")
        date_fid = field_ids.get("date")
        venue_fids = [field_ids.get(f) for f in
                      ("publicationTitle", "proceedingsTitle", "bookTitle") if field_ids.get(f)]
        abs_fid = field_ids.get("abstractNote")
        data_dir = Path(_rehome(cfg.get("zotero_data_dir", "")))
        base = cfg.get("zotero_base_attachment_path")   # 현재 미지원(None) — linked 절대경로만 해석
        atts = q.execute(
            "SELECT ia.itemID att_id, ia.parentItemID parent, ia.linkMode, ia.path, "
            "i.key att_key FROM itemAttachments ia JOIN items i ON i.itemID=ia.itemID "
            "WHERE ia.contentType='application/pdf'").fetchall()
        out = []
        for a in atts:
            parent = a["parent"]
            if a["att_id"] in deleted or (parent is not None and parent in deleted):
                continue
            pdf = resolve_pdf(data_dir, a["att_key"], a["linkMode"], a["path"], base)
            if parent is None:                  # 부모 없는 첨부(PDF만 드래그) → 첨부 자체를 항목으로
                zkey = None
                title = _zot_field(q, a["att_id"], title_fid) or (Path(str(pdf)).name if pdf else "(제목 없음)")
                creators, ck = [], None
                doi = year = venue = abstract = None
            else:
                pk = q.execute("SELECT key FROM items WHERE itemID=?", (parent,)).fetchone()
                zkey = pk[0] if pk else None
                title = _zot_field(q, parent, title_fid) or "(제목 없음)"
                creators = _zot_creators(q, parent)
                ck = _zot_field(q, parent, ck_fid) if ck_fid else None
                doi = str(_zot_field(q, parent, doi_fid) or "").strip()[:200] or None
                year = _zot_year(_zot_field(q, parent, date_fid))
                venue = next((v for v in (str(_zot_field(q, parent, f) or "").strip()
                                          for f in venue_fids) if v), None)
                abstract = str(_zot_field(q, parent, abs_fid) or "").strip()[:5000] or None
            out.append({
                "att_key": a["att_key"],        # 첨부 단위 식별 — 한 논문에 PDF 여러 개일 때 구분
                "zotero_key": zkey,
                "title": title,
                "creators": creators,
                "citekey": ck,
                "doi": doi,
                "year": year,
                "venue": venue,
                "abstract": abstract,
                "pdf": str(pdf) if pdf else None,
                "pdf_name": Path(str(pdf)).name if pdf else None,
                "resolved": pdf is not None,
            })
        return out
    finally:
        if con:
            con.close()
        shutil.rmtree(dbpath.parent, ignore_errors=True)


# ---- 스캔 캐시: 대기열 매칭·신착 배지·주석 감지는 화면을 열 때마다 돈다 — 매번 수십 MB
# sqlite 를 복사하면 느리다. zotero.sqlite 의 (경로, mtime, size)가 그대로면 재사용한다.
# Zotero 가 뭐든 바꾸면 mtime 이 바뀌므로 신선도는 파일시스템이 보장한다.

_ZOT_SCAN_CACHE = {"sig": None, "items": []}
_ZOT_ANN_CACHE = {}          # att_key -> {"sig": sig, "items": [...]}


def _zot_sqlite_sig(cfg: dict):
    src = Path(_rehome(cfg.get("zotero_data_dir", ""))) / "zotero.sqlite"
    try:
        st = src.stat()
        return (str(src), st.st_mtime, st.st_size)
    except OSError:
        return None


def scan_zotero_cached(cfg: dict) -> list:
    sig = _zot_sqlite_sig(cfg)
    if sig is None:
        return []
    if _ZOT_SCAN_CACHE["sig"] != sig:
        _ZOT_SCAN_CACHE.update(sig=sig, items=scan_zotero(cfg))
    return _ZOT_SCAN_CACHE["items"]


def scan_zotero_annotations_cached(att_key: str, cfg: dict) -> list:
    sig = _zot_sqlite_sig(cfg)
    if sig is None:
        return []
    e = _ZOT_ANN_CACHE.get(att_key)
    if not e or e["sig"] != sig:
        e = {"sig": sig, "items": scan_zotero_annotations(att_key, cfg)}
        _ZOT_ANN_CACHE[att_key] = e
    return e["items"]


def _lib_zotero_att_keys() -> set:
    return {ak for d in LIB.iterdir() if d.is_dir()
            for ak in [read_json(d / "meta.json", {}).get("zotero_att_key")] if ak}


def import_zotero(att_keys: list, cfg: dict) -> dict:
    scanned = {s["att_key"]: s for s in scan_zotero(cfg) if s["att_key"]}
    existing = set()                          # 첨부 단위로 중복 판정 (부모 단위면 v2가 v1을 막음)
    for d in LIB.iterdir():
        if d.is_dir():
            ak = read_json(d / "meta.json", {}).get("zotero_att_key")
            if ak:
                existing.add(ak)
    hashes = existing_pdf_hashes()            # {sha256: pid} — 업로드본 포함 내용중복 차단
    imported, skipped = [], []
    for k in att_keys:
        s = scanned.get(k)
        if not s or not s.get("resolved") or not s.get("pdf"):
            skipped.append({"key": k, "reason": "PDF 미해석"}); continue
        if k in existing:
            skipped.append({"key": k, "reason": "이미 있음"}); continue
        src = Path(s["pdf"])
        if not src.exists():
            skipped.append({"key": k, "reason": "PDF 파일 없음"}); continue
        h = file_sha256(src)
        if h in hashes:
            skipped.append({"key": k, "reason": "이미 있음(내용 동일)"}); continue
        pid = uuid.uuid4().hex[:8]
        nd = LIB / pid
        nd.mkdir()
        shutil.copy2(src, nd / "paper.pdf")
        meta = {
            "id": pid, "title": s["title"], "filename": src.name,
            "added": int(time.time()), "pages": 0,
            "zotero_key": s.get("zotero_key"), "zotero_att_key": k,
            "citekey": s.get("citekey"), "pdf_sha256": h,
        }
        # Zotero 가 이미 갖고 있는 서지는 그대로 흡수 — haiku 재추출 없이 인용·검증·트렌드가 바로 돈다.
        # creators 는 "Last, First" 형태라 meta.authors 규약(_split_author)과 호환.
        for src_key, dst_key in (("creators", "authors"), ("year", "year"),
                                 ("venue", "venue"), ("doi", "doi"), ("abstract", "abstract")):
            if s.get(src_key):
                meta[dst_key] = s[src_key]
        write_json(nd / "meta.json", meta)
        existing.add(k)
        hashes[h] = pid
        imported.append({"pid": pid, "title": s["title"]})
    _auto_prep([p["pid"] for p in imported])
    return {"imported": imported, "skipped": skipped}


# ---- achird → Zotero 저장: Zotero 7 이 로컬에 띄우는 커넥터 서버(127.0.0.1:23119)로
# 서지를 밀어 넣는다. 브라우저 커넥터가 쓰는 것과 같은 saveItems 엔드포인트라 별도 설정이
# 필요 없고, 항목은 Zotero 에서 현재 선택된 컬렉션에 담긴다. 쓰기지만 zotero.sqlite 를
# 직접 만지지 않는다 — 파일 직접 쓰기는 락·스키마 위험 때문에 하지 않는다.

ZOTERO_CONNECTOR = "http://127.0.0.1:23119"


def _zot_save_item(body: dict) -> dict:
    """요청 본문 → 커넥터 saveItems 페이로드의 item 하나(순수 함수)."""
    authors = []
    for a in (body.get("authors") or [])[:20]:
        last, _ini = _split_author(str(a))
        given = ""
        s = str(a).strip()
        if "," in s:
            given = s.split(",", 1)[1].strip()
        elif len(s.split()) > 1:
            given = " ".join(s.split()[:-1])
        if last:
            authors.append({"lastName": last, "firstName": given, "creatorType": "author"})
    item = {"itemType": "journalArticle", "title": str(body.get("title") or "").strip()}
    if authors:
        item["creators"] = authors
    year = _cite_year(body)
    if year:
        item["date"] = year
    for src, dst in (("venue", "publicationTitle"), ("doi", "DOI"),
                     ("url", "url"), ("abstract", "abstractNote")):
        v = str(body.get(src) or "").strip()
        if v:
            item[dst] = v
    return item


@app.post("/api/zotero/save")
def zotero_save(body: dict):
    item = _zot_save_item(body)
    if not item["title"]:
        raise HTTPException(400, "제목이 없는 항목은 저장할 수 없습니다")
    payload = json.dumps({"items": [item], "sessionID": uuid.uuid4().hex,
                          "uri": str(body.get("url") or URL)}).encode("utf-8")
    req = urllib.request.Request(
        f"{ZOTERO_CONNECTOR}/connector/saveItems", data=payload, method="POST",
        headers={"Content-Type": "application/json",
                 "X-Zotero-Connector-API-Version": "2"})
    try:
        with urllib.request.urlopen(req, timeout=10) as r:
            r.read()
    except urllib.error.HTTPError as e:
        raise HTTPException(502, f"Zotero가 저장을 거부했습니다 (HTTP {e.code}). "
                                 f"Zotero를 최신 버전으로 업데이트해보세요.")
    except (urllib.error.URLError, OSError):
        raise HTTPException(502, "Zotero에 연결할 수 없습니다 — Zotero 데스크톱 앱이 실행 중인지 확인하세요.")
    return {"ok": True, "title": item["title"]}


@app.get("/api/import/zotero")
def import_scan():
    """가져오기 대화상자용 — 이미 서재에 들어온 첨부는 목록에서 뺀다. '가져오기 → 1편 skip'
    보다 애초에 추천이 안 뜨는 게 맞다. 두 겹으로 거른다: ① 첨부 키(Zotero로 가져온 것)
    ② PDF 내용 해시(같은 파일을 손으로 업로드해 둔 것). DOI만 같은 다른 판본(프리프린트 등)은
    남긴다 — 근접 중복은 '막지 않고 알린다'가 앱 전체 규칙이다.
    (대기열 매칭 등 내부 호출은 scan_zotero 원본을 쓴다.)"""
    have = _lib_zotero_att_keys()
    have_sha = existing_pdf_hashes()
    out = []
    for s in scan_zotero(load_config()):
        if s.get("att_key") in have:
            continue
        if s.get("resolved") and s.get("pdf"):
            try:
                if file_sha256(Path(s["pdf"])) in have_sha:
                    continue
            except OSError:
                pass                    # 해시 실패면 목록에 남긴다 — 가져오기 단계가 다시 거른다
        out.append(s)
    return out


@app.get("/api/import/zotero/badge")
def zotero_badge():
    """아직 안 가져온 Zotero 항목 수 — 홈의 가져오기 버튼에 붙는 신착 배지.
    Zotero 미사용·경로 깨짐이면 조용히 0 — 배지는 있어도 그만인 정보라 실패로 시끄러우면 안 된다."""
    try:
        items = scan_zotero_cached(load_config())
        if not items:
            return {"new": 0}
        existing = _lib_zotero_att_keys()
        return {"new": sum(1 for s in items
                           if s.get("resolved") and s.get("att_key") not in existing)}
    except Exception:                # noqa: BLE001
        return {"new": 0}


@app.get("/api/papers/{pid}/zotero-annotations/new")
def zotero_annotations_new(pid: str):
    """이 논문에 아직 안 가져온 Zotero 주석 수. 강조 탭 버튼에 배지로 붙는다."""
    d = paper_dir(pid)
    att_key = read_json(d / "meta.json", {}).get("zotero_att_key")
    if not att_key:
        return {"new": 0}
    try:
        anns = scan_zotero_annotations_cached(att_key, load_config())
    except Exception:                # noqa: BLE001
        return {"new": 0}
    seen = {_norm_hl_text(h.get("text")) for h in
            read_json(d / "highlights.json", {"items": []}).get("items", [])}
    return {"new": sum(1 for a in anns if _norm_hl_text(a["text"]) not in seen)}


@app.post("/api/import/zotero")
async def import_do(req: Request):
    body = await req.json()
    return import_zotero(body.get("keys") or [], load_config())


# ---- Zotero 주석 단방향 가져오기: Zotero PDF 리더에서 그은 형광펜·밑줄을 achird 강조로.
# 역방향(achird → Zotero 주석)은 하지 않는다 — zotero.sqlite 쓰기는 락·스키마 위험이고,
# 커넥터 API 는 주석 쓰기를 지원하지 않는다.

def _norm_hl_text(s) -> str:
    """중복 판정용 정규화 — 프런트 normQuery 와 같은 규칙(소문자·합자 해체·공백/하이픈 제거).
    PDF 추출 아티팩트(줄바꿈 하이픈, 리가처)가 달라도 같은 문장이면 같은 키가 나온다."""
    return re.sub(r"[\s­​‐‑‒–—-]", "",
                  str(s or "").lower().replace("ﬁ", "fi").replace("ﬂ", "fl"))


def _zot_ann_page(position, page_label) -> int:
    """position JSON 의 pageIndex(0기준) 우선, 없으면 pageLabel 숫자. 못 구하면 0."""
    try:
        idx = json.loads(position or "{}").get("pageIndex")
        if isinstance(idx, int) and idx >= 0:
            return idx + 1
    except (ValueError, TypeError):
        pass
    m = re.match(r"\s*(\d{1,4})", str(page_label or ""))
    return int(m.group(1)) if m else 0


def scan_zotero_annotations(att_key: str, cfg: dict) -> list:
    """첨부 하나의 하이라이트(1)·밑줄(5) 주석. 텍스트 없는 이미지·잉크 주석은 대상이 아니다."""
    dbpath = zotero_db_copy(cfg)
    con = None
    try:
        con = sqlite3.connect(f"file:{dbpath}?mode=ro", uri=True)
        q = con.cursor()
        deleted = {r[0] for r in q.execute("SELECT itemID FROM deletedItems")}
        rows = q.execute(
            "SELECT ia.itemID, ia.text, ia.comment, ia.color, ia.pageLabel, ia.position "
            "FROM itemAnnotations ia "
            "JOIN items att ON att.itemID = ia.parentItemID "
            "WHERE att.key = ? AND ia.type IN (1, 5)", (att_key,)).fetchall()
        out = []
        for item_id, text, comment, color, page_label, position in rows:
            text = re.sub(r"\s+", " ", str(text or "")).strip()
            if item_id in deleted or not text:
                continue
            out.append({"text": text[:1200],
                        "comment": str(comment or "").strip(),
                        "color": str(color or ""),
                        "page": _zot_ann_page(position, page_label)})
        return out
    finally:
        if con:
            con.close()
        shutil.rmtree(dbpath.parent, ignore_errors=True)


@app.post("/api/papers/{pid}/import/zotero-annotations")
def import_zotero_annotations(pid: str):
    d = paper_dir(pid)
    meta = read_json(d / "meta.json", {})
    att_key = meta.get("zotero_att_key")
    if not att_key:
        raise HTTPException(400, "Zotero에서 가져온 논문이 아닙니다 — 주석을 연결할 첨부를 모릅니다")
    anns = scan_zotero_annotations(att_key, load_config())
    hl = read_json(d / "highlights.json", {"items": []})
    items = hl.get("items") or []
    seen = {_norm_hl_text(h.get("text")) for h in items}
    added = 0
    for a in anns:
        key = _norm_hl_text(a["text"])
        if not key or key in seen:
            continue
        item = {"id": uuid.uuid4().hex[:8], "page": a["page"], "text": a["text"],
                "source": "zotero"}
        if a["comment"]:
            item["reason"] = a["comment"]
        if a["color"]:
            item["zcolor"] = a["color"]
        items.append(item)
        seen.add(key)
        added += 1
    if added:
        write_json(d / "highlights.json", {"items": items})
    return {"found": len(anns), "added": added, "skipped": len(anns) - added}


# ---------------------------------------------------------------- static & launch

app.mount("/static", StaticFiles(directory=ROOT / "static"), name="static")


@app.middleware("http")
async def static_no_cache(request: Request, call_next):
    """정적 파일은 매 로드마다 재검증(no-cache) — 코드 업데이트 후 사용자가 하드 리프레시
    (Ctrl+Shift+R)를 해야 하는 문제 제거. localhost라 304 재검증 비용은 무시할 수준."""
    resp = await call_next(request)
    if request.url.path.startswith("/static") or request.url.path == "/":
        resp.headers["Cache-Control"] = "no-cache"
    return resp


@app.get("/")
def index():
    return FileResponse(ROOT / "static" / "index.html")


def open_window():
    """Edge/Chrome 앱모드 창으로 열고, 없으면 기본 브라우저."""
    candidates = [
        os.path.expandvars(r"%ProgramFiles(x86)%\Microsoft\Edge\Application\msedge.exe"),
        os.path.expandvars(r"%ProgramFiles%\Microsoft\Edge\Application\msedge.exe"),
        os.path.expandvars(r"%ProgramFiles%\Google\Chrome\Application\chrome.exe"),
        os.path.expandvars(r"%LocalAppData%\Google\Chrome\Application\chrome.exe"),
    ]
    for exe in candidates:
        if os.path.isfile(exe):
            try:
                subprocess.Popen([exe, f"--app={URL}"], creationflags=CREATE_NO_WINDOW)
                return
            except OSError as e:      # Timer 스레드에서 죽으면 그 스레드만 조용히 끝나 원인이 사라진다
                print(f"브라우저 실행 실패({exe}): {type(e).__name__}: {e}")
    import webbrowser
    webbrowser.open(URL)


def _self_check():
    """회귀 테스트 실행 (python app.py --self-check).

    케이스는 tests/ 로 옮겼다 — 예전엔 assert 한 덩어리라 첫 실패가 나머지를 전부 가렸다.
    여기서는 러너에 위임만 하고, 추가할 케이스는 tests/test_pure.py 한 군데에만 쓴다."""
    import run_tests
    if not run_tests.run():
        raise SystemExit(1)
    print("self-check OK")


if __name__ == "__main__":
    if "--self-check" in sys.argv:
        _self_check()
        sys.exit(0)
    import socket
    with socket.socket() as _s:
        _s.settimeout(0.5)
        running = _s.connect_ex((HOST, PORT)) == 0
    if running:
        if "--no-browser" not in sys.argv:
            open_window()
            print(f"Achird Local — 이미 실행 중, 창만 엽니다 ({URL})")
        else:
            print(f"Achird Local — 이미 실행 중 ({URL})")
        sys.exit(0)
    if "--no-browser" not in sys.argv:
        threading.Timer(1.0, open_window).start()
    print(f"Achird Local — {URL}")
    uvicorn.run(app, host=HOST, port=PORT, log_level="warning")
